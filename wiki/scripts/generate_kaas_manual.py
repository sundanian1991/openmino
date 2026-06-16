#!/usr/bin/env python3
"""生成 KaaS 智能知识库中台操作手册 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

doc = Document()

# ============ 样式设置 ============
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Microsoft YaHei'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)


def add_screenshot_placeholder(doc, caption=""):
    """添加截图占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ 截图：{caption} ]")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(10)
    run.font.italic = True
    # 给占位符加个边框效果
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("┌─────────────────────────────────────────────┐\n│                                             │\n│              请在此处插入截图                │\n│                                             │\n└─────────────────────────────────────────────┘")
    r2.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    r2.font.size = Pt(9)
    r2.font.name = 'Courier New'


def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def add_code_block(doc, text):
    """添加代码块样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 背景色通过段落底纹实现
    shading = p._element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F5F5F5',
        qn('w:val'): 'clear'
    })
    shading.append(shd)


# ============ 封面 ============
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("KaaS 智能知识库中台")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("操作手册")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("适用对象：电销服务组全体同事\n版本：V1.0\n日期：2026-06-16\n编写：sundanian")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============ 目录页 ============
doc.add_heading('目录', level=1)
toc_items = [
    "一、KaaS 是什么",
    "二、快速开始：5 分钟上手",
    "三、项目管理详解",
    "  3.1 新建项目",
    "  3.2 编辑项目",
    "  3.3 删除项目",
    "  3.4 同步 INDEX",
    "  3.5 搜索与筛选",
    "四、文档批量录入详解",
    "  4.1 Step 1：解析目录",
    "  4.2 Step 2：选择项目",
    "  4.3 Step 3：摘要与匹配",
    "  4.4 Step 4：写入完成",
    "五、AI 如何调用知识库（MCP 工具）",
    "  5.1 可用工具一览",
    "  5.2 检索场景示例",
    "  5.3 在 Claude Code 中使用",
    "六、AI 辅助编写 Wiki",
    "  6.1 用 AI 生成初稿",
    "  6.2 用 AI 补充关键词和描述",
    "  6.3 用 AI 做质量审查",
    "七、知识库维护最佳实践",
    "  7.1 文档命名规范",
    "  7.2 关键词与描述维护",
    "  7.3 定期健康检查",
    "八、常见问题 FAQ",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ============ 一、KaaS 是什么 ============
doc.add_heading('一、KaaS 是什么', level=1)

doc.add_paragraph(
    'KaaS（Knowledge as a Service）是金融科技事业群的智能知识库中台。'
    '它不只是一个文档存储系统，而是一个能让 AI 直接检索和调用的结构化知识平台。'
)

doc.add_paragraph('核心能力：')
items = [
    '结构化存储：按一级域 → 项目 → 文档三层组织，每篇文档都有关键词和描述',
    'AI 智能路由：根据关键词和描述，AI 能自动判断查询属于哪个项目/目录',
    '批量录入：从 JoySpace 目录一键导入文档，自动摘要 + 智能匹配项目',
    'MCP 协议调用：AI 通过标准 MCP 工具直接检索知识库，无需人工中转',
    'Git 版本管理：文档变更自动同步到 Git，可追溯、可回滚',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_screenshot_placeholder(doc, "KaaS 首页概览")

doc.add_page_break()

# ============ 二、快速开始 ============
doc.add_heading('二、快速开始：5 分钟上手', level=1)

doc.add_heading('Step 1：打开项目管理', level=2)
doc.add_paragraph('访问 http://ai-kaas.pre-apps.jd.com/project-management')
add_screenshot_placeholder(doc, "项目管理页面入口")

doc.add_heading('Step 2：新建项目', level=2)
doc.add_paragraph('点击右上角 "+ 新建项目"，填写基本信息后保存。')
doc.add_paragraph('（详见第三章 3.1 节）')

doc.add_heading('Step 3：同步 INDEX', level=2)
doc.add_paragraph('点击 "同步" 按钮，勾选 "同步后 git push"，让项目信息生效。')

doc.add_heading('Step 4：录入文档', level=2)
doc.add_paragraph('进入 "文档批量录入"，粘贴 JoySpace 目录 URL，四步完成导入。')
doc.add_paragraph('（详见第四章）')

doc.add_heading('Step 5：AI 调用', level=2)
doc.add_paragraph('知识库建好后，AI 会自动通过 MCP 工具检索和调用文档。你只需要正常提问即可。')
doc.add_paragraph('（详见第五章）')

doc.add_page_break()

# ============ 三、项目管理详解 ============
doc.add_heading('三、项目管理详解', level=1)

doc.add_heading('3.1 新建项目', level=2)
doc.add_paragraph('点击右上角 "+ 新建项目"，右侧滑出表单面板。')

add_screenshot_placeholder(doc, "新建项目表单")

doc.add_paragraph('各字段说明：')
add_table(doc, ['字段', '必填', '说明'], [
    ['一级域', '是', '下拉选择。电销运营选 data-science'],
    ['数据来源', '是', 'topic_system 或 legacy_api。创建后不可修改'],
    ['原始 ID', '否', '勾选后出现。对接外部系统时填写；不填则默认与目录名相同'],
    ['文件目录名', '是', '纯小写字母+数字+下划线+中划线，须以字母开头。同域下不可重复'],
    ['完整名称', '是', '中文，至少 2 个汉字、4 个字符以上'],
    ['文档路径', '自动', '生成规则：domains/{一级域}/{文件目录名}'],
    ['TOPIC_KEYWORDS', '否', '短句描述项目内容，直接影响 AI 检索路由质量'],
    ['分类描述', '否', '较长语句说明项目范围，AI 自动分类时参考此字段'],
    ['知识库 ID', '否', '高级选项。如 ds-tmk-ops'],
    ['层级', '否', '高级选项。数字，默认 1'],
    ['文档所有者', '否', '高级选项。ERP 账号，多个用逗号分隔'],
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('重要提示：')
run.font.bold = True
items = [
    'TOPIC_KEYWORDS 和分类描述决定 AI 检索质量，务必认真填写',
    '文件目录名一旦创建，数据来源不可修改',
    '同一一级域下目录名不可重复，系统会自动校验',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 编辑项目', level=2)
doc.add_paragraph('在项目列表中点击 "编辑"，可修改完整名称、关键词、描述、知识库 ID 等字段。')
doc.add_paragraph('注意：一级域和数据来源创建后不可修改。')
add_screenshot_placeholder(doc, "编辑项目面板")

doc.add_heading('3.3 删除项目', level=2)
doc.add_paragraph('点击 "删除" 后需二次确认。删除为软删除——列表不再显示，但数据库保留记录。')

doc.add_heading('3.4 同步 INDEX', level=2)
doc.add_paragraph('右上角 "同步" 按钮将项目信息写入知识库 INDEX 文件。')
items = [
    '勾选 "同步后 git push"：同步完成后自动推送到 Git 远端',
    '同步范围：默认同步当前筛选条件下的项目',
    '同步结果：显示在页面顶部，包含合并的域数、项目数、新建目录数',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.5 搜索与筛选', level=2)
items = [
    '一级域下拉：按域筛选',
    '数据来源下拉：按来源筛选',
    '关键词搜索：支持原始 ID、完整名称、关键词、描述的模糊搜索',
    '分页：每页 20 条，底部翻页',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============ 四、文档批量录入详解 ============
doc.add_heading('四、文档批量录入详解', level=1)

doc.add_paragraph('入口：左侧菜单 "文档批量录入"。四步向导式操作。')
add_screenshot_placeholder(doc, "批量录入入口")

doc.add_heading('4.1 Step 1：解析目录', level=2)
items = [
    '粘贴 JoySpace 目录 URL（如 https://joyspace.jd.com/folders/xxxxxxxxxx）',
    '可选勾选 "解析时调用 copyFolderShortcuts 在目标知识库创建快捷方式"',
    '点击 "解析并创建快捷方式"',
    '系统返回：目录树、文档列表（数量）、快捷方式创建结果',
]
for i, item in enumerate(items, 1):
    doc.add_paragraph(f'{i}. {item}')

add_screenshot_placeholder(doc, "Step 1 解析目录结果")

doc.add_heading('4.2 Step 2：选择项目', level=2)
items = [
    '系统列出所有已有项目，点击卡片勾选目标项目（可多选）',
    '也可点 "+ 新建项目" 直接在本步骤创建新项目',
    '已勾选项目数显示在页面顶部',
    '确认后点 "下一步：摘要与匹配"',
]
for i, item in enumerate(items, 1):
    doc.add_paragraph(f'{i}. {item}')

add_screenshot_placeholder(doc, "Step 2 选择项目")

doc.add_heading('4.3 Step 3：摘要与匹配', level=2)
doc.add_paragraph('系统自动为每个文档生成 100 字摘要，并根据项目描述智能推荐归属项目。')

items = [
    '摘要可手动编辑——直接在文本框中修改',
    '归属项目可手动调整——下拉选择其他项目',
    '推荐结果显示推荐分数和命中关键词',
    '确认后点 "确认并写入项目目录"',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_screenshot_placeholder(doc, "Step 3 摘要与匹配")

doc.add_heading('4.4 Step 4：写入完成', level=2)
items = [
    '显示写入成功数和跳过数',
    '写入成功的文档列出：文档名 → 知识库路径（项目名）',
    '跳过的文档列出原因',
    '点 "同步 GIT" 提交到远端',
    '点 "再录入一批" 回到 Step 1',
]
for i, item in enumerate(items, 1):
    doc.add_paragraph(f'{i}. {item}')

add_screenshot_placeholder(doc, "Step 4 写入完成")

doc.add_page_break()

# ============ 五、AI 如何调用知识库 ============
doc.add_heading('五、AI 如何调用知识库（MCP 工具）', level=1)

doc.add_paragraph(
    'KaaS 通过 MCP（Model Context Protocol）协议对外暴露能力。'
    'AI 助手（如 Claude Code）通过 MCP 工具直接检索和读取知识库文档，无需人工中转。'
)

doc.add_heading('5.1 可用工具一览', level=2)
add_table(doc, ['工具名', '用途', '输入', '输出'], [
    ['kb_describe', '查看知识库能力说明', '无', 'KaaS manifest（能力、工具、数据源清单）'],
    ['kb_search', '关键词检索', 'query, top_k, scope', '匹配的文档片段 + citations'],
    ['kb_resolve', '查询路由决策', 'query', '候选主题列表 + 文件清单（不读正文）'],
    ['kb_pack', '组装主题知识包', 'topic, role, intent', '模版 + wiki + docs 全文组装'],
    ['kb_get_artifact', '读取单篇文档', 'path', '指定文件的完整正文'],
    ['kb_get_artifacts', '批量读取文档', 'paths[]', '多篇文件的完整正文'],
])

doc.add_heading('5.2 检索场景示例', level=2)

doc.add_paragraph('场景 1：找质检标准')
add_code_block(doc, '用户："电销质检的扣分标准是什么？"\nAI 调用：kb_search(query="电销质检扣分标准")\n返回：质检相关文档片段，附来源路径')

doc.add_paragraph()
doc.add_paragraph('场景 2：了解培训流程')
add_code_block(doc, '用户："新人培训的流程是怎样的？"\nAI 调用：kb_resolve(query="新人培训流程") → 路由到 training 目录\n       kb_search(query="新人培训流程", scope="training")\n返回：培训体系相关文档')

doc.add_paragraph()
doc.add_paragraph('场景 3：写方案需要参考资料')
add_code_block(doc, '用户："帮我写一份赛马机制优化方案"\nAI 调用：kb_pack(topic="competition", role="pm", intent="write_prd")\n返回：赛马相关的制度、方法论、工具模板完整打包')

doc.add_paragraph()
doc.add_paragraph('场景 4：读取特定文档')
add_code_block(doc, 'AI 调用：kb_get_artifact(path="domains/data-science/telemarketing-ops/docs/QI-2026-001.md")\n返回：该文档完整正文')

doc.add_heading('5.3 在 Claude Code 中使用', level=2)

doc.add_paragraph('Claude Code 已配置 kaas-knowledge MCP 服务器，连接地址：')
add_code_block(doc, 'https://ai-analysis-api.jd.com/mcp')

doc.add_paragraph('配置方式（.mcp.json）：')
add_code_block(doc, '{\n  "mcpServers": {\n    "kaas-knowledge": {\n      "type": "http",\n      "url": "https://ai-analysis-api.jd.com/mcp"\n    }\n  }\n}')

doc.add_paragraph('使用方式：直接在对话中提问，AI 会自动判断是否需要检索知识库。')
items = [
    '不需要手动调用工具——AI 根据问题内容自动决定',
    '检索结果会标注来源路径，方便溯源',
    '如果知识库没有相关内容，AI 会如实告知',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============ 六、AI 辅助编写 Wiki ============
doc.add_heading('六、AI 辅助编写 Wiki', level=1)

doc.add_paragraph(
    '知识库的价值取决于内容质量。AI 可以辅助编写、补充和审查文档，但核心内容仍需人工把关。'
)

doc.add_heading('6.1 用 AI 生成初稿', level=2)
doc.add_paragraph('操作步骤：')
items = [
    '告诉 AI 你要写什么文档（如 "帮我写一份电销质检标准"）',
    'AI 会先检索知识库中已有的相关文档作为参考',
    'AI 生成初稿，包含结构化的内容框架',
    '你审查、修改、补充实际业务数据',
    '确认后通过批量录入或手动上传到知识库',
]
for i, item in enumerate(items, 1):
    doc.add_paragraph(f'{i}. {item}')

add_code_block(doc, '示例对话：\n你："参考知识库里已有的质检相关文档，帮我写一份 2026 年 Q3 的质检标准"\nAI：[检索知识库] → [参考已有标准] → [生成初稿]\n你：审查修改 → 上传到知识库')

doc.add_heading('6.2 用 AI 补充关键词和描述', level=2)
doc.add_paragraph(
    '新文档入库时，关键词和描述的质量直接影响后续检索。可以让 AI 帮你生成：'
)
add_code_block(doc, '你："这篇文档的关键词和描述帮我写一下"\nAI：[分析文档内容] → [生成 TOPIC_KEYWORDS 和分类描述]\n你：确认后填入项目配置')

doc.add_heading('6.3 用 AI 做质量审查', level=2)
doc.add_paragraph('定期让 AI 审查知识库内容：')
items = [
    '检查是否有过时内容（如旧版制度未更新）',
    '检查是否有冲突内容（不同文档对同一规则说法不一）',
    '检查是否有缺失内容（某个模块长期没有新文档入库）',
    '检查关键词覆盖度（是否存在检索盲区）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============ 七、知识库维护最佳实践 ============
doc.add_heading('七、知识库维护最佳实践', level=1)

doc.add_heading('7.1 文档命名规范', level=2)
add_table(doc, ['编码前缀', '含义', '示例'], [
    ['QI-2026-xxx', '质检类', 'QI-2026-001 质检标准V1'],
    ['TR-2026-xxx', '培训类', 'TR-2026-001 新人培训大纲'],
    ['SY-2026-xxx', '制度类', 'SY-2026-001 电销管理办法'],
    ['MG-2026-xxx', '管理类', 'MG-2026-001 绩效考核规则'],
    ['ST-2026-xxx', '结算类', 'ST-2026-001 结算规则'],
    ['CP-2026-xxx', '赛马类', 'CP-2026-001 赛马机制'],
])

doc.add_heading('7.2 关键词与描述维护', level=2)
items = [
    '每新增一篇文档，同步更新项目的 TOPIC_KEYWORDS',
    '关键词用逗号分隔，覆盖核心业务术语',
    '分类描述要写清业务场景和归属范围',
    '避免过于宽泛的关键词（如 "管理""流程"），要用具体术语',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.3 定期健康检查', level=2)
items = [
    '每月检查一次：是否有过时文档需要更新',
    '每季度检查一次：关键词覆盖度是否足够',
    '发现冲突内容：记录到 notes/冲突记录.md，明确采用标准',
    '发现缺失内容：及时补充，避免检索盲区',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============ 八、FAQ ============
doc.add_heading('八、常见问题 FAQ', level=1)

faqs = [
    ('Q: 文档上传后多久能被 AI 检索到？',
     'A: 同步 INDEX 并 git push 后立即生效。如果没勾选自动 push，需要手动同步。'),
    ('Q: 关键词写错了怎么办？',
     'A: 在项目管理中编辑项目，修改关键词后重新同步即可。'),
    ('Q: 一个文档可以属于多个项目吗？',
     'A: 不可以。每篇文档归属一个项目。如果内容跨多个模块，建议在最核心的模块下创建，关键词覆盖其他模块。'),
    ('Q: AI 检索不到我的文档怎么办？',
     'A: 检查三点：① 文档是否已同步到知识库 ② 项目的 TOPIC_KEYWORDS 是否覆盖了你的查询词 ③ 分类描述是否准确'),
    ('Q: 批量录入时文档被跳过了怎么办？',
     'A: 查看跳过原因。常见原因：文档已存在、格式不支持。针对性处理后重新录入。'),
    ('Q: 知识库 ID 有什么用？',
     'A: 知识库 ID 是项目的唯一标识，用于 MCP 工具的精确调用。建议填写，方便后续管理。'),
    ('Q: 如何删除一篇文档？',
     'A: 目前需要在 Git 仓库中手动删除文件后重新同步 INDEX。项目管理中的删除是删除整个项目。'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.font.bold = True
    doc.add_paragraph(a)
    doc.add_paragraph()  # 间距

# ============ 保存 ============
output_path = '/Users/sundanian/Documents/projects/ai-agents/my-agent/wiki/KaaS智能知识库中台_操作手册.docx'
doc.save(output_path)
print(f"已生成：{output_path}")
