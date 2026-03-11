#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
315 附件 Word 文档生成脚本
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 设置中文字体
def set_font(run, font_name='微软雅黑', size=10.5):
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run

def create_attachment_1():
    """附件 1：合规月通知"""
    doc = Document()

    # 标题
    title = doc.add_heading('2026 年 315 服务合规月通知', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '微软雅黑', 18)

    # 副标题
    subtitle = doc.add_paragraph('自查自纠 · 合规经营 · 持续发展')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, '微软雅黑', 12)

    # 基本信息表
    doc.add_paragraph('\n')
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '时间'
    table.cell(0, 1).text = '2026 年 3 月 1 日 - 3 月 31 日'
    table.cell(1, 0).text = '范围'
    table.cell(1, 1).text = '全体京东金融电销外呼 BPO 供应商'
    table.cell(2, 0).text = '级别'
    table.cell(2, 1).text = '强制要求'

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, '微软雅黑', 10.5)

    # 第一章
    doc.add_heading('一、背景与目的', level=2)

    doc.add_heading('1.1 315 背景说明', level=3)
    doc.add_paragraph('央视 3·15 消费者权益保护晚会每年 3 月 15 日举办，曝光侵害消费者权益的典型案例。历年来，电销行业多次被点名。')

    # 历史教训表
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['年份', '曝光问题', '涉及行业']
    data = [
        ['2025', 'AI 外呼骚扰、数据泄露', '电销/催收'],
        ['2024', '电话骚扰、个人信息倒卖', '金融电销'],
        ['2023', '虚假宣传、诱导消费', '金融营销'],
        ['2022', '电信诈骗、非法外呼', '通信服务'],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    doc.add_heading('1.2 为什么要做合规月', level=3)
    doc.add_paragraph('我们的风险：')
    risks = [
        '电销外呼业务天然敏感，易被投诉',
        '供应商职场管理不规范，易被暗访取证',
        '品牌使用不当，易被认定为"冒充京东"',
        '数据安全漏洞，易引发法律风险',
    ]
    for risk in risks:
        p = doc.add_paragraph(risk, style='List Bullet')

    doc.add_paragraph('合规月目的：')
    purposes = [
        '自查自纠：315 前后主动排查风险，避免被曝光',
        '重申规则：强化供应商合规意识，不碰红线',
        '持续管理：建立长效机制，不只是 315 期间',
    ]
    for purpose in purposes:
        p = doc.add_paragraph(purpose, style='List Bullet')

    # 第二章
    doc.add_heading('二、315 期间特别管控要求', level=2)

    doc.add_heading('2.1 品牌与标识管理（高危）', level=3)
    doc.add_paragraph('核心规则：严禁未经授权使用京东品牌')

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['检查项', '要求', '常见违规']
    data = [
        ['Logo 悬挂', '职场内外禁止悬挂京东 Logo', '前台背景墙、工牌、文化墙'],
        ['京东名称使用', '禁止以"京东金融电销中心"名义对外', '招聘海报、职场指示牌、工服'],
        ['宣传物料', '所有含京东元素的物料立即撤除', '展板、横幅、易拉宝'],
        ['工服标识', '工服不得有京东 Logo 或"京东"字样', '坐席工服、管理人员马甲'],
        ['职场命名', '职场名称不得使用"京东"', '门牌、楼层指引、公司铭牌'],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    doc.add_heading('2.2 访客与外来人员管理（高危）', level=3)
    doc.add_paragraph('核心规则：严防黑衣人暗访')
    doc.add_paragraph('315 期间，央视、媒体、监管部门可能派人暗访取证。供应商必须严格执行访客管理。')

    doc.add_heading('2.3 作业区合规管控', level=3)
    doc.add_paragraph('核心规则：私人设备禁止入内')

    doc.add_heading('2.4 数据安全与隐私保护', level=3)
    doc.add_paragraph('核心规则：客户数据零泄露')

    # 第三章
    doc.add_heading('三、自查自纠条例', level=2)

    doc.add_heading('3.1 自查时间表', level=3)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['阶段', '时间', '动作', '交付物']
    data = [
        ['自查阶段', '3 月 1 日-3 月 10 日', '供应商对照条例自查', '《自查报告》'],
        ['整改阶段', '3 月 11 日-3 月 20 日', '对发现问题进行整改', '《整改报告》+ 照片'],
        ['验收阶段', '3 月 21 日-3 月 25 日', '京东现场/远程验收', '《验收报告》'],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    doc.add_heading('3.2 自查清单（必须逐项检查）', level=3)

    doc.add_paragraph('一、品牌与标识（10 项）', style='Heading 4')
    items = [
        '职场入口无京东 Logo、"京东金融"等标识',
        '前台背景墙无京东元素',
        '工服无京东 Logo 或"京东"字样',
        '工牌无京东标识',
        '招聘海报无"京东金融电销"等字样',
        '培训室无未经授权使用的京东培训材料',
        '文化墙无京东价值观标语（可用供应商自己的）',
        '会议室命名不含"京东"',
        '职场指引牌无京东元素',
        '公司官网/公众号无虚假宣传"京东合作"',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

    doc.add_paragraph('二、访客管理（8 项）', style='Heading 4')
    items = [
        '前台设置访客登记柜台',
        '使用《京东项目访客登记表》（统一模板）',
        '登记表信息完整（姓名/机构/事由/手机号）',
        '访客进入作业区有专人陪同',
        '核心作业区张贴"禁止拍照/录像"标识',
        '敏感区域（机房、作业区）门禁正常',
        '近 3 个月访客登记表归档保存',
        '建立访客审批流程（双审批）',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

    doc.add_paragraph('三、作业区管理（10 项）', style='Heading 4')
    items = [
        '作业区入口设置手机保管柜/袋',
        '坐席入场前手机上交',
        '作业区无私人手机使用现象',
        '作业电脑 USB 存储功能已禁用',
        '无 U 盘、移动硬盘连接记录',
        '作业区监控全覆盖、无盲区',
        '监控系统运行正常，可实时查询',
        '外网访问管控生效（娱乐/社交/游戏屏蔽）',
        '作业区无私人录音录像设备',
        '坐席无私自抄录客户信息行为',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

    doc.add_paragraph('四、数据安全（8 项）', style='Heading 4')
    items = [
        '客户数据无法通过 USB 导出',
        '系统权限按最小化原则配置',
        '离职工号及时注销',
        '无客户数据打印带出现象',
        '废纸篓无含客户信息的纸质文件',
        '坐席签署保密协议',
        '定期开展数据安全意识培训',
        '有数据泄露应急预案',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

    doc.add_paragraph('五、职场合规（8 项）', style='Heading 4')
    items = [
        '核心作业区独立物理空间',
        '核心作业区专属门禁',
        '与培训室/休息区物理分隔',
        '视频监控分辨率≥1080P',
        '机房独立物理空间（或有机房保护措施）',
        '防火设备齐全有效',
        '职场展示供应商公司资质',
        '前台区域信息完整（公司名称、资质等）',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

    # 第四章
    doc.add_heading('四、处罚要求（合规月严肃处理）', level=2)

    doc.add_heading('4.1 处罚原则', level=3)
    doc.add_paragraph('零容忍：以下违规行为，一经发现，立即处罚')

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    headers = ['违规行为', '处罚措施', '说明']
    data = [
        ['未经授权使用京东 Logo', '罚款 2 万元 + 暂停新项目 1 个月', '包括职场悬挂、工服、宣传物料'],
        ['以京东名义对外招聘', '罚款 2 万元 + 暂停新项目 1 个月', '包括招聘海报、网站、公众号'],
        ['访客未登记进入作业区', '罚款 1 万元/次', '包括黑衣人暗访得手'],
        ['坐席私人手机带入作业区', '罚款 5000 元/人次', '监控发现或现场检查'],
        ['使用 USB 设备拷贝数据', '罚款 2 万元 + 清退责任人', '数据零容忍'],
        ['拍照/传播客户信息', '立即清退 + 追究法律责任', '数据零容忍'],
        ['转包/分包京东业务', '立即终止合作', '合同红线'],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    # 第五章
    doc.add_heading('五、持续管理要求', level=2)

    doc.add_heading('5.1 合规不是 315 的事', level=3)
    doc.add_paragraph('合规是日常，不是 315 期间应付检查。')

    # 第六章
    doc.add_heading('六、联系方式', level=2)
    doc.add_paragraph('发现以下情况，立即上报：')
    items = [
        '媒体/监管人员到访',
        '可疑人员试图进入作业区',
        '数据泄露风险',
        '重大合规隐患',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('上报流程：发现人 → 供应商负责人 → 京东对接人')

    # 第七章
    doc.add_heading('七、附件', level=2)
    doc.add_paragraph('附件 1：《京东项目访客登记表》模板')
    doc.add_paragraph('附件 2：《禁止拍照/录像》标识模板')
    doc.add_paragraph('附件 3：自查报告提交方式')
    doc.add_paragraph('  - 提交邮箱：sundanian@jd.com')
    doc.add_paragraph('  - 提交截止：2026 年 3 月 10 日')
    doc.add_paragraph('  - 邮件主题：[供应商名称] 2026 年 315 自查报告')
    doc.add_paragraph('  - 附件格式：PDF 盖章版 + Word 可编辑版')

    # 第八章
    doc.add_heading('八、关键时间节点', level=2)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    headers = ['时间', '事项', '责任方']
    data = [
        ['3 月 1 日', '合规月启动，本通知生效', '京东'],
        ['3 月 10 日', '供应商提交《自查报告》', '供应商'],
        ['3 月 11-20 日', '供应商整改', '供应商'],
        ['3 月 20 日', '提交《整改报告》+ 照片', '供应商'],
        ['3 月 21-25 日', '京东验收', '京东'],
        ['3 月 26-31 日', '飞行检查、随机抽检', '京东'],
        ['4 月 1 日', '合规月总结通报', '京东'],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    # 结尾
    doc.add_paragraph('\n')
    p = doc.add_paragraph('请各供应商务必重视，认真执行。')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('合规是底线，不是选择题。')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('docs/linshi/315-附件 1-合规月通知.docx')
    print('✓ 附件 1 已生成：315-附件 1-合规月通知.docx')

def create_attachment_2():
    """附件 2：自查报告模板"""
    doc = Document()

    # 标题
    title = doc.add_heading('2026 年 315 自查报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '微软雅黑', 18)

    # 基本信息
    doc.add_paragraph('\n')
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    fields = [
        ['供应商名称', ''],
        ['职场地址', ''],
        ['提交日期', '2026 年 3 月 10 日前'],
        ['提交至', 'sundanian@jd.com'],
        ['负责人姓名', ''],
        ['负责人联系方式', ''],
        ['合规官姓名', ''],
    ]
    for i, (label, value) in enumerate(fields):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value

    # 自查情况
    doc.add_heading('一、品牌与标识（10 项）', level=2)
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Table Grid'
    headers = ['序号', '检查项', '是否符合', '问题描述（如不符合）']
    items = [
        '职场入口无京东 Logo、"京东金融"等标识',
        '前台背景墙无京东元素',
        '工服无京东 Logo 或"京东"字样',
        '工牌无京东标识',
        '招聘海报无"京东金融电销"等字样',
        '培训室无未经授权使用的京东培训材料',
        '文化墙无京东价值观标语',
        '会议室命名不含"京东"',
        '职场指引牌无京东元素',
        '公司官网/公众号无虚假宣传"京东合作"',
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, item in enumerate(items):
        table.cell(i+1, 0).text = str(i+1)
        table.cell(i+1, 1).text = item
        table.cell(i+1, 2).text = '□是 □否'
        table.cell(i+1, 3).text = ''

    doc.add_paragraph('本项符合率：____/10 = ____%')
    doc.add_paragraph('发现问题：')
    doc.add_paragraph('整改措施：')
    doc.add_paragraph('整改完成时间：')

    doc.add_heading('二、访客管理（8 项）', level=2)
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    items = [
        '前台设置访客登记柜台',
        '使用《京东项目访客登记表》（统一模板）',
        '登记表信息完整（姓名/机构/事由/手机号）',
        '访客进入作业区有专人陪同',
        '核心作业区张贴"禁止拍照/录像"标识',
        '敏感区域（机房、作业区）门禁正常',
        '近 3 个月访客登记表归档保存',
        '建立访客审批流程（双审批）',
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, item in enumerate(items):
        table.cell(i+1, 0).text = str(i+1)
        table.cell(i+1, 1).text = item
        table.cell(i+1, 2).text = '□是 □否'
        table.cell(i+1, 3).text = ''

    doc.add_paragraph('本项符合率：____/8 = ____%')
    doc.add_paragraph('发现问题：')
    doc.add_paragraph('整改措施：')
    doc.add_paragraph('整改完成时间：')

    doc.add_heading('三、作业区管理（10 项）', level=2)
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Table Grid'
    items = [
        '作业区入口设置手机保管柜/袋',
        '坐席入场前手机上交',
        '作业区无私人手机使用现象',
        '作业电脑 USB 存储功能已禁用',
        '无 U 盘、移动硬盘连接记录',
        '作业区监控全覆盖、无盲区',
        '监控系统运行正常，可实时查询',
        '外网访问管控生效（娱乐/社交/游戏屏蔽）',
        '作业区无私人录音录像设备',
        '坐席无私自抄录客户信息行为',
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, item in enumerate(items):
        table.cell(i+1, 0).text = str(i+1)
        table.cell(i+1, 1).text = item
        table.cell(i+1, 2).text = '□是 □否'
        table.cell(i+1, 3).text = ''

    doc.add_paragraph('本项符合率：____/10 = ____%')
    doc.add_paragraph('发现问题：')
    doc.add_paragraph('整改措施：')
    doc.add_paragraph('整改完成时间：')

    doc.add_heading('四、数据安全（8 项）', level=2)
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    items = [
        '客户数据无法通过 USB 导出',
        '系统权限按最小化原则配置',
        '离职工号及时注销',
        '无客户数据打印带出现象',
        '废纸篓无含客户信息的纸质文件',
        '坐席签署保密协议',
        '定期开展数据安全意识培训',
        '有数据泄露应急预案',
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, item in enumerate(items):
        table.cell(i+1, 0).text = str(i+1)
        table.cell(i+1, 1).text = item
        table.cell(i+1, 2).text = '□是 □否'
        table.cell(i+1, 3).text = ''

    doc.add_paragraph('本项符合率：____/8 = ____%')
    doc.add_paragraph('发现问题：')
    doc.add_paragraph('整改措施：')
    doc.add_paragraph('整改完成时间：')

    doc.add_heading('五、职场合规（8 项）', level=2)
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    items = [
        '核心作业区独立物理空间',
        '核心作业区专属门禁',
        '与培训室/休息区物理分隔',
        '视频监控分辨率≥1080P',
        '机房独立物理空间（或有机房保护措施）',
        '防火设备齐全有效',
        '职场展示供应商公司资质',
        '前台区域信息完整（公司名称、资质等）',
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, item in enumerate(items):
        table.cell(i+1, 0).text = str(i+1)
        table.cell(i+1, 1).text = item
        table.cell(i+1, 2).text = '□是 □否'
        table.cell(i+1, 3).text = ''

    doc.add_paragraph('本项符合率：____/8 = ____%')
    doc.add_paragraph('发现问题：')
    doc.add_paragraph('整改措施：')
    doc.add_paragraph('整改完成时间：')

    # 整体自评
    doc.add_heading('六、整体自评', level=2)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ['检查大类', '符合项数', '总项数', '符合率']
    data = [
        ['品牌与标识', '', '10', ''],
        ['访客管理', '', '8', ''],
        ['作业区管理', '', '10', ''],
        ['数据安全', '', '8', ''],
        ['职场合规', '', '8', ''],
        ['总计', '', '44', ''],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    doc.add_paragraph('总符合率：____%')
    doc.add_heading('主要风险点：', level=4)
    doc.add_paragraph('')
    doc.add_heading('整改措施：', level=4)
    doc.add_paragraph('')

    # 承诺
    doc.add_heading('七、承诺', level=2)
    doc.add_paragraph('本公司承诺以上自查情况真实、准确，如有虚假，愿接受相应处罚。')
    doc.add_paragraph('\n负责人签字：________________')
    doc.add_paragraph('公司盖章：')
    doc.add_paragraph('日期：2026 年____月____日')

    # 附件清单
    doc.add_heading('八、附件清单', level=2)
    items = [
        '附件 1：职场照片（前台、作业区、手机保管区、访客登记区）',
        '附件 2：访客登记表样本（近 3 个月）',
        '附件 3：合规培训记录（签到表、课件、考核成绩）',
        '附件 4：其他佐证材料',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    # 提交说明
    doc.add_heading('九、提交说明', level=2)
    items = [
        '本报告需同时提交 PDF 盖章版和 Word 可编辑版',
        '邮件主题：[供应商名称] 2026 年 315 自查报告',
        '提交邮箱：sundanian@jd.com',
        '提交截止：2026 年 3 月 10 日',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

    doc.save('docs/linshi/315-附件 2-自查报告模板.docx')
    print('✓ 附件 2 已生成：315-附件 2-自查报告模板.docx')

def create_attachment_3():
    """附件 3：访客登记表模板"""
    doc = Document()

    # 标题
    title = doc.add_heading('京东项目访客登记表', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '微软雅黑', 18)

    doc.add_paragraph('\n')
    doc.add_paragraph('职场名称：________________________')
    doc.add_paragraph('登记日期：2026 年____月____日')

    # 访客登记表
    doc.add_heading('访客登记表', level=2)

    table = doc.add_table(rows=11, cols=10)
    table.style = 'Table Grid'
    headers = ['序号', '访客姓名', '所属机构', '来访事由', '联系电话', '进入区域', '陪同人员', '进入时间', '离场时间', '备注']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        # 设置列宽
        if i == 0:
            table.columns[i].width = Cm(0.8)
        elif i in [1, 2, 3]:
            table.columns[i].width = Cm(2)
        else:
            table.columns[i].width = Cm(1.5)

    for i in range(1, 11):
        table.cell(i, 0).text = str(i)

    # 访客类型说明
    doc.add_heading('访客类型说明', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['类型', '说明', '审批要求']
    data = [
        ['A 类 - 业务访客', '客户、合作伙伴、供应商', '供应商负责人审批'],
        ['B 类 - 技术访客', '设备维护、系统支持人员', '供应商负责人 + 京东对接人双审批'],
        ['C 类 - 其他访客', '面试人员、快递、外卖等', '禁止进入作业区'],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    # 特别注意
    p = doc.add_paragraph()
    run = p.add_run('特别注意：')
    run.bold = True
    items = [
        '所有访客必须提前预约，现场登记',
        '进入作业区必须由专人全程陪同',
        '严禁在作业区拍照、录像',
        '发现可疑人员（媒体/监管暗访）立即上报',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    # 可疑人员上报流程
    doc.add_heading('可疑人员上报流程', level=2)
    doc.add_paragraph('发现以下情况，立即上报：')
    items = [
        '声称"媒体采访"、"市场调研"的陌生人',
        '要求进入核心作业区拍照/录音',
        '询问敏感问题（薪资、话术、数据来源）',
        '未预约直接到访，要求"参观"',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('上报流程：发现人 → 供应商负责人 → 京东对接人')
    doc.add_paragraph('\n紧急联系人：')
    doc.add_paragraph('供应商负责人：________________ 电话：________________')
    doc.add_paragraph('京东对接人：________________ 电话：________________')

    # 填表说明
    doc.add_heading('填表说明', level=2)
    items = [
        '访客姓名：填写真实姓名，不得代填',
        '所属机构：填写公司/单位全称',
        '来访事由：具体说明来访目的（如：设备维护、业务洽谈等）',
        '联系电话：填写手机号码',
        '进入区域：填写具体区域（如：前台、会议室、作业区等）',
        '陪同人员：填写负责陪同的员工姓名',
        '进入/离场时间：精确到分钟（如：14:30）',
        '备注：其他需要说明的情况',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

    # 存档要求
    doc.add_heading('存档要求', level=2)
    items = [
        '本登记表由前台/行政主管负责保管',
        '每月汇总归档，保存期限≥1 年',
        '电子档 + 纸质档双备份',
        '京东有权随时调阅检查',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')

    doc.add_paragraph('\n')
    doc.add_paragraph('表格版本：2026-V1.0')
    doc.add_paragraph('生效日期：2026 年 3 月 1 日')

    doc.save('docs/linshi/315-附件 3-访客登记表模板.docx')
    print('✓ 附件 3 已生成：315-附件 3-访客登记表模板.docx')

def create_attachment_4():
    """附件 4：禁止拍照/录像标识模板"""
    doc = Document()

    # 标题
    title = doc.add_heading('《禁止拍照/录像》标识模板', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '微软雅黑', 18)

    # 版本一
    doc.add_heading('版本一：标准版（推荐）', level=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('⚠️\n\n敏 感 区 域\n\n禁 止 拍 照 / 录 像\n\n')
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('本区域涉及商业机密\n严禁使用手机、相机等设备进行拍照、录像\n\n违者将追究法律责任')
    run.font.size = Pt(12)

    doc.add_paragraph('\n尺寸建议：A4 或 A3')
    doc.add_heading('张贴位置：', level=4)
    items = ['作业区入口', '机房入口', '会议室（涉及敏感会议时）', '培训室（涉及敏感培训时）']
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    # 版本二
    doc.add_heading('版本二：简洁版', level=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('📵 禁止拍照 / 录像\nNo Photography / Recording')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph('\n尺寸建议：A5 或 A6')
    doc.add_heading('张贴位置：', level=4)
    items = ['工位隔板', '会议桌', '前台登记处']
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    # 版本三
    doc.add_heading('版本三：警示版', level=2)

    p = doc.add_paragraph()
    run = p.add_run('⚠️ 警 示\n\n')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph('本区域为核心作业区，涉及客户数据和商业秘密。')
    doc.add_paragraph('\n严禁以下行为：', style='List Bullet')
    items = ['拍照、录像', '使用手机等摄录设备', '携带 USB 存储设备']
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('\n违规行为将：', style='List Bullet')
    items = ['立即清退', '追究法律责任', '纳入行业黑名单']
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('\n举报热线：XXX-XXXX-XXXX')

    doc.add_paragraph('\n尺寸建议：A3')
    doc.add_heading('张贴位置：', level=4)
    items = ['作业区主入口', '职场文化墙', '新员工培训室']
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    # 制作说明
    doc.add_heading('制作说明', level=2)

    doc.add_heading('颜色规范', level=3)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['元素', '颜色', '色值']
    data = [
        ['背景', '白色', '#FFFFFF'],
        ['边框', '红色', '#E74C3C'],
        ['标题', '黑色', '#000000'],
        ['警示图标', '黄色 + 黑色', '#F39C12 + #000000'],
        ['重点文字', '红色', '#E74C3C'],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    doc.add_heading('字体规范', level=3)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['内容', '字体', '字号']
    data = [
        ['主标题', '黑体/思源黑体 Bold', '48pt'],
        ['副标题', '黑体/思源黑体 Regular', '32pt'],
        ['正文', '宋体/思源宋体', '24pt'],
        ['英文', 'Arial', '24pt'],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    doc.add_heading('材质建议', level=3)
    items = [
        '室内：背胶写真 + 冷裱',
        '长期张贴：KT 板 + 边框',
        '临时张贴：A4 纸打印 + 透明胶',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('张贴高度', level=3)
    doc.add_paragraph('标识中心点距地面：150-160cm（视线平齐）')
    doc.add_paragraph('入口处：门框上方或侧边')

    # 使用场景
    doc.add_heading('使用场景', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['场景', '推荐版本', '数量']
    data = [
        ['作业区主入口', '版本三（警示版）', '1 个'],
        ['机房入口', '版本一（标准版）', '1 个'],
        ['前台登记处', '版本二（简洁版）', '1 个'],
        ['会议室', '版本二（简洁版）', '每间 1 个'],
        ['培训室', '版本一（标准版）', '1 个'],
        ['工位隔板', '版本二（简洁版）', '每排 1 个'],
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = cell

    doc.add_paragraph('\n')
    doc.add_paragraph('版本：2026-V1.0')
    doc.add_paragraph('生效日期：2026 年 3 月 1 日')

    doc.save('docs/linshi/315-附件 4-禁止拍照录像标识模板.docx')
    print('✓ 附件 4 已生成：315-附件 4-禁止拍照录像标识模板.docx')

if __name__ == '__main__':
    print('开始生成 315 附件 Word 文档...\n')
    create_attachment_1()
    create_attachment_2()
    create_attachment_3()
    create_attachment_4()
    print('\n✓ 全部生成完成！')
