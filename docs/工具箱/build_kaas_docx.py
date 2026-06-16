#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 KaaS 知识库中台操作手册 Word 文档
依赖: python-docx (pip install python-docx)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG_DIR = os.path.dirname(os.path.abspath(__file__)) + "/kaas-images"
OUT = os.path.dirname(os.path.abspath(__file__)) + "/KaaS知识库中台-操作手册.docx"

# 颜色
JD_RED = RGBColor(0xD7, 0x1E, 0x2E)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x1A, 0x6E, 0xD8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1A, 0x88, 0x3A)
ORANGE = RGBColor(0xE8, 0x7A, 0x1E)

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.font.color.rgb = DARK
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35

# 页边距
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

def shade(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10, white=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if white:
        run.font.color.rgb = WHITE
    elif color:
        run.font.color.rgb = color

def add_table(headers, rows, col_widths=None, header_color="D71E2E"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Light Grid'
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, white=True, size=9.5)
        shade(hdr[i], header_color)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def h(text, level=1, color=None, num=None):
    p = doc.add_paragraph()
    p.space_before = Pt(14 if level==1 else 10)
    p.space_after = Pt(6)
    if level == 1:
        # 红色横条标题
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '18')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'D71E2E')
        pbdr.append(bottom)
        pPr.append(pbdr)
        run = p.add_run(("  " if num else "") + (f"{num}  " if num else "") + text)
        run.font.size = Pt(17)
        run.font.color.rgb = JD_RED
        run.bold = True
    elif level == 2:
        run = p.add_run((f"{num}  " if num else "") + text)
        run.font.size = Pt(13.5)
        run.font.color.rgb = JD_RED
        run.bold = True
    else:
        run = p.add_run((f"{num}  " if num else "") + text)
        run.font.size = Pt(11.5)
        run.font.color.rgb = BLUE
        run.bold = True
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def para(text, bold=False, color=None, size=10.5, after=4, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level*0.6)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.name = '微软雅黑'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = '微软雅黑'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def num_list(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.name = '微软雅黑'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = '微软雅黑'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_image(filename, width_cm=15.5, caption=None):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        para(f"[图片缺失: {filename}]", color=ORANGE, size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        crun = cap.add_run(caption)
        crun.font.size = Pt(9)
        crun.font.color.rgb = GRAY
        crun.italic = True
        crun.font.name = '微软雅黑'
        crun.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def callout(title, text, kind="tip"):
    """提示框: tip=绿, warn=橙, info=蓝"""
    colors = {"tip": ("1A883A", "E8F5E9"), "warn": ("E87A1E", "FFF3E0"), "info": ("1A6ED8", "E3F2FD"), "danger": ("D71E2E", "FFEBEE")}
    border_c, fill_c = colors.get(kind, colors["info"])
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Cm(16)
    shade(cell, fill_c)
    # 左边框加粗上色
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24'); left.set(qn('w:color'), border_c)
    borders.append(left)
    for side in ['top','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4'); b.set(qn('w:color'), border_c)
        borders.append(b)
    tcPr.append(borders)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if title:
        r1 = p.add_run(title + "  ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = RGBColor.from_string(border_c)
        r1.font.name = '微软雅黑'
        r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.name = '微软雅黑'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pbdr.append(bottom)
    pPr.append(pbdr)

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

print("样式定义完成, 开始填充内容...")

# ============================================================
# 封面
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("KaaS 智能知识库中台")
r.font.size = Pt(34)
r.font.color.rgb = JD_RED
r.bold = True
r.font.name = '微软雅黑'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("操作手册")
r.font.size = Pt(26)
r.font.color.rgb = DARK
r.bold = True
r.font.name = '微软雅黑'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
r = p.add_run("—— 如何用 AI 写 wiki、调用知识库、维护文档 ——")
r.font.size = Pt(13)
r.font.color.rgb = GRAY
r.font.name = '微软雅黑'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(6):
    doc.add_paragraph()

# 信息表
info_t = doc.add_table(rows=5, cols=2)
info_t.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("归属", "京东集团 · 京东科技 · 金融科技事业群"),
    ("管理台", "http://ai-kaas.pre-apps.jd.com"),
    ("MCP 接口", "https://ai-analysis-api.jd.com/mcp/"),
    ("服务版本", "kaas-knowledge v1.27.2"),
    ("文档日期", "2026 年 6 月"),
]
for i, (k, v) in enumerate(info_data):
    set_cell_text(info_t.rows[i].cells[0], k, bold=True, white=True, size=10)
    shade(info_t.rows[i].cells[0], "D71E2E")
    set_cell_text(info_t.rows[i].cells[1], v, size=10)
    info_t.rows[i].cells[0].width = Cm(3.5)
    info_t.rows[i].cells[1].width = Cm(10)

doc.add_page_break()

# ============================================================
# 目录提示
# ============================================================
h("本手册怎么读", level=1)
para("本手册面向第一次接触 KaaS 的同事,目标是用最低的学习成本上手。按下面的角色快速定位你需要看的章节:", after=8)

add_table(
    ["你的角色", "先看哪几章", "5 分钟能干嘛"],
    [
        ["产品经理 / 需求方", "第二章概念 → 第三章新增项目 → 第六章AI写PRD", "建好项目,让 AI 帮你起草二期 PRD"],
        ["前端 / 后端工程师", "第二章概念 → 第四章录入文档 → 第六章AI写代码", "录入历史文档,让 AI 带上下文改代码"],
        ["测试工程师", "第二章概念 → 第四章录入文档 → 第六章AI写用例", "录入用例库,让 AI 补回归范围"],
        ["团队 Lead / 平台维护者", "通读全文 → 第七章协作规范", "制定团队接入规范"],
    ],
    col_widths=[3.8, 6.2, 5.5]
)
para("")
callout("💡 三个最重要的结论(先记住)", "", "tip")
bullet("KaaS = 事业群的\"历史文档大脑\"。写 PRD、改代码、写用例前先问它,能拿到一期决策、历史坑、接口契约。", bold_prefix="它是什么: ")
bullet("人用 Web 管理台(录入/维护),AI 用 MCP 接口(检索/写作)。两边是同一个知识库。", bold_prefix="两种用法: ")
bullet("写任何东西之前,先 kb_search → 再 kb_pack 拿全量,最后才动笔。只用搜索片段是已知反模式。", bold_prefix="铁律: ")

doc.add_page_break()
print("封面+导读完成")

# ============================================================
# 第一章: 这是什么
# ============================================================
h("一、KaaS 是什么?为什么要用?", level=1, num="")

para("KaaS(Knowledge as a Service)是金融科技事业群级的知识库中台。它把散落在各团队的 PRD、前后端实现文档、测试用例、架构说明,统一沉淀到一个 Git 仓库,再通过两种方式对外提供能力:", after=6)

add_table(
    ["", "Web 管理台", "MCP 服务"],
    [
        ["给谁用", "人(产品/研发/测试)", "AI Agent(ZCode / Cursor / 其他)"],
        ["干什么", "建项目、录入文档、维护 wiki、调试召回", "检索知识、起草 PRD、生成代码、补充用例"],
        ["地址", "http://ai-kaas.pre-apps.jd.com", "https://ai-analysis-api.jd.com/mcp/"],
        ["背后", "同一个 Git 知识库(241 个文件)", "← 同一个"],
    ],
    col_widths=[2.5, 6.5, 6.5]
)
para("")

h("解决什么痛点", level=2, num="1.1")
bullet("新需求来了,找不到一期的 PRD 和决策记录 —— 现在让 AI 直接从知识库召回全量上下文。")
bullet("换人接手老项目,不知道历史踩过什么坑 —— wiki 里沉淀了\"决策因果\"和\"避坑指南\"。")
bullet("前后端联调,接口契约散落在各处 —— 知识库按项目归档,一次 kb_pack 全拿到。")
bullet("写二期方案只能凭记忆 —— AI 带着一期全文帮你起草,覆盖入口枚举、用户心智、卖点字段。")

h("当前已收录的项目", level=2, num="1.2")
add_image("01-project-management.png", 15.5, "▲ 项目管理页 —— 当前已收录的 4 个项目")
para("截至 2026 年 6 月,知识库已收录以下项目(在项目管理页可见):", after=4)

add_table(
    ["项目 slug", "中文名", "一级域", "路径", "关键词(示例)"],
    [
        ["voiceprint", "支付-声纹识别", "payment", "domains/payment/voiceprint2", "声纹、声纹支付、语音支付、生物识别"],
        ["aipay", "京东AI付", "payment", "domains/payment/aipay", "京东AI购、AI任务功能、AI 眼镜"],
        ["payment-basic", "支付基础的业务", "payment", "domains/payment/payment-basic", "支付SDK、支付基础工具"],
        ["rental-3c", "企业金融中3C租赁", "enterprise-finance", "domains/enterprise-finance/rental-3c", "企业金融、3C租赁、电子产品"],
    ],
    col_widths=[2.8, 2.8, 2.2, 4.5, 3.2]
)
para("")
callout("⚠️ 注意", "KaaS 只收录\"金融科技事业群\"知识库仓库内的文档。实时生产数据、线上日志、其他事业群的文档不在范围内。", "warn")

doc.add_page_break()
print("第一章完成")

# ============================================================
# 第二章: 核心概念
# ============================================================
h("二、先搞懂三个概念(3 分钟)", level=1, num="")

para("KaaS 的知识按\"一级域 → 项目 → 文件\"三层组织。理解这三层,后面所有操作都顺理成章:", after=8)

add_table(
    ["层级", "叫什么", "是什么", "举例"],
    [
        ["第 1 层", "一级域 (L1 Domain)", "业务大块,目录的第一层", "payment(支付)、enterprise-finance(企业金融)"],
        ["第 2 层", "项目 (Topic)", "具体业务主题,检索的最小单元", "voiceprint(声纹)、aipay(AI付)"],
        ["第 3 层", "文件", "项目下的各类文档", "docs/历史PRD、wiki/角色知识、template/模板"],
    ],
    col_widths=[1.8, 2.8, 5, 5.9]
)
para("")

h("仓库目录结构(重要)", level=2, num="2.1")
para("每个项目(topic)在 Git 仓库里有固定的 4 类子目录,这就是 AI 检索时拿到的内容:")
code_block("""domains/
└── payment/                       ← 一级域
    └── voiceprint/                ← 项目(topic)
        ├── template/              ← 模板:PRD 章节骨架,写新 PRD 的参照
        ├── wiki/                  ← 角色知识:用户心智、规范、决策因果、避坑指南
        │   ├── pm-wiki.md         ← 产品经理视角的知识
        │   ├── fe-wiki.md         ← 前端视角(组件规范、页面约定)
        │   ├── be-wiki.md         ← 后端视角(接口契约、部署约定)
        │   └── evolution-wiki.md  ← 演进记录(版本决策、ADR)
        ├── docs/                  ← 历史 PRD、设计文档(按日期命名)
        │   ├── 240520_声纹开通H5组件化.md
        │   ├── 250514_声纹支付-支付侧H5能力建设.md
        │   └── 260302_生物支付新增声纹识别.md
        └── artifacts/             ← 附件:架构图、ADR、Runbook
            └── PRD规约模版.md""")
para("")

h("知识库(可选)", level=2, num="2.2")
para("项目之外还有个\"知识库(Knowledge Base)\"概念,是可选的。一个项目可以绑定一个知识库(用于向量检索),在项目的高级选项里填\"知识库 ID\"(如 kb_e4bbdf31)。不绑定也能用,只是检索会退化为关键词匹配。", after=6)
add_image("06-knowledge-base.png", 13, "▲ 知识库管理页 —— 每个知识库显示文档数、切片数、embedding 模型")

callout("💡 一句话总结", "一级域是大类,项目是你要检索的主题,项目下的 4 类文件(template/wiki/docs/artifacts)是 AI 能读到的全部内容。你的任务就是把文档放对位置。", "tip")

doc.add_page_break()
print("第二章完成")

# ============================================================
# 第三章: 新增项目(操作步骤)
# ============================================================
h("三、如何新增一个项目(图文分步)", level=1, num="")

para("新增项目是所有同事都会做的第一个操作。这一章用真实截图带你走一遍。", after=6)

h("第 1 步:进入项目管理", level=2, num="3.1")
num_list("打开 http://ai-kaas.pre-apps.jd.com,用京东 ERP 账号登录。")
num_list("左侧导航点击「项目管理 —— 一级域下项目维护」。")
num_list("页面右上角点击「+ 新建项目」按钮。")
add_image("step-01-locate-btn.png", 15.5, "▲ 第①步:在项目管理页找到右上角红色「+ 新建项目」按钮")
add_image("step-02-fill-form.png", 15.5, "▲ 第②步:弹出表单,填写一级域/项目slug/名称/关键词/描述(带✅为必填)")
add_image("step-03-advanced.png", 15.5, "▲ 第③步(可选):勾选「高级选项」展开知识库ID/层级/所有者")

h("第 2 步:填写表单", level=2, num="3.2")
para("点击新建后弹出表单,字段如下表。带 ✅ 的是必填:")
para("")

add_table(
    ["字段", "必填", "怎么填", "示例"],
    [
        ["一级域", "✅", "下拉选业务大类", "payment"],
        ["来源", "✅", "默认 topic_system;对接旧系统才选 legacy_api", "topic_system"],
        ["项目 slug", "✅", "小写英文单词,会成为目录名", "voiceprint"],
        ["原始 ID", "默认同 slug", "勾选才出现,对接外部系统时用", "voiceprint2"],
        ["名称", "✅", "项目中文名", "支付-声纹识别"],
        ["关键词", "✅", "逗号分隔的业务词,影响检索打分", "声纹,声纹支付,语音支付"],
        ["描述", "✅", "说清业务场景与边界,文档自动分类会参考", "声纹支付开通、多入口、H5组件化"],
    ],
    col_widths=[2.2, 1.5, 5.5, 5.3]
)
para("")

h("第 3 步:高级选项(可选)", level=2, num="3.3")
para("勾选「填写高级选项」可展开更多字段:")
para("")
bullet("知识库 ID:绑定一个已创建的知识库(如 kb_e4bbdf31),启用向量检索。")
bullet("层级:文档归属的业务层级(如\"支付域-内场交易域-交易域\")。")
bullet("文档所有者:指定该项目文档的负责人。")

h("第 4 步:创建与同步", level=2, num="3.4")
num_list("确认信息无误后点击「创建」。")
num_list("如果需要立即推送到 Git 仓库,创建前在表格上方勾选「同步后 git push」再点「同步」。")
callout("⚠️ 关于「同步后 git push」", "勾选后,你在管理台的所有改动会自动 commit + push 到 Git 仓库。这意味着所有接入知识库的 AI Agent 会立即看到新内容。如果只是测试,先不要勾。", "warn")

h("关键词和描述怎么写(关键!)", level=2, num="3.5")
para("这两个字段直接决定检索质量。常见错误是写得太简略。对比一下:")
add_table(
    ["", "❌ 差的写法", "✅ 好的写法"],
    [
        ["关键词", "声纹", "声纹,声纹支付,语音支付,朗读验证,按住朗读,生物识别支付,子频道"],
        ["描述", "声纹相关", "声纹支付开通流程,支持多入口(APP/H5/眼镜),按住朗读方案,个性化设置"],
    ],
    col_widths=[2, 6.5, 7]
)
callout("💡 写法口诀", "关键词 = 用户会搜的所有同义词,用逗号分隔;描述 = 一句话说清\"这个项目管什么事、边界在哪\"。", "tip")

doc.add_page_break()
print("第三章完成")

# ============================================================
# 第四章: 录入文档
# ============================================================
h("四、如何录入文档(从 JoySpace 批量导入)", level=1, num="")

para("项目建好后,要把历史文档灌进去。KaaS 支持从京东 JoySpace(云文档)批量导入,这是最常用的录入方式。", after=6)

h("标准接入流程(3 步)", level=2, num="4.1")
para("「接入指南」页已经画好了完整流程,核心是三步:")
add_image("03-onboarding.png", 15.5, "▲ 接入指南页 —— 标准接入流程")

num_list("进入项目管理,确认或新建目标项目(上一章已讲)。", bold_prefix="")
num_list("进入「文档批量录入」,粘贴 JoySpace 文件夹路径,系统自动解析目录树。", bold_prefix="")
num_list("按步骤为每份文档归属到对应项目,系统自动写入 Git 知识库 + RAG 知识库(双写)。", bold_prefix="")
para("")
callout("💡 双写机制", "文档录入后会同时写入两个地方:① Git 原始副本(版本追溯)② RAG 知识库(向量检索)。所以你录入后,AI 既能搜到片段,也能读全文。", "info")

h("文档批量录入操作(四步向导)", level=2, num="4.2")
add_image("05-knowledge-input.png", 15.5, "▲ 文档批量录入页 —— 四步向导")

add_table(
    ["步骤", "做什么", "说明"],
    [
        ["① 解析目录", "粘贴 JoySpace 文件夹 URL", "系统自动展开目录树,列出子文件夹和文档清单"],
        ["② 选择项目", "为文档选择目标项目", "可新建项目或选已有项目"],
        ["③ 摘要与匹配", "系统自动生成摘要并智能匹配归属", "可人工调整每份文档归属"],
        ["④ 写入完成", "自动写入对应项目目录", "触发加工流水线:解析 → Wiki 生成 → ES 灌入"],
    ],
    col_widths=[2.5, 4.5, 8.5]
)
para("")

h("加工流水线:文档怎么被消化", level=2, num="4.3")
para("文档录入后不是直接可用,而是经过三层加工。在「加工流水线」页可以监控进度:")
add_image("04-pipeline.png", 15.5, "▲ 加工流水线监控页 —— 看到每篇文档的处理状态")

add_table(
    ["处理阶段", "做什么", "状态查看"],
    [
        ["1. 解析", "解析文档结构、提取正文", "解析中 / 已完成 / 失败"],
        ["2. Wiki 生成", "AI 自动提炼角色 wiki(用户心智、决策因果)", "Wiki生成中 / 已完成"],
        ["3. ES 灌入", "写入向量索引,供语义检索", "ES灌入中 / 已完成"],
    ],
    col_widths=[2.5, 7, 5.5]
)
callout("⚠️ 处理失败怎么办", "如果看到「失败」状态,通常是文档格式问题(扫描件 PDF、特殊排版)。联系平台维护者或在文档录入页重试。失败的文档不会被检索到。", "warn")

h("文档命名规范(重要!)", level=2, num="4.4")
para("为了让 AI 正确归类和检索,文档文件名建议遵循:")
code_block("""✅ 推荐命名: YYMMDD_文档标题.md
   示例: 250514_声纹支付-支付侧H5能力建设.md
        260302_生物支付新增声纹识别.md

✅ JoySpace 云文档: 会被自动加上 [云文档] 前缀
   示例: [云文档]PRD-京小贝接入AI付

❌ 避免: 新建文档.md / 未命名.md / 复件_xxx.md""")
para("命名清晰的好处:kb_resolve 和 kb_search 返回结果时,标题就是文件名,同事一眼能看懂是哪篇。")

doc.add_page_break()
print("第四章完成")

# ============================================================
# 第五章: 日常维护
# ============================================================
h("五、如何维护知识库(新增/更新/废弃)", level=1, num="")

para("知识库不是建好就完了,需要持续维护。这一章讲三种最常见的维护场景。", after=6)

h("场景 1:新增文档", level=2, num="5.1")
num_list("在「文档批量录入」页,粘贴新的 JoySpace 文件夹路径。")
num_list("按四步向导归属到项目(参考第四章)。")
num_list("在「加工流水线」确认状态变成「已完成」。")
num_list("在「召回调试」页用一个相关查询验证能搜到新文档。")

h("场景 2:更新已有文档", level=2, num="5.2")
para("文档内容有变化时(比如 PRD 改版),有两种方式:")
add_table(
    ["方式", "操作", "适用场景"],
    [
        ["重新录入", "在文档批量录入页重新导入,JoySpace 最新内容会覆盖旧版", "文档在 JoySpace 有更新"],
        ["Git 直接改", "clone 仓库 git@coding.jd.com:ai-infra-apps/knowledge-base.git,改完后 push", "需要精细修改 wiki 或 template"],
    ],
    col_widths=[2.5, 8, 5]
)
callout("💡 关于 git push 权限", "管理台的「同步」按钮 + 勾选「同步后 git push」会自动提交。如果你直接操作 Git 仓库,需要有 ai-infra-apps/knowledge-base 的写权限。不确定就联系平台维护者。", "info")

h("场景 3:废弃过时文档", level=2, num="5.3")
bullet("在 JoySpace 删除或归档原文档后,知识库里的旧副本不会自动清除。")
bullet("需要联系平台维护者,或在 Git 仓库手动删除对应文件并 push。")
bullet("废弃前建议在 wiki 里记一笔(如\"v1 方案已于 X 月废弃,原因:YYY\"),保留决策痕迹。")

h("场景 4:手动触发同步", level=2, num="5.4")
para("如果发现知识库内容滞后(别人改了你搜不到),可以手动触发 git pull:")
num_list("在「项目管理」页顶部点击「同步」按钮。")
num_list("或调用 REST API: POST http://ai-analysis-api.jd.com/api/agent/sync")
callout("💡 同步 vs 录入的区别", "「同步」= 拉取 Git 仓库最新内容(别人提交的改动);「录入」= 你主动把新文档写入知识库。两者方向相反。", "info")

h("召回调试:验证你的文档能被搜到", level=2, num="5.5")
para("录入完成后,务必验证检索效果。「召回调试」页支持 REST / MCP / SKILL 三种方式测试:")
add_image("08-retrieval-debug.png", 13, "▲ 召回调试页 —— 输入查询测试召回效果")
num_list("进入「召回调试」页。")
num_list("输入一个你刚录入文档里会出现的查询(如\"声纹支付开通流程\")。")
num_list("点击发送,查看返回的文档列表和打分。")
num_list("如果搜不到 → 检查文档是否在「加工流水线」显示已完成;检查项目的关键词/描述是否充分。")

doc.add_page_break()
print("第五章完成")

# ============================================================
# 第六章: AI 调用(MCP) - 核心
# ============================================================
h("六、AI 如何调用知识库(MCP 完整指南)", level=1, num="")

para("这是本手册最重要的一章。讲清楚:AI 怎么接入知识库、怎么检索、怎么帮你写 wiki / 写 PRD / 改代码。", after=6)

callout("🎯 本章目标", "看完后你能回答:① AI 怎么连上知识库 ② 有哪些工具可用 ③ 不同任务该用哪条流程 ④ 怎么让 AI 写出高质量 wiki。", "info")

h("第一步:接入 MCP(一次性配置)", level=2, num="6.1")
para("在你的 AI Agent(如 ZCode、Cursor)配置文件 .mcp.json 里加入:")
code_block("""{
  "mcpServers": {
    "kaas-knowledge": {
      "type": "http",
      "url": "https://ai-analysis-api.jd.com/mcp/"
    }
  }
}""")
callout("⚠️ URL 必须带尾斜杠", "地址是 https://ai-analysis-api.jd.com/mcp/(结尾有 /)。不带斜杠会触发 307 重定向,导致连接失败。这是最常见的接入错误。", "danger")
para("配置后重启 AI Agent,工具列表里会出现 kb_ 开头的 6 个工具。", after=4)

h("可用工具一览(6 个)", level=2, num="6.2")
add_table(
    ["工具", "用途", "什么时候用"],
    [
        ["kb_describe", "返回服务能力自描述", "第一次接入时,了解服务能力"],
        ["kb_search", "关键词检索,返回片段+citations", "快速定位、探索阶段"],
        ["kb_resolve", "路由决策+候选主题文件清单", "不确定 topic 归属时"],
        ["kb_pack ★", "一次性组装模板+wiki+历史PRD全文", "★ 写作前首选(写PRD/代码/用例)"],
        ["kb_get_artifact", "读取单篇完整原文", "精读某一篇文档"],
        ["kb_get_artifacts", "批量读取多篇原文", "补齐 pack 里没加载的文件"],
    ],
    col_widths=[3.2, 5.3, 7]
)
para("")

h("第二步:选对工作流(按任务类型)", level=2, num="6.3")
para("服务端强制要求按任务类型选流程,违反是已知反模式。一共有三条:", after=6)

# 流程 A
h("流程 A:写 PRD / 写 wiki / 二期需求(首选)", level=3, num="")
para("这是最常用、也是最容易做错的流程。核心原则:写之前先 kb_pack 拿全量,不要只用 kb_search 片段。", after=4)
add_table(
    ["步骤", "调用", "拿到什么"],
    [
        ["1", "kb_search(query=\"需求名\")", "命中哪个 topic,心中有数"],
        ["2 ★", "kb_pack(topic, role=\"pm\", intent=\"write_prd\", query=需求名)", "模板+角色wiki+一期PRD全文"],
        ["3", "AI 按模板章节顺序起草", "完整 PRD 初稿"],
        ["4", "kb_get_artifacts(paths=[...]) 补缺口", "pack 里没加载的文件"],
    ],
    col_widths=[1.2, 8.5, 5.8]
)
para("")

# 流程 B
h("流程 B:改前端 / 改后端 / 写测试用例", level=3, num="")
add_table(
    ["步骤", "调用", "拿到什么"],
    [
        ["1", "kb_search(query=..., role=\"fe\"/\"be\"/\"qa\")", "topic + related_paths"],
        ["2", "kb_pack(topic=..., role=同上, intent=frontend/backend/test)", "角色 wiki + 关联文档"],
        ["3", "kb_get_artifacts(paths=[...])", "补齐缺口"],
    ],
    col_widths=[1.2, 8.5, 5.8]
)
para("")

# 流程 C
h("流程 C:仅探索 / 调试路由", level=3, num="")
code_block("kb_describe → kb_resolve(query) → kb_resolve 看文件清单 → kb_search 看片段")
para("只看不写,用于了解知识库覆盖范围。", after=4)

h("第三步:理解 kb_pack 的返回结构", level=2, num="6.4")
para("kb_pack 是核心工具,它一次返回四样东西。理解这个结构,才知道 AI 拿到了什么:")
add_table(
    ["字段", "是什么", "真实示例(voiceprint 项目)"],
    [
        ["template", "PRD 模板(章节骨架)", "PRD规约模版.md,372行,基于11份PRD逆向归纳"],
        ["wikis", "角色知识(按 role 过滤)", "pm-wiki.md(用户心智)、evolution-wiki.md(决策因果)"],
        ["docs", "历史 PRD 全文(按相关性排序)", "云文档-京小贝-AI付接入调研、京东AI付营销方案"],
        ["siblings", "未加载的文件清单", "docs_not_loaded(9篇)、wikis_not_loaded(4篇)、artifacts(8个)"],
    ],
    col_widths=[2.2, 4.5, 8.8]
)
callout("💡 预算机制", "kb_pack 有字数预算(docs_budget=40000 字符)。文档太多会截断,被截断的列在 siblings 里,需要时用 kb_get_artifacts 补。可通过 docs_top_k / wikis_top_k 调整数量。", "info")

print("第六章上半完成")

h("第四步:角色(role)怎么选", level=2, num="6.5")
para("kb_search 和 kb_pack 都有 role 参数,它决定 AI 拿到哪类 wiki。选错角色会拿到不相关的知识:")
add_table(
    ["role 值", "角色", "拿到的 wiki 重点", "适合谁"],
    [
        ["pm", "产品经理", "用户心智、权益状态机、入口枚举、卖点字段", "写 PRD / 二期方案"],
        ["fe", "前端工程师", "Vue/React 组件规范、页面约定、接口联调", "改前端代码"],
        ["be", "后端工程师", "服务设计、接口契约、数据库变更、部署约定", "改 Java 服务"],
        ["qa", "测试工程师", "测试方案、用例库、缺陷与回归范围", "写测试用例"],
        ["ops", "运维", "部署、配置、Runbook", "部署/排障"],
    ],
    col_widths=[1.5, 2.2, 7.5, 4.3]
)
para("")

h("真实案例:用 AI 写声纹识别二期 PRD", level=2, num="6.6")
para("下面是完整的端到端流程,展示 AI 如何利用知识库写 PRD。这是同事可以直接照做的示范:", after=6)
add_image("demo-ai-read-kb.png", 14, "▲ AI 调用知识库的真实演示:kb_search 定位 → kb_pack 拿全量上下文")

para("① AI 先搜索,定位 topic:", bold=True, size=10, color=BLUE, after=2)
code_block("""AI 调用: kb_search(query="声纹支付 多入口 个性化", top_k=5, role="pm")

返回: 命中 5 篇文档,全部属于 domains/payment/voiceprint
  [#1] 260302_生物支付新增声纹识别.md
  [#2] 250514_声纹支付-支付侧H5能力建设.md
  [#3] 260424_新终端声纹开通-双端收音长按方案.md
  [#4] 251125_AI秒付—开通声纹流程移至眼镜.md
  [#5] 260424_双端收音双开声纹-BPRD.md
→ AI 确认 topic = voiceprint""")

para("② AI 拉取知识包,拿全量上下文:", bold=True, size=10, color=BLUE, after=2)
code_block("""AI 调用: kb_pack(topic="voiceprint", role="pm", intent="write_prd", query="多入口 个性化")

返回(45KB):
  ├── template:  PRD规约模版.md (372行,含[*必填][*条件必填]标记)
  ├── wikis:     pm-wiki.md(用户心智/权益状态机)
  │             evolution-wiki.md(决策因果/避坑)
  │             evolution-log.md(版本演进记录)
  ├── docs:      云文档-京小贝-AI付接入调研.md
  │             云文档-京东AI付营销方案.md
  └── siblings:  还有9篇docs+4篇wiki未加载(按需补)
→ AI 拿到了:模板骨架 + 历史决策 + 一期完整PRD""")

para("③ AI 按模板起草 PRD:", bold=True, size=10, color=BLUE, after=2)
code_block("""AI 现在可以写出高质量二期PRD:
  - 按照 template 的章节顺序(概述→背景→流程→字段→异常...)
  - 引用 wikis 里的用户心智和权益状态机(不会漏掉)
  - 基于 docs 里的一期PRD做增量(而不是从零写)
  - 标注每个引用的来源 path(citations)

输出示例:
  "基于 PRD规约模版.md 第2章(产品流程概述),
   一期 250514_声纹支付H5能力建设.md 已实现APP端开通,
   二期需新增眼镜端(参考 251125_AI秒付眼镜方案.md)..."
""")

para("④ 如有缺口,补齐:", bold=True, size=10, color=BLUE, after=2)
code_block("""AI 调用: kb_get_artifacts(paths=[
    "domains/payment/voiceprint/docs/260302_生物支付新增声纹识别.md",
    "domains/payment/voiceprint/artifacts/PRD规约模版.md"
])
→ 拿到完整原文,继续完善PRD""")
para("")

h("反模式:不要这么做", level=2, num="6.7")
callout("🚫 反模式 1", "只用 kb_search 的片段就写二期 PRD / 改代码。片段不足以覆盖入口枚举、用户心智、历史决策、卖点字段。评测显示这样会严重漏内容。", "danger")
callout("🚫 反模式 2", "只读一篇 doc 就动笔。应该同时读 template + wiki + 至少一篇 doc。", "danger")
callout("🚫 反模式 3", "调完 kb_search 后忽略返回里的 next_step_hint 和 related_paths。这两个字段是服务端给你的行动建议。", "danger")

doc.add_page_break()

# ============================================================
# 第七章: 协作规范
# ============================================================
h("七、团队协作规范(谁做什么)", level=1, num="")

para("KaaS 是团队共建的知识库。这一章明确分工,避免\"谁都觉得别人会维护、最后没人维护\"。", after=6)

h("角色分工表", level=2, num="7.1")
add_table(
    ["角色", "职责", "在 KaaS 做什么", "频率"],
    [
        ["产品经理", "需求方、PRD 作者", "新建项目、录入一期PRD、用AI起草二期", "每个需求"],
        ["前端/后端", "开发方", "录入接口文档/实现说明、用AI带上下文改代码", "每次开发"],
        ["测试工程师", "质量方", "录入用例库、用AI补回归范围", "每轮测试"],
        ["项目 Lead", "知识库负责人", "审核项目归属、维护wiki、定期清理过时文档", "每周"],
        ["平台维护者", "KaaS 平台方", "处理失败的文档、Git仓库权限、整体同步", "日常"],
    ],
    col_widths=[2.2, 2.8, 7, 2.5]
)
para("")

h("一个需求的协作流程示例", level=2, num="7.2")
para("以\"声纹支付二期\"为例,展示团队如何配合:", after=4)
add_table(
    ["阶段", "谁做", "在 KaaS 做什么"],
    [
        ["需求启动", "产品经理", "确认项目已存在 → AI 用 kb_pack 拿一期PRD+模板 → 起草二期PRD"],
        ["PRD 评审", "产品+研发", "AI 生成的PRD经人工审核修改 → 定稿后录入 docs/ 目录"],
        ["开发", "前端+后端", "AI 用 kb_pack(role=fe/be) 拿组件规范+接口契约 → 改代码"],
        ["开发完成", "前端+后端", "把新增的接口文档/实现说明录入知识库"],
        ["测试", "测试工程师", "AI 用 kb_pack(role=qa) 拿历史用例 → 补回归用例"],
        ["上线后", "项目Lead", "在 evolution-wiki.md 记录本次决策和避坑经验"],
    ],
    col_widths=[2.2, 2.5, 10]
)
para("")

h("wiki 谁来写?AI 写还是人写?", level=2, num="7.3")
para("这是同事最常问的问题。答案是:AI 负责初稿,人负责审核和补充。", after=4)
add_table(
    ["wiki 类型", "谁生成", "内容", "人要做什么"],
    [
        ["pm-wiki.md", "AI 自动消化(加工流水线)", "用户心智、权益状态机、入口枚举", "审核准确性,补充AI漏掉的"],
        ["fe/be-wiki.md", "AI 自动消化", "组件规范、接口约定", "同上"],
        ["evolution-wiki.md", "★ 人写为主", "决策因果(ADR)、避坑指南、版本取舍", "每次上线后主动记录"],
        ["evolution-log.md", "AI 自动+人补", "版本变更记录", "检查是否有遗漏"],
    ],
    col_widths=[3, 3, 5.5, 4]
)
callout("💡 关键认知", "带\"(自动消化)\"标记的 wiki 是 AI 在文档录入时自动生成的,你不需要手写。但 evolution-wiki(决策因果、避坑)AI 写不好,需要人在每次重要决策后手动补充。这是知识库最有价值的部分。", "tip")

doc.add_page_break()
print("第六章+第七章完成")

# ============================================================
# 第八章: 外部集成 API
# ============================================================
h("八、外部集成:API 与三种调用方式", level=1, num="")

para("除了在 AI Agent 里用 MCP,KaaS 还提供 REST API 和 SKILL 两种调用方式,供外部系统集成。", after=6)

add_image("07-api-integration.png", 15.5, "▲ 外部集成页 —— 三种集成方式与 API 清单")

h("三种集成方式", level=2, num="8.1")
add_table(
    ["方式", "适合谁", "怎么用", "示例"],
    [
        ["MCP", "AI Agent(ZCode/Cursor)", "在 .mcp.json 配置,AI 自动调用 kb_* 工具", "AI 写 PRD、改代码"],
        ["RESTful API", "外部系统/脚本", "直接 HTTP 调用,需 Bearer 鉴权", "CI 流水线、内部工具"],
        ["SKILL", "技能化封装", "把知识库能力封装成可复用技能", "批量PRD生成器"],
    ],
    col_widths=[2.2, 3.5, 6, 3.8]
)
para("")

h("REST API 完整清单", level=2, num="8.2")
para("Base URL: http://ai-analysis-api.jd.com", bold=True, after=2)
add_table(
    ["方法", "Endpoint", "说明"],
    [
        ["GET", "/api/agent/manifest", "服务能力自描述(角色/工具/数据源)"],
        ["GET", "/api/agent/llm-config", "LLM 配置快照(是否Mock/已配Key)"],
        ["GET", "/api/agent/sources", "已注册 Retriever 及状态"],
        ["GET", "/api/agent/artifact", "按仓库路径读原文(支持行范围)"],
        ["POST", "/api/agent/sync", "手动 git pull 同步远端"],
        ["POST ★", "/api/agent/sync-index", "pull → 重写 INDEX → push(重建索引)"],
        ["POST ★", "/api/agent/push", "本地改动 git add → commit → push"],
        ["POST", "/api/agent/run", "LangGraph 多层 Agent 问答"],
        ["POST", "/api/agent/resume", "人工审核断点恢复(approve/reject)"],
        ["POST", "/api/ingestion/run", "处理队首一条 pending 文档"],
        ["POST ★", "/api/digest/ingest", "消化:git pull → 生成Wiki+L1-draft"],
        ["GET", "/api/digest/items", "查询消化元数据列表"],
        ["GET", "/api/digest/items/{id}", "消化详情 + 引用 + 事件"],
        ["GET", "/api/digest/items/{id}/events", "仅生命周期事件"],
    ],
    col_widths=[1.5, 6.5, 7.5]
)
para("")
callout("💡 带 ★ 的是写入操作", "sync-index / push / digest/ingest 会修改 Git 仓库。调用前确认你了解影响。日常检索用 GET 类接口和 MCP 即可,不需要碰这些。", "warn")

h("调用示例(REST)", level=2, num="8.3")
para("手动触发同步:", bold=True, after=2)
code_block("""curl -X POST http://ai-analysis-api.jd.com/api/agent/sync \\
  -H "Authorization: Bearer <你的token>\"""")
para("读取某篇文档原文:", bold=True, after=2)
code_block("""curl "http://ai-analysis-api.jd.com/api/agent/artifact?\\
path=domains/payment/voiceprint/docs/250514_声纹支付H5.md" \\
  -H "Authorization: Bearer <你的token>\"""")

doc.add_page_break()

# ============================================================
# 第九章: 快速上手 + FAQ
# ============================================================
h("九、5 分钟上手 + 常见问题", level=1, num="")

h("新人 5 分钟跑通清单", level=2, num="9.1")
callout("🎯 照着做,5 分钟完成第一次调用", "", "tip")
para("")
num_list("打开 http://ai-kaas.pre-apps.jd.com,ERP 登录。", bold_prefix="第1步 登录: ")
num_list("左侧「项目管理」→ 找到你的业务项目(如 voiceprint),确认存在。", bold_prefix="第2步 找项目: ")
num_list("左侧「召回调试」→ 输入一个业务关键词(如\"声纹支付开通\")→ 看到命中文档。", bold_prefix="第3步 验证: ")
num_list("在你的 .mcp.json 加入 kaas-knowledge 配置(见 6.1)→ 重启 Agent。", bold_prefix="第4步 接AI: ")
num_list("对 AI 说\"帮我查一下声纹支付的历史PRD\"→ AI 自动调用 kb_search + kb_pack。", bold_prefix="第5步 用起来: ")
para("")

h("常见问题(FAQ)", level=2, num="9.2")

faqs = [
    ("Q: MCP 连接失败怎么办?",
     "检查 URL 是否带尾斜杠(https://ai-analysis-api.jd.com/mcp/)。其次确认在内网环境。如果报 401/403 说明需要 token,联系平台维护者。"),
    ("Q: 页面打开是空白的?",
     "多数页面需要登录。点右上角「登录」,用京东 ERP 账号。"),
    ("Q: kb_search 搜不到我的文档?",
     "① 检查文档是否在「加工流水线」显示「已完成」(失败的不会被检索);② 检查项目的关键词/描述是否充分;③ 调大 top_k(默认8,二期建议10-20);④ 在「召回调试」页测试。"),
    ("Q: AI 写的 PRD 质量不行?",
     "大概率是没用 kb_pack,只用了 kb_search 片段。务必走\"流程A\":先 search 定位 topic,再 pack 拿全量,最后才写。"),
    ("Q: 文档录入后多久能被检索到?",
     "取决于加工流水线速度。通常几分钟内完成(解析→Wiki生成→ES灌入)。在「加工流水线」页看实时状态。"),
    ("Q: 我能直接改 Git 仓库吗?",
     "可以,如果你有 ai-infra-apps/knowledge-base 的写权限。但更推荐用管理台操作,会自动触发加工流水线。直接改 Git 的话,改完要在管理台点「同步」。"),
    ("Q: role 参数选错了有影响吗?",
     "会拿到不相关的 wiki(如 pm 拿不到 fe 的组件规范)。但检索打分不受影响,role 只影响 hint 和 wiki 过滤。改对重调即可。"),
    ("Q: 两个人同时录入文档会冲突吗?",
     "不会。录入走的是 JoySpace 目录解析,系统自动处理。但直接操作 Git 仓库时,多人同时 push 可能冲突,建议协调。"),
    ("Q: 怎么知道知识库里有哪些文档?",
     "用 kb_resolve(query=\"你的业务词\") 会返回候选主题的完整文件清单。或在管理台「加工流水线」看所有已录入文档。"),
    ("Q: evolution-wiki 一定要写吗?",
     "强烈建议。这是知识库最有价值的部分(决策因果、避坑指南),AI 写不好,只有亲历者能写。每次重要上线后花5分钟记一笔,长期价值巨大。"),
]
for q, a in faqs:
    para(q, bold=True, color=BLUE, size=10.5, after=2)
    para(a, size=10, after=8, indent=True)

doc.add_page_break()

# ============================================================
# 附录: 速查卡
# ============================================================
h("附录:速查卡", level=1, num="")

h("常用地址", level=2, num="")
add_table(
    ["用途", "地址"],
    [
        ["管理台", "http://ai-kaas.pre-apps.jd.com"],
        ["MCP 接口", "https://ai-analysis-api.jd.com/mcp/"],
        ["REST API Base", "http://ai-analysis-api.jd.com"],
        ["Git 仓库", "git@coding.jd.com:ai-infra-apps/knowledge-base.git"],
    ],
    col_widths=[3.5, 12]
)
para("")

h("MCP 配置(复制即用)", level=2, num=""
)
code_block("""{
  "mcpServers": {
    "kaas-knowledge": {
      "type": "http",
      "url": "https://ai-analysis-api.jd.com/mcp/"
    }
  }
}""")
para("")

h("三条铁律", level=2, num=""
)
num_list("写 PRD/wiki 前必走 kb_pack,不要只用 kb_search 片段。", bold_prefix="铁律一: ")
num_list("文档命名带日期前缀(YYMMDD_标题.md),关键词描述写充分。", bold_prefix="铁律二: ")
num_list("每次重要决策后在 evolution-wiki 记一笔(避坑指南/决策因果)。", bold_prefix="铁律三: ")
para("")

h("遇到问题找谁", level=2, num=""
)
add_table(
    ["问题类型", "找谁"],
    [
        ["页面打不开/登录问题", "平台维护者(金融科技事业群 KaaS 团队)"],
        ["文档加工失败", "平台维护者"],
        ["Git 仓库权限", "平台维护者"],
        ["AI 调用报错", "先看本手册 FAQ,再找平台维护者"],
        ["项目归属不确定", "你的项目 Lead"],
    ],
    col_widths=[5, 10.5]
)
para("")
para("")

divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— 本手册基于 KaaS 管理台实际界面(2026-06-16)和 MCP 服务 v1.27.2 整理 —")
r.font.size = Pt(9)
r.font.color.rgb = GRAY
r.italic = True
r.font.name = '微软雅黑'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("如界面有更新,以实际页面为准。有问题反馈给平台维护者。")
r.font.size = Pt(9)
r.font.color.rgb = GRAY
r.italic = True
r.font.name = '微软雅黑'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

print("全部内容写入完成, 保存文档...")
doc.save(OUT)
print(f"✅ 已保存: {OUT}")
import os
print(f"   文件大小: {os.path.getsize(OUT)/1024:.1f} KB")
