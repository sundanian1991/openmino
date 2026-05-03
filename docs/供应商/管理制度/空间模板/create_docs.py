#!/usr/bin/env python3
"""生成两个Word文档：供应商填报指南（对外）+ 京东内部使用指南（对内）"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
GREEN = "2E8B6E"
DARK = "333333"
LIGHT_GRAY = "666666"

def style_doc(doc):
    """Apply default styles"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    # Set East Asian font
    rPr = style.element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(20)
    p.space_after = Pt(20)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.name = 'Arial'
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Arial'
    return p

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.name = 'Arial'
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.name = 'Arial'
    return p

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.name = 'Arial'
    return p

def add_blank(doc):
    return doc.add_paragraph()

def add_list(doc, items, numbered=True):
    for item in items:
        if numbered:
            p = doc.add_paragraph(item, style='List Number')
        else:
            p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.name = 'Arial'

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Arial'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Green background
        shading = cell._element.get_or_add_tcPr()
        shade = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): GREEN,
            qn('w:val'): 'clear'
        })
        shading.append(shade)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Arial'

    # Column widths
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): 'CCCCCC'
        })
        borders.append(b)
    tblPr.append(borders)

    add_blank(doc)
    return table

# ============================================================
# Document 1: 供应商填报指南 (External)
# ============================================================
def create_filling_guide():
    doc = Document()
    style_doc(doc)

    add_title(doc, '供应商工作空间填报指南')
    add_subtitle(doc, '适用对象：供应商对接人 / 项目主管')
    add_subtitle(doc, '生效日期：2026-04-18  |  版本：v1.2')

    add_h1(doc, '一、工作空间说明')
    add_para(doc, '这是贵司在京东的供应商管理工作空间。所有与项目管理相关的数据、文档、沟通记录都统一存放在此。')
    add_para(doc, '原则：数据透明、信息同步、问题早发现早解决。', italic=True)

    add_h1(doc, '二、文件夹结构')
    add_para(doc, '每个供应商拥有独立的工作空间，文件夹结构如下：')
    make_table(doc,
        ['文件夹', '类型', '说明'],
        [
            ['01-人力管理', '必填', '供应商基础信息表、人力明细表、主管信息表、离职登记'],
            ['02-质检合规', '必填', '质检台账、合规自查表、质检月报'],
            ['03-沟通留痕', '选填', '约谈记录、工作沟通、邮件归档'],
            ['04-整改跟踪', '必填', '整改计划表、整改进度跟踪'],
            ['05-整改通知', '只读', '京东发布的巡检整改通知'],
            ['06-供应商自主使用', '自选', '贵司可自行上传内部资料'],
            ['07-通知与公告', '只读', '京东发布的通知、制度更新'],
        ],
        col_widths=[3.5, 2, 10]
    )

    add_h1(doc, '三、必须填报的表格清单')
    make_table(doc,
        ['表格', '填报频率', '截止时间', '填报人', '不填报的影响'],
        [
            ['供应商基础信息表', '准入时填写', '准入时', '对接人', ''],
            ['人力明细表', '周报', '每周一 12:00 前', '对接人', ''],
            ['主管信息表', '月度/变更', '每月5日/变更24h内', '对接人', ''],
            ['离职原因登记表', '事件触发', '员工离职当日', '对接人', ''],
            ['质检台账', '日常', '发现问题当日', '对接人', ''],
            ['合规自查表', '月报', '每月5日前', '对接人', ''],
            ['整改计划表', '事件触发', '约谈后48h内', '对接人', ''],
            ['整改进度跟踪', '周报', '每周五 18:00 前', '对接人', ''],
        ],
        col_widths=[3, 2.2, 3, 1.8, 4.5]
    )

    add_h1(doc, '四、各表填报说明')

    # 4.1
    add_h2(doc, '4.1 供应商基础信息表')
    add_para(doc, '填写位置：01-人力管理/供应商基础信息表.xlsx', bold=True)
    add_para(doc, '填报频率：准入时填写 / 人员变更时24小时内更新')
    add_list(doc, [
        '合作准入时填写首次信息',
        '供应商名称：公司全称',
        '统一社会信用代码：18位代码',
        '法人代表：法定代表人姓名',
        '联系人姓名/职务：日常对接人及其职务',
        '联系电话/微信号/邮箱：至少填写一项',
        '合作起始日期：首次合作日期',
        '服务业务线：从下拉菜单选择',
        '团队规模：当前团队人数',
        '人员变更时，在最后一行新增一条记录，备注栏说明变更内容',
    ], numbered=True)

    # 4.2
    add_h2(doc, '4.2 人力明细表（周报）')
    add_para(doc, '填写位置：01-人力管理/人力明细表.xlsx', bold=True)
    add_para(doc, '每周一 12:00 前提交，填写上周数据。')
    add_list(doc, [
        '填报日期：填写提交当天的日期',
        '填报人：填写对接人姓名',
        '周期起始/结束：填写上周周一到周日',
        '总人数：当日在职总人数（含熟手+新人）',
        '熟手人数：入职超过30天的人数',
        '新人人数：入职30天以内（含30天）的人数',
        '本周新入职：本周内新入职的人数',
        '本周离职：本周内离职的人数',
        '日均有效出勤率：日均有效出勤人数 \u00f7 目标人数',
        '日均有效出勤人数：日均通话时长超过30分钟的人数',
        '备注：异常情况说明，无异常留空',
    ], numbered=True)
    add_para(doc, '注意：熟手人数 + 新人人数 = 总人数', bold=True, italic=True)

    # 4.3
    add_h2(doc, '4.3 主管信息表（月度确认/变更）')
    add_para(doc, '填写位置：01-人力管理/主管信息表.xlsx', bold=True)
    add_list(doc, [
        '每月5日前确认信息是否有变化',
        '无变化："本次是否变更"填"否"，其他不动',
        '有变化（更换主管/联系方式变更）："本次是否变更"填"是"，填写变更类型、变更原因、生效日期',
    ], numbered=True)

    # 4.4
    add_h2(doc, '4.4 离职原因登记表（事件触发）')
    add_para(doc, '填写位置：01-人力管理/离职管理/离职原因登记表.xlsx', bold=True)
    add_para(doc, '员工离职当天登记。')
    add_list(doc, [
        '离职日期：实际离职日期',
        '姓名/岗位/入职日期：基本信息',
        '在职天数：自动计算（离职日期 - 入职日期），无需手填',
        '离职类型：从下拉菜单选择（主动离职/被动辞退/合同到期不续签/其他）',
        '离职原因（一级）：从下拉菜单选择大类别',
        '离职原因（二级）：具体描述',
        '主管是否面谈：是/否',
        '面谈记录摘要：面谈了解到的核心原因',
    ], numbered=True)

    # 4.5
    add_h2(doc, '4.5 质检台账（日常记录）')
    add_para(doc, '填写位置：02-质检合规/质检台账.xlsx', bold=True)
    add_para(doc, '发现问题当日登记，即使一天只有几条也要记录。')
    add_list(doc, [
        '登记日期：发现问题的日期',
        '业务线/坐席姓名/工号：涉事坐席信息',
        '通话/工单编号：关联的通话或工单',
        '问题类型：从下拉菜单选择（A类错误/B类错误/C类错误/服务态度/流程违规/其他）',
        '问题描述：具体描述问题',
        '严重程度：从下拉菜单选择（A类/B类/C类/提醒）',
        '处理结果：已采取的处理措施',
        '整改措施：后续整改计划',
        '质检来源：从下拉菜单选择（质检抽检/客户投诉/内部自查/甲方反馈）',
    ], numbered=True)

    # 4.6
    add_h2(doc, '4.6 合规自查表（月报）')
    add_para(doc, '填写位置：02-质检合规/合规自查表.xlsx', bold=True)
    add_para(doc, '每月5日前填报上月自查结果。')
    add_list(doc, [
        '供应商名称：公司全称',
        '自查月份：填报哪个月的数据（如 2026-04）',
        '填报日期：提交当天日期',
        '检查项编号、名称、内容已预设，不需要改',
        '自查结果：从下拉菜单选择（是/否/不适用）',
        '如果选"否"，必须填写：不符合说明、整改计划、整改完成时间',
    ], numbered=True)

    # 4.7
    add_h2(doc, '4.7 整改计划表（事件触发）')
    add_para(doc, '填写位置：04-整改跟踪/整改计划表.xlsx', bold=True)
    add_para(doc, '约谈后48小时内提交。')
    add_list(doc, [
        '基本信息：供应商名称、约谈日期、提交日期',
        '触发原因：从下拉菜单选择',
        '核心问题描述：约谈中指出的问题',
        '整改项：每条整改项填写一行（整改项编号、整改目标、具体动作、责任人、完成时间、验收标准）',
        '具体动作：3-5条具体动作，用序号分隔',
    ], numbered=True)

    # 4.8
    add_h2(doc, '4.8 整改进度跟踪表（周报）')
    add_para(doc, '填写位置：04-整改跟踪/整改进度跟踪.xlsx', bold=True)
    add_para(doc, '每周五18:00前更新。')
    add_list(doc, [
        '每个整改项对应一行',
        '填报日期：提交当天日期',
        '整改项编号：对应整改计划表中的编号',
        '整改目标：从整改计划表复制',
        '本周进度(%)：0-100的数字',
        '本周具体完成情况：做了什么、结果如何',
        '是否按期：是/否',
        '阻碍/风险：如有阻碍填写，无阻碍留空',
        '预计完成日期：预计哪天能完成',
    ], numbered=True)

    # 五
    add_h1(doc, '五、填报规范')
    add_h2(doc, '5.1 命名规范')
    make_table(doc,
        ['项目', '命名规则', '示例'],
        [
            ['填报日期', 'YYYY-MM-DD', '2026-04-14'],
            ['周期', 'YYYY-MM-DD ~ YYYY-MM-DD', '2026-04-07 ~ 2026-04-13'],
            ['月份', 'YYYY-MM', '2026-04'],
        ],
        col_widths=[3.5, 6, 6]
    )

    add_h2(doc, '5.2 数据质量要求')
    add_list(doc, [
        '真实性：数据必须真实，不得造假。一旦发现数据造假，直接记入供应商诚信档案',
        '及时性：按截止时间填报，迟报影响排名中的配合度评分',
        '完整性：必填项不能留空',
        '一致性：人力明细表的本周离职数 = 离职登记表的本月登记数',
    ], numbered=False)

    # 六
    add_h1(doc, '六、常见问题')
    qas = [
        ('Q1：人力明细表中的"有效出勤"怎么定义？', 'A：当日通话时长超过30分钟的坐席才算有效出勤。'),
        ('Q2：熟手和新人的划分标准是什么？', 'A：入职超过30天为熟手，30天以内（含）为新人。'),
        ('Q3：如果没有离职，离职登记表需要填吗？', 'A：不需要。离职登记表是事件触发，没有离职事件就不填。'),
        ('Q4：整改计划没按时完成怎么办？', 'A：在整改进度跟踪表的"是否按期"填"否"，并在"阻碍/风险"栏说明原因。'),
        ('Q5：表格填错了怎么修改？', 'A：直接在原表格修改，不要删除原行重新插入。'),
        ('Q6：对接人换了怎么办？', 'A：在主管信息表中更新对接人信息，并提前通知京东对接人。'),
    ]
    for q, a in qas:
        add_para(doc, q, bold=True)
        add_para(doc, a)

    # 七
    add_h1(doc, '七、联系方式')
    make_table(doc,
        ['角色', '姓名', '联系方式'],
        [['供应商管理者', '京东对接人', '[电话/微信]']],
        col_widths=[4, 4, 4]
    )

    add_blank(doc)
    add_para(doc, '如有疑问，请及时沟通。', italic=True)

    path = os.path.join(BASE, '供应商工作空间填报指南.docx')
    doc.save(path)
    print(f'OK: {path}')

# ============================================================
# Document 2: 京东内部使用指南
# ============================================================
def create_internal_guide():
    doc = Document()
    style_doc(doc)

    add_title(doc, '供应商工作空间 \u2014 使用指南')
    add_subtitle(doc, '适用对象：京东供应商管理者')
    add_subtitle(doc, '版本：v1.2  |  日期：2026-04-18')

    add_h1(doc, '一、模板包说明')
    add_para(doc, '这是一套可直接上传到云端的供应商管理工作空间模板包。每个供应商一个独立的工作空间文件夹，用于统一存放人力、质检、整改等数据。')

    add_h1(doc, '二、使用流程')
    steps = [
        ('第一步：复制模板', '将"[供应商名称]-工作空间"文件夹复制一份，重命名为实际供应商名称，如：毅航-工作空间/、毛毛虫-工作空间/'),
        ('第二步：上传云端', '将每个供应商的文件夹上传到云端共享目录'),
        ('第三步：通知供应商', '将《供应商工作空间填报指南》发给供应商对接人，告知需填报的表格及截止时间'),
        ('第四步：日常维护', '供应商按频率填报，京东侧定期收集数据、汇总基础信息、发布整改通知'),
    ]
    for title, desc in steps:
        add_para(doc, title, bold=True)
        add_para(doc, desc)

    add_h1(doc, '三、日常维护责任矩阵')
    make_table(doc,
        ['动作', '执行方', '频率'],
        [
            ['填报基础信息表', '供应商', '准入/变更时'],
            ['填报人力明细', '供应商', '每周一'],
            ['确认主管信息', '供应商', '每月5日'],
            ['登记离职', '供应商', '事件触发'],
            ['登记质检台账', '供应商', '日常'],
            ['合规自查', '供应商', '每月5日'],
            ['提交整改计划', '供应商', '约谈后48h'],
            ['更新整改进度', '供应商', '每周五'],
            ['发布整改通知', '京东侧', '事件触发'],
            ['发布通知', '京东侧', '按需'],
            ['汇总基础信息', '京东侧', '按需'],
        ],
        col_widths=[4, 2.5, 5]
    )

    add_h1(doc, '四、根目录汇总表 \u2014 供应商基础信息汇总')
    add_para(doc, '根目录的"供应商基础信息汇总表.xlsx"由京东侧维护，从各供应商的供应商基础信息表中汇总。')
    add_para(doc, '该表不发送给供应商，仅内部使用。', bold=True)

    add_h1(doc, '五、供应商空间文件夹结构')
    make_table(doc,
        ['文件夹', '维护方', '内容'],
        [
            ['01-人力管理', '供应商填写', '基础信息表、人力明细表、主管信息表、离职登记'],
            ['02-质检合规', '供应商填写', '质检台账、合规自查表、质检月报'],
            ['03-沟通留痕', '双方填写', '约谈记录、工作沟通、邮件归档'],
            ['04-整改跟踪', '供应商填写', '整改计划表、整改进度跟踪'],
            ['05-整改通知', '京东侧发布', '巡检整改通知（供应商只读）'],
            ['06-供应商自主使用', '供应商自选', '内部管理制度、培训材料等'],
            ['07-通知与公告', '京东侧发布', '制度更新、排名通报、重要通知'],
        ],
        col_widths=[3.5, 3, 8]
    )

    add_h1(doc, '六、数据来源分类')
    make_table(doc,
        ['来源', '数据', '采集方式'],
        [
            ['供应商填报', '人力、离职、质检台账、合规自查、整改', '按频率填报对应Excel'],
            ['质检团队', '质检月报', '推送至02-质检合规/质检月报/'],
            ['京东侧记录', '约谈记录、整改通知', '写入03-沟通留痕/、05-整改通知/'],
        ],
        col_widths=[3, 6, 5]
    )

    add_h1(doc, '七、分阶段推广计划')
    add_h2(doc, '第一期（试点）')
    add_list(doc, ['供应商基础信息表', '人力明细表', '整改计划表（触发时）'], numbered=False)
    add_h2(doc, '第二期（2-4周后）')
    add_list(doc, ['主管信息表', '离职原因登记表', '质检台账', '合规自查表', '整改进度跟踪'], numbered=False)

    add_h1(doc, '八、注意事项')
    add_list(doc, [
        '现场巡检（FCI评分、巡检表单）不在供应商空间内，这是京东侧内部使用的',
        '05-整改通知文件夹存放京东侧发布的巡检整改通知，供应商只能看到通知，不能修改',
        '03-沟通留痕中，约谈记录由京东侧填写，工作沟通双方都可以记录',
        '每个Excel模板都有数据验证（下拉选项）和示例数据，供应商按示例格式填写',
        '根目录的"供应商基础信息汇总表.xlsx"由京东侧从各供应商的基础信息表中汇总维护，不发送给供应商',
    ], numbered=False)

    add_h1(doc, '九、文件清单')
    add_list(doc, [
        '供应商基础信息汇总表.xlsx \u2014 根目录，京东侧维护的汇总',
        '[供应商名称]-工作空间/ \u2014 模板文件夹（复制后重命名）',
        '供应商工作空间填报指南.docx \u2014 给供应商看的填报说明',
        '供应商工作空间使用指南.docx \u2014 本文件，京东侧使用指南',
        '数据汇总表设计.md \u2014 数据汇总方案',
        'gen_templates.py \u2014 模板生成脚本',
    ], numbered=False)

    path = os.path.join(BASE, '供应商工作空间使用指南.docx')
    doc.save(path)
    print(f'OK: {path}')

if __name__ == '__main__':
    create_filling_guide()
    create_internal_guide()
    print('两个Word文档生成完成!')
