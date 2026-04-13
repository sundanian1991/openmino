#!/usr/bin/env python3
"""生成供应商分层分级管理办法 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# ── 默认样式 ──
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'SimSun'
font.size = Pt(12)
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style_normal.paragraph_format
pf.space_after = Pt(6)
pf.space_before = Pt(0)
pf.line_spacing = 1.5

for level, size in [(1, Pt(22)), (2, Pt(16)), (3, Pt(14))]:
    hs = doc.styles[f'Heading {level}']
    hf = hs.font
    hf.name = 'SimHei'
    hf.size = size
    hf.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    hs.paragraph_format.space_after = Pt(8)
    hs.paragraph_format.line_spacing = 1.5

# ── 辅助函数 ──
def cell_txt(cell, text, bold=False, size=Pt(10), align=None, east='宋体'):
    cell.text = ''
    p = cell.paragraphs[0]
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.size = size
    r.font.name = 'SimHei' if bold else 'SimSun'
    r.font.bold = bold
    r.element.rPr.rFonts.set(qn('w:eastAsia'), east)

def style_table(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tbl_pr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>')
    tbl_pr.append(borders)

def shade_row(cells, color='1A1A2E'):
    for c in cells:
        c._tc.get_or_add_tcPr().append(parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'))

def header_row(table, headers):
    for i, h in enumerate(headers):
        cell_txt(table.rows[0].cells[i], h, bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER, east='黑体')
    shade_row(table.rows[0].cells)
    for i in range(len(headers)):
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def fill_table(table, data, bold_cols=None, centers=None):
    """填充表格数据行。bold_cols={row_idx: [col_idxs]}, centers={row_idx: [col_idxs]}"""
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            is_bold = bold_cols and ri in bold_cols and ci in bold_cols[ri]
            is_center = centers and ri in centers and ci in centers[ri]
            align = WD_ALIGN_PARAGRAPH.CENTER if is_center else WD_ALIGN_PARAGRAPH.LEFT
            cell_txt(table.rows[ri + 1].cells[ci], val, bold=is_bold, size=Pt(9), align=align)

def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=Pt(3), sa=Pt(3)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = sb
    p.paragraph_format.space_after = sa
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.name = 'SimSun'; r.font.bold = bold
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if bold else '宋体')
    return p

def add_mixed(doc, segments, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=Pt(3), sa=Pt(3)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = sb
    p.paragraph_format.space_after = sa
    p.paragraph_format.line_spacing = 1.5
    for text, bold in segments:
        r = p.add_run(text)
        r.font.size = Pt(12); r.font.name = 'SimSun'; r.font.bold = bold
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if bold else '宋体')
    return p

def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.0 + indent * 0.75)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.name = 'SimSun'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pt)
    p.paragraph_format.line_height = Pt(pt)

# ═══════════ 封面 ═══════════
spacer(doc, 72)

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_after = Pt(12)
r = tp.add_run('供应商分层分级管理办法')
r.font.size = Pt(28); r.font.bold = True; r.font.name = 'SimHei'
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.paragraph_format.space_before = Pt(6); sp.paragraph_format.space_after = Pt(6)
r = sp.add_run('京东金融电销业务供应商管理 V2.0')
r.font.size = Pt(16); r.font.name = 'SimSun'; r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.paragraph_format.space_before = Pt(24)
r = dp.add_run('2026年4月')
r.font.size = Pt(14); r.font.name = 'SimSun'; r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

spacer(doc, 24)
lp = doc.add_paragraph()
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = lp.add_run('\u2500' * 40)
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

doc.add_page_break()

# ═══════════ 第一章 ═══════════
doc.add_heading('第一章 总则', level=1)
doc.add_heading('1.1 目的', level=2)
add_para(doc, '为建立科学的供应商管理体系，实现供应商的分类分级管理，提升供应商整体效能，特制定本办法。')

doc.add_heading('1.2 适用范围', level=2)
add_para(doc, '本办法适用于京东金融电销业务（金条、企金、信用卡、财富等产线）所有合作供应商。')

doc.add_heading('1.3 管理原则', level=2)
for bp, tp in [('客观公正：', '以数据为依据，定量为主、定性为辅'),
                ('动态管理：', '季度评估、动态升降级'),
                ('分类施策：', '不同级别差异化管理，资源向优质供应商倾斜'),
                ('限期改善：', 'C级供应商强制进入PIP，改善不合格启动退出')]:
    add_mixed(doc, [(bp, True), (tp, False)])

doc.add_page_break()

# ═══════════ 第二章 ═══════════
doc.add_heading('第二章 评估体系设计', level=1)
doc.add_heading('2.1 评估维度与权重', level=2)

t = doc.add_table(rows=5, cols=5, style='Table Grid'); style_table(t)
header_row(t, ['评估维度', '权重', '性质', '评估主体', '细项数量'])
fill_table(t, [
    ['业务督导', '60分', '定量', '业务督导岗', '6项'],
    ['供应商管理', '30分', '定性', '供应商管理岗', '6项'],
    ['质检', '10分', '定量', '质检培训运营岗', '4项'],
    ['合计', '100分', '\u2014', '\u2014', '16项'],
], bold_cols={3: [0,1,2,3,4]}, centers={r: [0,1,2,3,4] for r in range(4)})
spacer(doc, 6)

doc.add_heading('2.2 业务督导细项（60分，定量）', level=2)
t = doc.add_table(rows=7, cols=7, style='Table Grid'); style_table(t)
header_row(t, ['序号', '细项名称', '权重', '维度', '满分条件', '及格条件', '不及格条件'])
fill_table(t, [
    ['1', '季度目标达成率', '15分', '目标达成', '达成率\u2265100%', '85%\u2264达成率<100%', '达成率<85%'],
    ['2', '季度业绩贡献占比', '10分', '业绩产出', '人均产出\u2265均值110%', '90%\u2264人均产出<110%', '人均产出<90%'],
    ['3', '转化效率指数', '10分', '转化效率', '\u2265均值110%', '95%\u2264指数<110%', '指数<95%'],
    ['4', '月度业绩增长趋势', '10分', '趋势变化', '月均增长\u22655%', '-2%\u2264增长<5%', '月均增长<-2%'],
    ['5', '季度人力稳定贡献', '8分', '人力贡献', '稳定留存率\u226590%', '75%\u2264留存率<90%', '留存率<75%'],
    ['6', '有效人力占比', '7分', '人力贡献', '有效占比\u226595%', '85%\u2264有效占比<95%', '有效占比<85%'],
], centers={r: [0,1,2] for r in range(6)})
spacer(doc, 6)

add_para(doc, '评分规则：', bold=True, sa=Pt(2))
for r in ['满分与及格之间、及格与不及格之间按线性插值计算',
           '季度结束后第3个工作日锁定数据，第5个工作日出分',
           '出分后2个工作日内可发起数据异议，逾期视为认可',
           '新准入供应商季度运营不足1个月不参与当季排名']:
    add_bullet(doc, r)
spacer(doc, 6)

doc.add_heading('2.3 供应商管理细项（30分，定性）', level=2)
add_para(doc, '采用5分制评估（优秀5分/良好4分/合格3分/需改进2分/差1分），每项权重乘得分系数。', sa=Pt(6))

t = doc.add_table(rows=7, cols=4, style='Table Grid'); style_table(t)
header_row(t, ['序号', '细项名称', '权重', '评估标准（5分=优秀）'])
fill_table(t, [
    ['1', '配合意愿与响应速度', '6分', '需求2小时内响应，24小时内出方案，主动跟进，零推诿'],
    ['2', '管理团队能力', '5分', '驻场经理经验丰富，能独立解决现场问题，管理例会高质量'],
    ['3', '招聘与培训体系', '5分', '招聘渠道多元，月度满编率\u226595%，新人7天达标率\u226585%'],
    ['4', '合规与风险管理', '5分', '季度零合规事故，主动识别上报风险，合规承诺书100%'],
    ['5', '资源投入与稳定性', '5分', '资源充足，核心团队年度流失率<10%，产能波动<5%'],
    ['6', '创新与改进意愿', '4分', '季度主动提出\u22653条优化建议并被采纳，积极参与试点'],
], centers={r: [0,1,2] for r in range(6)})
spacer(doc, 6)

doc.add_heading('2.4 质检细项（10分，定量）', level=2)
t = doc.add_table(rows=5, cols=5, style='Table Grid'); style_table(t)
header_row(t, ['序号', '细项名称', '权重', '满分条件', '零分条件'])
fill_table(t, [
    ['1', '质检合格率', '4分', '合格率\u226595%', '合格率<85%'],
    ['2', '合规违规次数', '3分', '零A类违规且B类\u22642次', '出现A类违规（红线）'],
    ['3', '客户投诉率', '2分', '投诉率\u22643/万', '投诉率>8/万'],
    ['4', '话术规范执行率', '1分', '执行率\u226595%', '执行率<85%'],
], centers={r: [0,1,2] for r in range(4)})

doc.add_page_break()

# ═══════════ 第三章 ═══════════
doc.add_heading('第三章 ABC分级体系', level=1)
doc.add_heading('3.1 分级规则（相对排名法）', level=2)

t = doc.add_table(rows=4, cols=4, style='Table Grid'); style_table(t)
header_row(t, ['级别', '排名范围', '占比', '管理定位'])
fill_table(t, [
    ['A级', '排名前30%', '约30%', '战略合作伙伴'],
    ['B级', '排名30%-70%', '约40%', '稳定合作方'],
    ['C级', '排名后30%', '约30%', '重点管控对象'],
], centers={r: [0,1,2,3] for r in range(3)})
spacer(doc, 6)

add_para(doc, '为什么用相对排名而非绝对分数：', bold=True, sa=Pt(3))
for r in ['自动适应供应商总数变化', '与现有赛马排名逻辑一致',
           '避免绝对分数导致的级别固化', '保持适度的内部竞争张力']:
    add_bullet(doc, r)
spacer(doc, 6)

doc.add_heading('3.2 分级公布', level=2)
for a in ['每季度结束后第7个工作日公布分级结果',
           '分级结果以书面形式通知各供应商',
           '分级结果同步抄送服务组负责人（王易人）']:
    add_bullet(doc, a)

doc.add_page_break()

# ═══════════ 第四章 ═══════════
doc.add_heading('第四章 分级管理策略', level=1)

# 4.1 A级
doc.add_heading('4.1 A级供应商管理策略', level=2)
add_mixed(doc, [('管理定位：', True), ('战略合作伙伴 \u2014 少管多给，用信任和机会绑定', False)], sa=Pt(6))

t = doc.add_table(rows=7, cols=2, style='Table Grid'); style_table(t)
header_row(t, ['管理维度', '具体措施'])
data_a = [
    ['分量策略', '分量优先倾斜，新项目优先准入，单项目配额可放宽至40%上限'],
    ['沟通频率', '月度正式沟通 + 季度战略对话'],
    ['质检频次', '下调至50%（信任替代检查）'],
    ['结算政策', '结算账期优惠'],
    ['激励措施', '年度优秀供应商优先推荐、联合创新试点资格'],
    ['豁免政策', '豁免常规整改通知（问题直接沟通，不发正式文书）'],
]
fill_table(t, data_a, centers={r: [0] for r in range(6)}, bold_cols={r: [0] for r in range(6)})
spacer(doc, 6)

add_para(doc, '管理动作清单：', bold=True, sa=Pt(3))
for i, a in enumerate(['月度经营回顾会议（30分钟，聚焦数据和机会）',
    '季度战略合作对话（对齐半年目标、探讨新合作模式）',
    '季度评分结果一对一反馈（肯定成绩，指出潜在风险）',
    '年度联合规划（次年的目标、投入、预期收益）',
    '异常事件快速响应通道（专属对接人，2小时内响应）'], 1):
    add_bullet(doc, f'{i}. {a}')
spacer(doc, 6)

# 4.2 B级
doc.add_heading('4.2 B级供应商管理策略', level=2)
add_mixed(doc, [('管理定位：', True), ('稳定合作方 \u2014 常规管理 + 定向辅导，促其向上', False)], sa=Pt(6))

t = doc.add_table(rows=6, cols=2, style='Table Grid'); style_table(t)
header_row(t, ['管理维度', '具体措施'])
fill_table(t, [
    ['分量策略', '分量维持稳定，单项目配额20%-30%'],
    ['沟通频率', '双周正式沟通 + 月度评分反馈'],
    ['质检频次', '正常（100%覆盖）'],
    ['升级通道', '连续两期排名进入前30%自动升A级'],
    ['预警机制', '连续两期跌入后30%触发预警谈话'],
], centers={r: [0] for r in range(5)}, bold_cols={r: [0] for r in range(5)})
spacer(doc, 6)

add_para(doc, '管理动作清单：', bold=True, sa=Pt(3))
for i, a in enumerate(['双周运营回顾（检查关键指标、进度、问题）',
    '月度评分反馈（书面通知，标注各维度得分和排名变化）',
    '季度能力提升计划（针对弱项制定改进措施，跟踪执行）',
    '半年度综合评估（趋势分析：连续上升/下降/平稳）',
    '风险预警机制（排名连续下滑时提前介入）',
    '弱项专项辅导（质检差安排培训、配合度差约谈负责人）'], 1):
    add_bullet(doc, f'{i}. {a}')
spacer(doc, 6)

# 4.3 C级
doc.add_heading('4.3 C级供应商管理策略', level=2)
add_mixed(doc, [('管理定位：', True), ('重点管控对象 \u2014 重管严控，限期改善，否则退出', False)], sa=Pt(6))

t = doc.add_table(rows=7, cols=2, style='Table Grid'); style_table(t)
header_row(t, ['管理维度', '具体措施'])
fill_table(t, [
    ['分量策略', '分量严格压缩，存量配额降至15%以下，核心业务逐步剥离'],
    ['新项目', '全面暂停准入'],
    ['沟通频率', '周度正式沟通 + 日度数据跟踪'],
    ['质检频次', '翻倍（200%覆盖，重点抽查）'],
    ['结算政策', '结算账期延长（改善后恢复）'],
    ['强制措施', '强制进入PIP流程（1个月改善期）'],
], centers={r: [0] for r in range(6)}, bold_cols={r: [0] for r in range(6)})
spacer(doc, 6)

add_para(doc, '管理动作清单：', bold=True, sa=Pt(3))
for i, a in enumerate(['周度改善进度会（对照PIP目标逐项检查，形成书面纪要）',
    '日度数据跟踪（业绩、质检、合规关键指标日报）',
    '供应商负责人约谈（C级确认后3个工作日内完成）',
    '根因分析报告（要求供应商5个工作日内提交）',
    '改善计划审核（审核供应商提交的改善方案是否可行）',
    '风险预案准备（同步准备替代供应商）',
    '存量项目风险评估（评估交接难度和风险点）'], 1):
    add_bullet(doc, f'{i}. {a}')

doc.add_page_break()

# ═══════════ 第五章 ═══════════
doc.add_heading('第五章 PIP（绩效改进计划）', level=1)

doc.add_heading('5.1 触发条件', level=2)
add_para(doc, '季度评分落入C级（排名后30%）即自动触发PIP，无需额外审批。')

doc.add_heading('5.2 PIP流程（30个自然日）', level=2)
t = doc.add_table(rows=6, cols=3, style='Table Grid'); style_table(t)
header_row(t, ['阶段', '时间', '核心动作'])
fill_table(t, [
    ['启动与根因诊断', '第1-3天', '发出PIP通知书、召开启动会、根因诊断、确定改善目标、签署确认书'],
    ['改善计划制定', '第4-7天', '供应商提交改善计划、我方审核、确认最终版、分量压缩、建立日度跟踪'],
    ['密集执行跟踪', '第8-21天', '周度改善会、日度指标跟踪、过程纠偏、中期评估、保留改善证据'],
    ['终期评估', '第22-28天', '期末评分、改善效果对比、综合判断、编写评估报告'],
    ['结果判定执行', '第29-30天', '输出判定结果、正式通报、执行对应措施、归档全流程文档'],
], centers={r: [0,1] for r in range(5)})
spacer(doc, 6)

doc.add_heading('5.3 改进目标', level=2)
for t_item in ['总分提升\u226510分', '核心失分项（导致C级的维度）改善\u226550%',
                '排名脱离后30% 或 评分绝对值\u226570分']:
    add_bullet(doc, t_item)
spacer(doc, 6)

doc.add_heading('5.4 验收标准', level=2)
t = doc.add_table(rows=4, cols=3, style='Table Grid'); style_table(t)
header_row(t, ['结果', '条件', '处理'])
fill_table(t, [
    ['通过', '期末评分\u226570分 且 排名脱离后30% 且 核心失分项改善\u226550%', '解除PIP，升级为B级'],
    ['有条件通过', '期末评分60-69分 或 核心失分项改善未达50%', '延长1个月观察期'],
    ['不通过', '期末评分<60分 或 评分提升<5分 或 拒绝执行', '启动退出流程'],
])
spacer(doc, 6)

doc.add_heading('5.5 PIP结果执行', level=2)
add_mixed(doc, [('通过：', True), ('解除PIP，升级为B级管理策略，分量逐步恢复，次月恢复正常质检频次。', False)])
add_mixed(doc, [('有条件通过：', True), ('延长1个月观察期，观察期内按B-管理（分量恢复50%、双周沟通），观察期末达标则正式升B，不达标则降回C并启动清退评估。', False)])
add_mixed(doc, [('不通过：', True), ('', False)], sa=Pt(2))
for i, s in enumerate(['约谈供应商高管，明确告知公司制度和行业影响',
    '给予7天考虑期，引导主动退出',
    '若7天内未主动退出，正式启动清退流程（30天内完成）',
    '该供应商6个月内不接受重新准入'], 1):
    add_bullet(doc, f'{i}. {s}')

doc.add_page_break()

# ═══════════ 第六章 ═══════════
doc.add_heading('第六章 特殊场景处理', level=1)

doc.add_heading('6.1 新准入供应商', level=2)
add_para(doc, '新准入不足一个季度的供应商：', sa=Pt(3))
for item in ['不纳入ABC分级，设置"观察期"管理', '首个季度独立排名，不与存量供应商混排',
    '观察期内按B-级标准管理（双周沟通、正常质检、分量从10%起步）',
    '观察期末给出首次正式评分，参与下一次ABC分级',
    '若观察期表现极差（首次评分低于50分），可直接启动退出流程']:
    add_bullet(doc, item)
spacer(doc, 6)

doc.add_heading('6.2 连续两个季度C级', level=2)
for item in ['不再进入标准PIP流程，直接进入加速退出程序',
    '第二次数确认为C级后3个工作日内发出"加速退出通知书"',
    '给予5个工作日考虑期，引导主动退出',
    '若5个工作日内未主动退出，立即启动正式清退（15天内完成）',
    '该供应商永久不得重新准入', '同时启动替代供应商遴选']:
    add_bullet(doc, item)
spacer(doc, 6)

doc.add_heading('6.3 A级骤降至C级', level=2)
add_para(doc, '属于重大异常事件，处理方式与常规C级不同：', sa=Pt(3))
add_mixed(doc, [('1. ', True), ('立即启动异常诊断', True), ('（5个工作日内完成根因分析）', False)])
add_mixed(doc, [('2. ', True), ('区分原因', True), ('：', False)], sa=Pt(2))
add_bullet(doc, '系统性原因（项目调整、政策变化、人员大规模流失）\u2192 给予"保护期"1个月，保护期内不压缩分量、不启动PIP', indent=1)
add_bullet(doc, '供应商自身原因（管理懈怠、配合度下降、战略转移）\u2192 给予"警告期"2周，高管层面约谈', indent=1)
add_bullet(doc, '3. 必须向上级（王易人）报备，部门内通报（不点名）')
add_bullet(doc, '4. 恢复机制：恢复正常后不直接回A级，先降至B级观察1个季度')

doc.add_page_break()

# ═══════════ 第七章 ═══════════
doc.add_heading('第七章 与现有机制的衔接', level=1)

doc.add_heading('7.1 与赛马排名的对接', level=2)
t = doc.add_table(rows=4, cols=2, style='Table Grid'); style_table(t)
header_row(t, ['机制', '对接方式'])
fill_table(t, [
    ['数据共享', '赛马4维度数据直接映射到100分体系对应细项，避免重复打分'],
    ['趋势联动', '赛马排名连续2期垫底，即使总分未到C级线，也触发预警谈话'],
    ['快速通道', '赛马排名连续2期第一，总分在A级临界值附近（前35%），可破格升A'],
], centers={r: [0] for r in range(3)}, bold_cols={r: [0] for r in range(3)})
spacer(doc, 6)

doc.add_heading('7.2 与集中度管控的对接', level=2)
t = doc.add_table(rows=4, cols=2, style='Table Grid'); style_table(t)
header_row(t, ['级别', '单项目配额上限'])
fill_table(t, [
    ['A级', '35%-40%（接近上限作为激励）'],
    ['B级', '20%-30%'],
    ['C级', '\u226415%（强制压缩）'],
], centers={r: [0] for r in range(3)}, bold_cols={r: [0] for r in range(3)})
spacer(doc, 6)

doc.add_heading('7.3 与清退流程的对接', level=2)
t = doc.add_table(rows=5, cols=3, style='Table Grid'); style_table(t)
header_row(t, ['触发条件', '清退周期', '备注'])
fill_table(t, [
    ['PIP不通过+7天考虑期满', '30天', '标准清退'],
    ['连续两季度C级+5天考虑期满', '15天', '加速退出'],
    ['重大合规事故', '15天', '快速通道'],
    ['供应商主动申请', '30天', '自愿退出'],
])

doc.add_page_break()

# ═══════════ 第八章 ═══════════
doc.add_heading('第八章 附则', level=1)

doc.add_heading('8.1 数据管理', level=2)
for item in ['所有评分数据由业务组统一收集、审核、归档',
    '供应商对评分结果有异议的，可在出分后2个工作日内提出数据异议',
    '异议由业务组在3个工作日内复核并答复']:
    add_bullet(doc, item)
spacer(doc, 6)

doc.add_heading('8.2 制度修订', level=2)
for item in ['本办法由数据科技业务部服务组负责解释',
    '根据实际执行情况，每季度评估修订一次']:
    add_bullet(doc, item)
spacer(doc, 6)

doc.add_heading('8.3 生效时间', level=2)
for item in ['本办法自发布之日起生效',
    '原有供应商评估规则与本办法不一致的，以本办法为准']:
    add_bullet(doc, item)

# ── 保存 ──
output = '/Users/sundanian/Documents/projects/ai-agents/my-agent/plans/2026-04-13-供应商分层分级体系/供应商分层分级管理办法.docx'
doc.save(output)
print(f'Done: {output}')
