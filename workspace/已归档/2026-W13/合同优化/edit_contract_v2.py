#!/usr/bin/env python3
"""
编辑客户服务外包合同：在附件四保密协议中插入新增保密条款。

段落索引（已确认）：
- [348] "保密信息"不应包括 → 修改1在这之前插入
- [357] §2(e) 第三方存储限制 → 修改2/4/5在这之后插入
- [367] 以上违约金不足以弥补 → 修改3在这之后插入（层次化+禁令）
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import re

INPUT = "/Users/sundanian/Documents/projects/ai-agents/my-agent/workspace/⏳待处理/合同优化/客户服务外包合同-202511.docx.docx"
OUTPUT = "/Users/sundanian/Documents/projects/ai-agents/my-agent/workspace/⏳待处理/合同优化/客户服务外包合同-202511-修改版.docx"

doc = Document(INPUT)

# ── 工具函数 ──────────────────────────────────────────────

def find_para(doc, text_fragment, start=0, end=None):
    """找到包含指定文本片段的段落索引"""
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, min(end, len(doc.paragraphs))):
        if text_fragment in doc.paragraphs[i].text:
            return i
    return -1

def clone_paragraph_after(doc, ref_para_idx, text):
    """在指定段落之后插入一个新段落，复制其样式和格式"""
    ref_para = doc.paragraphs[ref_para_idx]

    # 克隆段落元素
    new_para_elem = deepcopy(ref_para._element)

    # 清空所有 run 内容
    for child in list(new_para_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'hyperlink'):
            new_para_elem.remove(child)

    # 创建新 run
    r_elem = OxmlElement('w:r')

    # 复制 run properties
    ref_runs = ref_para.runs
    if ref_runs:
        ref_rpr = ref_runs[0]._element.find(qn('w:rPr'))
        if ref_rpr is not None:
            new_rpr = deepcopy(ref_rpr)
            r_elem.append(new_rpr)

    # 添加文本
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    r_elem.append(t_elem)
    new_para_elem.append(r_elem)

    # 插入到 ref_para 之后
    ref_para._element.addnext(new_para_elem)

    return new_para_elem

def insert_multiple_after(doc, start_idx, texts):
    """在start_idx段落之后依次插入多段文字，返回最后插入的段落索引"""
    current_idx = start_idx
    for text in texts:
        clone_paragraph_after(doc, current_idx, text)
        # 新段落插入后，在document.paragraphs中的索引是 current_idx+1
        current_idx = current_idx + 1
    return current_idx

# ── 定位关键段落 ──────────────────────────────────────────

# 在附件四区域(段落343+)查找
ANNEX4_START = 343

idx_s1_exclude = find_para(doc, '不应包括以下信息和资料', ANNEX4_START)
if idx_s1_exclude == -1:
    idx_s1_exclude = find_para(doc, '不应包括', ANNEX4_START + 5)
idx_s2e = find_para(doc, '接收方不得将在交易过程中获取的披露方的任何资料和数据存储在第三方', ANNEX4_START)
idx_s3_damages = find_para(doc, '以上违约金不足以弥补损失的', ANNEX4_START)
idx_s3f = find_para(doc, '接收方因其违约向披露方按照上述约定支付违约金', ANNEX4_START)
idx_s4 = find_para(doc, '所有保密信息是且应是披露方的专属财产', ANNEX4_START)

print(f"§1(b) 排除项位置: {idx_s1_exclude}")
print(f"§2(e) 第三方存储: {idx_s2e}")
print(f"§3(e) 违约金不足: {idx_s3_damages}")
print(f"§3(f) 继续履行: {idx_s3f}")
print(f"§4 其他: {idx_s4}")

# ── 修改1: 保密信息范围细化 ──
# 在 "保密信息不应包括" 之前插入细化内容
print("\n=== 修改1: 保密信息范围细化 ===")

if idx_s1_exclude > 0:
    # 在排除项之前插入
    texts_m1 = [
        '前述保密信息包括但不限于以下类别：',
        '(i) 项目信息：项目需求、技术方案、实施计划、预算、报价、成本结构、利润水平；',
        '(ii) 商业信息：甲方的外呼策略、市场计划、运营数据、供应商名单、客户信息；',
        '(iii) 技术信息：系统功能描述、算法逻辑、专有营销工具；',
        '(iv) 任何被标明为"保密"或虽未标明但依其性质应被合理认定为保密的信息。',
    ]
    # 从排除项的前一个段落开始插入
    insert_multiple_after(doc, idx_s1_exclude - 1, texts_m1)
    print(f"  在段落 {idx_s1_exclude} 前插入了 {len(texts_m1)} 段")

# 重新定位（插入后索引偏移了）
offset_m1 = len(texts_m1) if idx_s1_exclude > 0 else 0

idx_s2e = find_para(doc, '接收方不得将在交易过程中获取的披露方的任何资料和数据存储在第三方', ANNEX4_START)
idx_s3_damages = find_para(doc, '以上违约金不足以弥补损失的', ANNEX4_START)
idx_s3f = find_para(doc, '接收方因其违约向披露方按照上述约定支付违约金', ANNEX4_START)

print(f"\n偏移后 §2(e): {idx_s2e}, §3(e)违约金: {idx_s3_damages}, §3(f): {idx_s3f}")

# ── 修改2+4+5: 在 §2(e) 之后插入 (f)(g)(h)(i)(j) ──
print("\n=== 修改2+4+5: 信息流程控制+审计+不竞争+分包 ===")

if idx_s2e > 0:
    texts_m245 = [
        '(f) 接收方仅能为履行本协议项下交易之目的，将保密信息披露给其必须知悉该信息的雇员、高管（"需知"原则），并确保这些人员受到同等保密义务的约束。接收方应与所有获得或有权访问保密信息的代表订立书面协议，其保密限制内容不得低于本协议的约定。',
        '(g) 未经披露方事前书面同意，接收方不得将本协议项下交易的任何信息（包括但不限于合作事实本身、项目内容、成果）用于任何其他用途，包括但不限于向第三方（尤其是同行竞对公司）进行业务推介、投标、报价或作为案例宣传。如接收方拟将本协议项下的信息或成果用于其他项目的投标文件，必须至少提前【10】个工作日将该投标文件的相关部分提交给披露方审核。仅在获得披露方明确的书面同意后，方可使用。',
        '(h) 披露方有权不定期对接收方就本协议项下信息的安全保护措施进行审计和检查，接收方应予以积极配合，并提供必要的便利和资料。',
        '(i) 在本协议有效期内及本协议终止或解除后【6】个月内，接收方不得为披露方的直接竞争对手提供与本协议项下相同或类似的服务。',
        '(j) 未经披露方书面同意，接收方不得将本协议项下的核心业务分包给任何第三方。',
    ]
    insert_multiple_after(doc, idx_s2e, texts_m245)
    print(f"  在段落 {idx_s2e} 后插入了 {len(texts_m245)} 段")

# 重新定位
offset_m245 = len(texts_m245)
idx_s3_damages = find_para(doc, '以上违约金不足以弥补损失的', ANNEX4_START)
idx_s3f = find_para(doc, '接收方因其违约向披露方按照上述约定支付违约金', ANNEX4_START)

print(f"\n偏移后 §3(e)违约金: {idx_s3_damages}, §3(f): {idx_s3f}")

# ── 修改3: 在 §3(e) 违约金条款后补充层次化 + 禁令救济 ──
print("\n=== 修改3: 多层次法律后果 + 禁令救济 ===")

if idx_s3_damages > 0:
    texts_m3 = [
        '就前述违约金，双方进一步约定如下层次化标准：(I) 一般违约：违反保密义务但未造成重大损失的（如无意泄露），违约金为人民币10万元，或按本协议项下月结算总金额的10%计算，以较高者为准；(II) 严重违约：接收方将保密信息透露给披露方的直接竞争对手，或利用该信息为竞争对手服务，从而给披露方造成重大损失的，违约金为人民币100万元。',
        '接收方同意，一旦发生违约披露，金钱赔偿可能不足以救济。因此，披露方有权申请法院禁令，要求接收方立即停止违约行为，并采取必要措施防止损害扩大。',
    ]
    insert_multiple_after(doc, idx_s3_damages, texts_m3)
    print(f"  在段落 {idx_s3_damages} 后插入了 {len(texts_m3)} 段")

# ── 保存 ──
print(f"\n=== 保存到: {OUTPUT} ===")
doc.save(OUTPUT)
print("完成！")

# ── 验证：打印修改后的附件四段落 ──
print("\n=== 验证修改后的附件四 ===")
doc2 = Document(OUTPUT)
annex4_title = -1
for i, p in enumerate(doc2.paragraphs):
    if '保密协议' in p.text and i > 300 and '附件' not in p.text:
        if p.text.strip() == '保密协议':
            annex4_title = i
            break

if annex4_title > 0:
    for i in range(annex4_title, min(annex4_title + 60, len(doc2.paragraphs))):
        text = doc2.paragraphs[i].text.strip()
        if text:
            print(f"  [{i:3d}] {text[:130]}")
