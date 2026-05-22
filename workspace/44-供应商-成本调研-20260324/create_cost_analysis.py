#!/usr/bin/env python3
"""
小号开卡成本分析文档生成器
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_cost_analysis_doc():
    """创建成本分析Word文档"""
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)

    # 标题
    title = doc.add_heading('小号开卡方式成本分析', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # 副标题
    subtitle = doc.add_paragraph('基于200元套餐的实际成本测算')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # 空行

    # 一、方案基本信息
    doc.add_heading('一、方案基本信息', level=1)
    doc.add_paragraph('1. 套餐规格：200元/月，含2000分钟通话时长', style='List Bullet')
    doc.add_paragraph('2. 开卡费用：100元/卡（一次性）', style='List Bullet')
    doc.add_paragraph('3. 人员配置：每人2张卡（实体卡+虚拟卡）', style='List Bullet')
    doc.add_paragraph('4. 日均通话：50-60分钟（实体卡占比80%，虚拟卡占比20%）', style='List Bullet')

    doc.add_paragraph()

    # 二、成本计算过程
    doc.add_heading('二、成本计算过程', level=1)

    p = doc.add_paragraph()
    p.add_run('1. 名义单价计算：').bold = True
    doc.add_paragraph('   200元 ÷ 2000分钟 = 0.10元/分钟')

    p = doc.add_paragraph()
    p.add_run('2. 实际使用情况（按日均55分钟计算）：').bold = True
    doc.add_paragraph('   月通话时长：55分钟 × 30天 = 1650分钟')
    doc.add_paragraph('   - 实体卡：1650 × 80% = 1320分钟')
    doc.add_paragraph('   - 虚拟卡：1650 × 20% = 330分钟')

    p = doc.add_paragraph()
    p.add_run('3. 月度成本构成：').bold = True
    doc.add_paragraph('   套餐费：200元 × 2卡 = 400元/月')
    doc.add_paragraph('   开卡费分摊：100元 × 2卡 ÷ 6个月 = 33元/月')
    doc.add_paragraph('   月度总成本：400 + 33 = 433元/月')

    p = doc.add_paragraph()
    p.add_run('4. 实际每分钟成本：').bold = True
    doc.add_paragraph('   433元 ÷ 1650分钟 = 0.26元/分钟')

    doc.add_paragraph()

    # 三、成本对比汇总
    doc.add_heading('三、成本对比汇总', level=1)

    # 创建表格
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '数值'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    data = [
        ('名义单价', '0.10元/分钟'),
        ('实际单价', '0.26元/分钟'),
        ('成本倍率', '2.6倍'),
        ('月度总成本', '433元/人'),
        ('月通话时长', '1650分钟'),
        ('套餐利用率', '41%'),
    ]

    for i, (key, val) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = key
        row_cells[1].text = val
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 四、风险说明
    doc.add_heading('四、风险说明', level=1)

    p = doc.add_paragraph()
    p.add_run('1. 实体卡限制：').bold = True
    doc.add_paragraph('   - 单人名下最多5张卡')
    doc.add_paragraph('   - 跨省使用政策存在地区差异')

    p = doc.add_paragraph()
    p.add_run('2. 虚拟卡风险：').bold = True
    doc.add_paragraph('   - 投诉罚款：1500元/次')
    doc.add_paragraph('   - 30张卡可免除1次投诉处罚')

    p = doc.add_paragraph()
    p.add_run('3. 标记风险：').bold = True
    doc.add_paragraph('   - 号码被标记后提供消除途径')
    doc.add_paragraph('   - 需关注标记频率对业务的影响')

    doc.add_paragraph()

    # 五、核心结论
    doc.add_heading('五、核心结论', level=1)

    doc.add_paragraph('1. 实际成本（0.26元/分钟）是名义单价（0.10元/分钟）的2.6倍')
    doc.add_paragraph('2. 主要成本损耗来源：')
    doc.add_paragraph('   - 套餐利用率低（仅41%）：双卡配置导致每张卡用量不足', style='List Bullet 2')
    doc.add_paragraph('   - 开卡费摊销：一次性成本按月分摊', style='List Bullet 2')
    doc.add_paragraph('   - 虚拟卡投诉风险：需预留风险准备金', style='List Bullet 2')
    doc.add_paragraph('3. 优化建议：')
    doc.add_paragraph('   - 提高单卡使用率，考虑减少虚拟卡配置', style='List Bullet 2')
    doc.add_paragraph('   - 或寻找按实际用量计费的模式', style='List Bullet 2')

    # 保存文档
    output_path = '/Users/sundanian/Documents/projects/ai-agents/my-agent/workspace/小号开卡成本分析.docx'
    doc.save(output_path)
    print(f"Word文档已保存: {output_path}")
    return output_path

def create_svg_chart():
    """创建成本对比SVG图表"""
    svg_content = '''<?xml version="1.0" encoding="UTF-8"?>
<svg width="600" height="400" viewBox="0 0 600 400" xmlns="http://www.w3.org/2000/svg">
  <!-- Background -->
  <rect width="600" height="400" fill="#fafafa"/>

  <!-- Title -->
  <text x="300" y="30" text-anchor="middle" font-size="18" font-weight="bold" fill="#333">
    成本对比分析
  </text>

  <!-- Chart Area -->
  <g transform="translate(80, 60)">
    <!-- Y Axis -->
    <line x1="0" y1="0" x2="0" y2="280" stroke="#ccc" stroke-width="1"/>
    <!-- X Axis -->
    <line x1="0" y1="280" x2="480" y2="280" stroke="#ccc" stroke-width="1"/>

    <!-- Y Axis Labels -->
    <text x="-10" y="280" text-anchor="end" font-size="11" fill="#666">0</text>
    <text x="-10" y="224" text-anchor="end" font-size="11" fill="#666">0.1</text>
    <text x="-10" y="168" text-anchor="end" font-size="11" fill="#666">0.2</text>
    <text x="-10" y="112" text-anchor="end" font-size="11" fill="#666">0.3</text>
    <text x="-10" y="56" text-anchor="end" font-size="11" fill="#666">0.4</text>
    <text x="-10" y="10" text-anchor="end" font-size="11" fill="#666">元/分钟</text>

    <!-- Grid lines -->
    <line x1="0" y1="224" x2="480" y2="224" stroke="#eee" stroke-width="1"/>
    <line x1="0" y1="168" x2="480" y2="168" stroke="#eee" stroke-width="1"/>
    <line x1="0" y1="112" x2="480" y2="112" stroke="#eee" stroke-width="1"/>
    <line x1="0" y1="56" x2="480" y2="56" stroke="#eee" stroke-width="1"/>

    <!-- Bar 1: 名义单价 -->
    <rect x="60" y="252" width="80" height="28" fill="#70a8d8" rx="4"/>
    <text x="100" y="245" text-anchor="middle" font-size="12" font-weight="bold" fill="#333">0.10</text>
    <text x="100" y="305" text-anchor="middle" font-size="12" fill="#666">名义单价</text>

    <!-- Bar 2: 实际单价 -->
    <rect x="220" y="173" width="80" height="107" fill="#c25030" rx="4"/>
    <text x="260" y="165" text-anchor="middle" font-size="12" font-weight="bold" fill="#333">0.26</text>
    <text x="260" y="305" text-anchor="middle" font-size="12" fill="#666">实际单价</text>

    <!-- Bar 3: 对比线 -->
    <line x1="140" y1="266" x2="220" y2="266" stroke="#666" stroke-width="1" stroke-dasharray="4,2"/>
    <text x="180" y="260" text-anchor="middle" font-size="10" fill="#c25030">+160%</text>

    <!-- Key metrics boxes -->
    <g transform="translate(340, 80)">
      <rect x="0" y="0" width="120" height="60" fill="#fff" stroke="#ddd" rx="6"/>
      <text x="60" y="22" text-anchor="middle" font-size="11" fill="#666">成本倍率</text>
      <text x="60" y="45" text-anchor="middle" font-size="20" font-weight="bold" fill="#c25030">2.6x</text>
    </g>

    <g transform="translate(340, 160)">
      <rect x="0" y="0" width="120" height="60" fill="#fff" stroke="#ddd" rx="6"/>
      <text x="60" y="22" text-anchor="middle" font-size="11" fill="#666">套餐利用率</text>
      <text x="60" y="45" text-anchor="middle" font-size="20" font-weight="bold" fill="#c25030">41%</text>
    </g>

    <g transform="translate(340, 240)">
      <rect x="0" y="0" width="120" height="60" fill="#fff" stroke="#ddd" rx="6"/>
      <text x="60" y="22" text-anchor="middle" font-size="11" fill="#666">月度成本</text>
      <text x="60" y="45" text-anchor="middle" font-size="20" font-weight="bold" fill="#333">¥433</text>
    </g>
  </g>

  <!-- Footer note -->
  <text x="300" y="380" text-anchor="middle" font-size="10" fill="#999">
    数据来源：200元套餐 + 100元开卡费 + 双卡配置
  </text>
</svg>'''

    output_path = '/Users/sundanian/Documents/projects/ai-agents/my-agent/workspace/成本对比图表.svg'
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(svg_content)
    print(f"SVG图表已保存: {output_path}")
    return output_path

if __name__ == '__main__':
    # 激活虚拟环境
    import subprocess
    subprocess.run(['source', '.venv/bin/activate'], shell=True, executable='/bin/bash')

    doc_path = create_cost_analysis_doc()
    svg_path = create_svg_chart()

    print(f"\n文件已生成:")
    print(f"  - Word文档: {doc_path}")
    print(f"  - SVG图表: {svg_path}")
