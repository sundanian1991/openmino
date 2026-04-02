#!/usr/bin/env python3
"""完整检查DOCX段落结构"""
from docx import Document

INPUT = "/Users/sundanian/Documents/projects/ai-agents/my-agent/workspace/⏳待处理/合同优化/客户服务外包合同-202511.docx.docx"
doc = Document(INPUT)

print(f"总段落数: {len(doc.paragraphs)}")
print(f"总表格数: {len(doc.tables)}")
print()

# 打印所有段落
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else 'None'
    if text:  # 只打印非空段落
        print(f"[{i:3d}] ({style:20s}) {text[:150]}")
