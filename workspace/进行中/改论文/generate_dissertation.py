#!/usr/bin/env python3
"""
Fan Hui Dissertation - Temple University Format
Generates a properly formatted DOCX according to G1-G4 templates
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=1.0):
    """Set paragraph spacing"""
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing

def add_centered_paragraph(doc, text, bold=False, font_size=12, space_before=0, space_after=0):
    """Add a centered paragraph"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    set_paragraph_spacing(p, space_before, space_after)
    return p

def add_title_page(doc):
    """G1: Title Page with Inverted Pyramid"""
    # Title in inverted pyramid format
    add_centered_paragraph(doc, "DIGITAL TECHNOLOGIES", bold=True, font_size=14, space_after=0)
    add_centered_paragraph(doc, "AND RURAL FINANCE", bold=True, font_size=14, space_after=72)

    add_centered_paragraph(doc, "A Dissertation", font_size=12, space_before=24, space_after=12)
    add_centered_paragraph(doc, "Submitted to", font_size=12, space_after=12)
    add_centered_paragraph(doc, "the Temple University Graduate Board", font_size=12, space_after=12)
    add_centered_paragraph(doc, "In Partial Fulfillment", font_size=12, space_after=12)
    add_centered_paragraph(doc, "of the Requirements for the Degree", font_size=12, space_after=12)
    add_centered_paragraph(doc, "DOCTOR OF SCIENCE", bold=True, font_size=12, space_after=12)
    add_centered_paragraph(doc, "by", font_size=12, space_before=24, space_after=12)
    add_centered_paragraph(doc, "Hui Fan", font_size=12, space_after=12)
    add_centered_paragraph(doc, "May 2026", font_size=12, space_after=48)

    # Committee members
    add_centered_paragraph(doc, "Examining Committee Members:", font_size=12, space_after=12)
    add_centered_paragraph(doc, "Oleg Rytchkov, Advisory Chair, Finance", font_size=12, space_after=6)
    add_centered_paragraph(doc, "Xiaohui Gao Bakshi, Finance", font_size=12, space_after=6)
    add_centered_paragraph(doc, "Lalitha Naveen, Finance", font_size=12, space_after=6)
    add_centered_paragraph(doc, "Jayanthi Krishnan, External Reader, Accounting", font_size=12, space_after=6)

    doc.add_page_break()

def add_copyright_page(doc):
    """Copyright page"""
    # Center copyright vertically
    for _ in range(10):
        doc.add_paragraph()

    add_centered_paragraph(doc, "©", font_size=12, space_after=12)
    add_centered_paragraph(doc, "Copyright", font_size=12, space_after=12)
    add_centered_paragraph(doc, "2026", font_size=12, space_after=12)
    add_centered_paragraph(doc, "by", font_size=12, space_after=12)
    add_centered_paragraph(doc, "Hui Fan", font_size=12, space_after=12)

    doc.add_page_break()

def add_abstract(doc):
    """Abstract page"""
    add_centered_paragraph(doc, "ABSTRACT", bold=True, font_size=14, space_after=24)

    abstract_text = """Since the reform and opening-up policy, especially since the 18th National Congress of the Communist Party of China, China's rural financial reform has continuously deepened. It has actively played a role in supporting agriculture, rural areas, and farmers, achieving positive results in supporting the precise implementation of poverty alleviation strategies, realizing the comprehensive development of rural areas, and advancing comprehensive rural revitalization. The 20th National Congress of the Communist Party of China proposed the goal of "comprehensively promoting rural revitalization" and the need to "improve the rural financial service system," providing further guidance for the reform and development of rural finance in the new era.

Digital rural development is a strategic direction for rural revitalization, and the exploration of the integration of digital technology and rural finance holds significant importance in practical financial services for rural revitalization. Long-term observation shows that the application of digital technology in rural finance can effectively address long-standing issues in rural financial development, such as high transaction costs, information asymmetry, and difficulties in financing.

In this dissertation, centering on the application of digital technology in rural financial development, we firstly analyze the challenges and reasons for the difficulties in rural financial development in China. Through empirical research, the paper concludes that there is a positive correlation between the level of rural financial development and the coverage of rural internet users in various provinces across the country.

Key words: Digitalization; rural finance; rural credit; Rural Revitalization"""

    p = doc.add_paragraph(abstract_text)
    p.paragraph_format.line_spacing = 2.0  # Double spacing
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_page_break()

def add_dedication(doc):
    """Dedication page"""
    for _ in range(8):
        doc.add_paragraph()

    add_centered_paragraph(doc, "To my family and mentors", font_size=12, space_after=12)
    add_centered_paragraph(doc, "who have supported me throughout this journey", font_size=12, space_after=12)

    doc.add_page_break()

def add_acknowledgments(doc):
    """Acknowledgments page"""
    add_centered_paragraph(doc, "ACKNOWLEDGMENTS", bold=True, font_size=14, space_after=24)

    text = """Time flies, and my years of studying at GFD are coming to an end. I want to express my gratitude to the teachers for their unwavering guidance throughout these years. Your rich knowledge, rigorous academic attitude, and dedication to excellence in your work are all exemplary qualities that I need to continue learning throughout my life.

Secondly, I would like to thank the teachers at the Wudaokou School of Finance for your rigorous academic attitude and hard work, which have enabled us to gain knowledge and apply it. I also sincerely thank the teachers and fellow students who helped me during the process of writing my thesis.

Finally, I want to express my gratitude to numerous mentors, fellow students, and friends for their selfless help. Please accept my heartfelt thanks. In the days ahead, I will continue my journey with the wealth of knowledge and spirit that my teachers and fellow students have given me."""

    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 2.0
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_page_break()

def add_toc(doc):
    """Table of Contents"""
    add_centered_paragraph(doc, "TABLE OF CONTENTS", bold=True, font_size=14, space_after=24)

    toc_entries = [
        ("ABSTRACT", "iii"),
        ("DEDICATION", "iv"),
        ("ACKNOWLEDGMENTS", "v"),
        ("TABLE OF CONTENTS", "vi"),
        ("LIST OF TABLES", "vii"),
        ("LIST OF FIGURES", "viii"),
        ("", ""),
        ("CHAPTER", ""),
        ("1. INTRODUCTION", "1"),
        ("2. REVIEW OF LITERATURE", "7"),
        ("3. THEORETICAL BASIS", "17"),
        ("4. MAJOR EXPERIENCE IN DEVELOPMENT OF RURAL FINANCE IN CHINA AND ABROAD", "28"),
        ("5. CURRENT DEVELOPMENT SITUATION OF RURAL FINANCE IN CHINA", "34"),
        ("6. EMPIRICAL STUDY ON DRIVING FUNCTION OF INTERNET DIGITAL INCLUSION", "48"),
        ("7. THE CURRENT STATUS OF THE APPLICATION OF DIGITAL TECHNOLOGY", "55"),
        ("8. RECOMMENDATIONS FOR THE APPLICATION OF DIGITAL TECHNOLOGY", "66"),
        ("9. RESEARCH CONCLUSION AND CONTRIBUTION", "75"),
        ("", ""),
        ("REFERENCES", "78"),
    ]

    for entry, page in toc_entries:
        if entry == "" and page == "":
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        if entry == "CHAPTER":
            p.add_run(entry).bold = True
        else:
            # Add tab between entry and page
            p.add_run(entry)
            if page:
                p.add_run("\t" + page)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    doc.add_page_break()

def add_chapter_heading(doc, chapter_num, title):
    """G2: Chapter heading format"""
    # Page break before chapter
    doc.add_page_break()

    # Chapter number line
    add_centered_paragraph(doc, f"CHAPTER {chapter_num}", bold=True, font_size=12, space_after=24)

    # Chapter title
    add_centered_paragraph(doc, title.upper(), bold=True, font_size=12, space_after=48)

def main():
    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set margins (1 inch = 2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

    # G1: Preliminary Pages
    add_title_page(doc)
    add_copyright_page(doc)
    add_abstract(doc)
    add_dedication(doc)
    add_acknowledgments(doc)
    add_toc(doc)

    # G2: Chapters (placeholders for structure)
    chapters = [
        ("1", "INTRODUCTION"),
        ("2", "REVIEW OF LITERATURE"),
        ("3", "THEORETICAL BASIS"),
        ("4", "MAJOR EXPERIENCE IN DEVELOPMENT OF RURAL FINANCE IN CHINA AND ABROAD"),
        ("5", "CURRENT DEVELOPMENT SITUATION OF RURAL FINANCE IN CHINA"),
        ("6", "EMPIRICAL STUDY ON DRIVING FUNCTION OF INTERNET DIGITAL INCLUSION ON DEVELOPMENT OF RURAL FINANCE"),
        ("7", "THE CURRENT STATUS OF THE APPLICATION OF DIGITAL TECHNOLOGY IN RURAL FINANCIAL DEVELOPMENT"),
        ("8", "RECOMMENDATIONS FOR THE APPLICATION OF DIGITAL TECHNOLOGY IN RURAL FINANCIAL DEVELOPMENT"),
        ("9", "RESEARCH CONCLUSION AND CONTRIBUTION"),
    ]

    for num, title in chapters:
        add_chapter_heading(doc, num, title)
        p = doc.add_paragraph(f"[Content for Chapter {num}: {title} to be inserted here]")
        p.paragraph_format.line_spacing = 2.0

    # G3: References
    doc.add_page_break()
    add_centered_paragraph(doc, "REFERENCES", bold=True, font_size=14, space_after=24)
    p = doc.add_paragraph("[References to be inserted here with hanging indent]")
    p.paragraph_format.line_spacing = 2.0

    # G4: Appendices
    doc.add_page_break()
    add_centered_paragraph(doc, "APPENDIX A", bold=True, font_size=12, space_after=12)
    add_centered_paragraph(doc, "RAW DATA", bold=True, font_size=12, space_after=24)
    p = doc.add_paragraph("[Appendix content to be inserted here]")
    p.paragraph_format.line_spacing = 2.0

    # Save document
    output_path = "/Users/sundanian/Documents/projects/ai-agents/my-agent/workspace/进行中/改论文/Fan_Hui_Dissertation_Temple_Format.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    main()
