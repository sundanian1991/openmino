#!/usr/bin/env python3
"""Generate Q1 开门红颁奖站点选择方案 Word document"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Page Setup ───
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# ─── Helper: set cell shading ───
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

# ─── Helper: add formatted table ───
def add_styled_table(data, col_widths=None, header_color='1A5A90'):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if i == 0:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
                    else:
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            if i == 0:
                set_cell_shading(cell, header_color)
    return table

# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Q1 开门红颁奖\n站点选择与活动方案')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x90)
run.bold = True
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('数据科技业务部 · 服务组')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('编制日期：2026年4月22日\n编制人：服务组')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. 背景与目的
# ═══════════════════════════════════════════════════════════

h1 = doc.add_heading('一、背景与目的', level=1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x90)

p = doc.add_paragraph()
p.add_run('为表彰2026年第一季度各业务线"开门红"优秀团队和个人，服务组计划前往供应商站点举行集中颁奖活动。本次活动覆盖').font.size = Pt(11)
p.add_run('企业金融、信用卡、金条、财富').font.size = Pt(11).bold = True
p.add_run('四条业务线，采用"按业务分批出行"的方式，逐站完成颁奖。').font.size = Pt(11)

doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════
# 2. 供应商总体情况
# ═══════════════════════════════════════════════════════════

h1 = doc.add_heading('二、供应商总体情况', level=1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x90)

p = doc.add_paragraph('当前合作供应商共计34家，站点108个，覆盖全国五大区域。各业务线站点分布如下：')
p.runs[0].font.size = Pt(11)

add_styled_table([
    ['业务线', '站点数', 'Q1总人力', '主要供应商', '已承办站点'],
    ['企业金融', '42', '约760人', '中乾/汇讯/德辰', '海腾·西安、锦瑞·定州'],
    ['信用卡', '8', '约500人', '创博/速讯达/中弘', '创博·承德、速讯达·沈阳'],
    ['金条', '56', '约1,400人', '毅航/毛毛虫/伽玛', '毛毛虫·合肥、毅航·合肥'],
    ['财富', '4', '约230人', '毅航/毛毛虫/嘉信', '毛毛虫·合肥、毅航·合肥'],
], header_color='1A5A90')

doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════
# 3. 站点选择方案
# ═══════════════════════════════════════════════════════════

h1 = doc.add_heading('三、站点选择方案', level=1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x90)

p = doc.add_paragraph('选择原则：优先选择人力规模大、尚未承办过的供应商站点；兼顾地理分布，便于分批出行。')
p.runs[0].font.size = Pt(11)

# ─── 3.1 企业金融 ───
h2 = doc.add_heading('3.1 企业金融', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['选项', '供应商', '城市', 'Q1人力', '站点数', '状态', '推荐理由'],
    ['推荐', '中乾', '洛阳', '143人', '8个', '本期承办', '企金最大规模供应商，8个职场覆盖最广'],
    ['备选', '汇讯', '多站点', '120人', '4个', '未承办', '跨多业务线，人力较分散'],
    ['备选', '德辰', '多站点', '103人', '4个', '未承办', '多站点分布，单站规模偏小'],
    ['已承办', '海腾', '西安', '47人', '1个', '不重复', 'Q1已承办过'],
    ['已承办', '锦瑞', '定州', '91人', '3个', '不重复', 'Q1已承办过'],
], header_color='1A5A90')

doc.add_paragraph('')
p = doc.add_paragraph('推荐承办方：中乾（洛阳）。站点规模大（143人，8个职场），覆盖面广，颁奖仪式气派，能集中展示企金业务成果。')
p.runs[0].font.size = Pt(11)
p.runs[0].bold = True

# ─── 3.2 信用卡 ───
h2 = doc.add_heading('3.2 信用卡', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['选项', '供应商', '城市', 'Q1人力', '站点数', '状态', '推荐理由'],
    ['推荐', '中弘', '沈阳', '68人', '4个', '本期承办', '除已承办外最大信用卡供应商，4个职场集中'],
    ['备选', '靖涵', '多站点', '58人', '3个', '未承办', '规模较小，可作为备选'],
    ['已承办', '创博', '承德', '183人', '4个', '不重复', 'Q1已承办过'],
    ['已承办', '速讯达', '沈阳', '95人', '2个', '不重复', 'Q1已承办过，可在沈阳回访'],
], header_color='1A5A90')

doc.add_paragraph('')
p = doc.add_paragraph('推荐承办方：中弘（沈阳）。中弘在沈阳有4个站点，68人规模适合集中颁奖。速讯达也在沈阳，可顺访回访。')
p.runs[0].font.size = Pt(11)
p.runs[0].bold = True

# ─── 3.3 金条 ───
h2 = doc.add_heading('3.3 金条', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['选项', '供应商', '城市', 'Q1人力', '站点数', '状态', '推荐理由'],
    ['方案A', '广达', '西安', '67人(金条)', '8个', '候选', '金条67人在正定/宿豫/鸡西，西安仅企金80人'],
    ['方案B', '伽玛', '亳州', '279人', '6个', '本期承办', '金条第三大，A类优质，利辛站点126人'],
    ['备选', '赛维斯', '石家庄', '189人', '6个', '未承办', '规模大，但石家庄区域覆盖偏北'],
    ['备选', '岐力', '南昌', '176人', '2个', '未承办', '江西区域，与合肥有一定距离'],
    ['已承办', '毛毛虫', '合肥', '257人', '4个', '不重复', 'Q1已承办过'],
    ['已承办', '毅航', '合肥', '308人', '4个', '不重复', 'Q1已承办过'],
], header_color='1A5A90')

doc.add_paragraph('')
p = doc.add_paragraph('重要说明：广达西安站点仅有企金业务（80人），金条业务67人分布在正定、宿豫、鸡西，不在西安。')
p.runs[0].font.size = Pt(11)
p.runs[0].bold = True
p.runs[0].font.color.rgb = RGBColor(0xC2, 0x50, 0x30)

doc.add_paragraph('')
p = doc.add_paragraph('方案A（合并西安）：金条与企金、财富三业务线合并颁奖，在广达西安会场统一举办，金条获奖者远程参与或集中表彰。')
p.runs[0].font.size = Pt(11)

p = doc.add_paragraph('方案B（独立金条）：伽玛亳州（利辛站点，126人）独立承办。伽玛为A类优质供应商，总人力279人，金条业务规模大，颁奖仪式独立且有气势。')
p.runs[0].font.size = Pt(11)

# ─── 3.4 财富 ───
h2 = doc.add_heading('3.4 财富', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['选项', '供应商', '城市', 'Q1人力', '站点数', '状态', '推荐理由'],
    ['推荐', '嘉信', '西安', '71人', '2个', '本期承办', '财富唯一有规模的非承办供应商'],
    ['备选', '汇讯', '单站点', '28人', '1个', '未承办', '人力偏少，场地可能不够气派'],
    ['已承办', '毛毛虫', '合肥', '59人', '1个', '不重复', 'Q1已承办过'],
    ['已承办', '毅航', '合肥', '78人', '1个', '不重复', 'Q1已承办过'],
], header_color='1A5A90')

doc.add_paragraph('')
p = doc.add_paragraph('推荐承办方：嘉信（西安）。嘉信为财富业务非承办供应商中规模最大的（71人），西安作为西北中心城市，交通便利。')
p.runs[0].font.size = Pt(11)
p.runs[0].bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 4. 出行批次方案
# ═══════════════════════════════════════════════════════════

h1 = doc.add_heading('四、出行批次方案', level=1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x90)

h2 = doc.add_heading('方案A：三站走完全程（推荐合并）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['批次', '城市', '业务', '承办方', 'Q1人力', '天数', '说明'],
    ['第1站', '洛阳', '企金', '中乾', '143人', '1天', '高铁直达，当天可返'],
    ['第2站', '沈阳', '信用卡', '中弘', '68人', '1-2天', '飞沈阳，可顺访速讯达回访'],
    ['第3站', '西安', '金条+财富', '广达+嘉信', '151人', '1天', '合并颁奖，一场覆盖两条业务线'],
], header_color='2E8B6E')

doc.add_paragraph('')
p = doc.add_paragraph('优势：只跑3个城市，西安一场合并金条+财富颁奖，节省差旅成本。广达虽无金条团队，但可作为企金供应商承办联合颁奖。')
p.runs[0].font.size = Pt(11)

h2 = doc.add_heading('方案B：四站独立颁奖（金条独立）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['批次', '城市', '业务', '承办方', 'Q1人力', '天数', '说明'],
    ['第1站', '洛阳', '企金', '中乾', '143人', '1天', '高铁直达，当天可返'],
    ['第2站', '沈阳', '信用卡', '中弘', '68人', '1-2天', '飞沈阳，可顺访速讯达回访'],
    ['第3站', '亳州', '金条', '伽玛', '279人', '2天', '飞合肥转亳州，规模大仪式感强'],
    ['第4站', '西安', '财富', '嘉信', '71人', '1天', '飞西安，可顺访海腾回访'],
], header_color='2E8B6E')

doc.add_paragraph('')
p = doc.add_paragraph('优势：金条独立颁奖，仪式感更强；伽玛279人规模大、A类优质供应商，配得上独立颁奖。四条业务线各有一次专属颁奖，供应商感受更好。')
p.runs[0].font.size = Pt(11)
p.runs[0].bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 5. 颁奖活动流程
# ═══════════════════════════════════════════════════════════

h1 = doc.add_heading('五、颁奖活动流程', level=1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x90)

h2 = doc.add_heading('5.1 活动前（准备阶段）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['序号', '环节', '负责人', '时间节点', '说明'],
    ['1', '确定颁奖站点', '服务组', 'T-21天', '与各供应商确认档期和意愿'],
    ['2', '确认颁奖名单', '服务组+数据', 'T-14天', '各业务线Q1排名/业绩数据出炉'],
    ['3', '奖杯/证书制作', '行政', 'T-10天', '按业务线+供应商定制'],
    ['4', '活动流程确定', '服务组', 'T-7天', '每个站点的详细议程'],
    ['5', '物料准备', '行政', 'T-5天', '横幅、PPT、签到表、伴手礼'],
    ['6', '供应商沟通', '服务组', 'T-5天', '提前同步活动安排，让对方准备场地'],
    ['7', '差旅安排', '行政', 'T-3天', '机票/高铁/酒店预订'],
], header_color='6F8660')

doc.add_paragraph('')

h2 = doc.add_heading('5.2 活动当天（每个站点）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['时段', '环节', '时长', '内容'],
    ['到达', '到场签到', '15分钟', '到达供应商场地，与负责人寒暄'],
    ['开场', '领导致辞', '10分钟', '供应商负责人开场，介绍团队'],
    ['回顾', 'Q1业绩回顾', '15分钟', '展示该供应商Q1数据（排名、贡献、亮点）'],
    ['颁奖', '颁奖仪式', '20分钟', '按业务线逐一颁奖，合影留念'],
    ['交流', '座谈交流', '30分钟', '听取供应商反馈、困难、诉求'],
    ['展望', 'Q2目标对齐', '15分钟', '传达Q2方向和期望'],
    ['结束', '合影留念+用餐', '30分钟', '大合影，团队聚餐（可选）'],
], header_color='6F8660')

doc.add_paragraph('')

h2 = doc.add_heading('5.3 活动后（跟进阶段）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['序号', '环节', '说明'],
    ['1', '新闻稿/内部通报', '在部门群或内部平台发布颁奖消息'],
    ['2', '照片整理归档', '照片存档到共享盘或workspace'],
    ['3', '供应商反馈整理', '座谈中提到的问题和诉求，形成跟进清单'],
    ['4', 'Q2目标追踪', '把颁奖时对齐的目标纳入日常追踪'],
], header_color='6F8660')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 6. 各站点详细信息
# ═══════════════════════════════════════════════════════════

h1 = doc.add_heading('六、各承办站点详细信息', level=1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x90)

# 中乾
h2 = doc.add_heading('6.1 中乾 · 洛阳（企金承办）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['项目', '信息'],
    ['供应商全称', '中乾（具体全称待确认）'],
    ['承办业务', '企业金融（企业主贷、企业金采等）'],
    ['Q1平均人力', '143人'],
    ['站点数量', '8个职场'],
    ['区域分布', '洛阳为核心区域'],
    ['推荐理由', '企金最大规模供应商，8个职场覆盖面广'],
], header_color='1A5A90')

doc.add_paragraph('')

# 中弘
h2 = doc.add_heading('6.2 中弘 · 沈阳（信用卡承办）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['项目', '信息'],
    ['供应商全称', '中弘（具体全称待确认）'],
    ['承办业务', '信用卡（菁卡拉新、商家白条等）'],
    ['Q1平均人力', '68人'],
    ['站点数量', '4个职场'],
    ['城市', '辽宁省沈阳市'],
    ['推荐理由', '除已承办外最大信用卡供应商，4个职场集中'],
    ['备注', '速讯达也在沈阳，可顺访回访'],
], header_color='1A5A90')

doc.add_paragraph('')

# 伽玛
h2 = doc.add_heading('6.3 伽玛 · 亳州（金条承办，方案B）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['项目', '信息'],
    ['供应商全称', '湖北伽玛科技有限公司'],
    ['承办业务', '金条、金企融合'],
    ['Q1总人力', '279人'],
    ['站点数量', '6个职场'],
    ['主会场推荐', '利辛站点（亳州）· 126人'],
    ['站点地址', '安徽省亳州市利辛县城关镇前进路78号交通局西侧二楼'],
    ['其他站点', '樊城(襄阳)54人、宜城20人、廊坊80人'],
    ['供应商分层', 'A类：优质供应商'],
    ['推荐理由', '金条第三大供应商，A类优质，规模大仪式感强'],
], header_color='1A5A90')

doc.add_paragraph('')

# 嘉信
h2 = doc.add_heading('6.4 嘉信 · 西安（财富承办）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['项目', '信息'],
    ['供应商全称', '嘉信（具体全称待确认）'],
    ['承办业务', '财富'],
    ['Q1平均人力', '71人'],
    ['站点数量', '2个职场'],
    ['城市', '陕西省西安市'],
    ['推荐理由', '财富唯一有规模的非承办供应商'],
    ['备注', '海腾也在西安（企金已承办），可一天内顺访'],
], header_color='1A5A90')

doc.add_paragraph('')

# 广达（方案A）
h2 = doc.add_heading('6.5 广达 · 西安（金条+财富合并，方案A）', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x8B, 0x6E)

add_styled_table([
    ['项目', '信息'],
    ['供应商全称', '陕西广达信息服务有限公司'],
    ['承办业务', '企金（西安站点）、金条+财富（联合颁奖）'],
    ['西安站点人力', '80人（企金），金条67人不在西安'],
    ['金条业务分布', '正定48人、宿豫7人、鸡西12人'],
    ['站点地址', '陕西省西安市雁塔区含光路南段216号嘉翔大厦12楼'],
    ['供应商分层', 'B级供应商（近期从C升至B）'],
    ['备注', '西安无金条团队，需考虑金条颁奖的落地方式'],
], header_color='1A5A90')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 7. 风险与建议
# ═══════════════════════════════════════════════════════════

h1 = doc.add_heading('七、风险与建议', level=1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x90)

risks = [
    ('广达西安无金条团队', '广达金条67人分布在正定、宿豫、鸡西，西安站点仅有企金业务。如选择方案A，金条颁奖需在广达西安会场以"联合颁奖"形式进行，金条获奖者可能无法全部到场。建议：如金条颁奖需要全员参与的仪式感，选择方案B（伽玛亳州）。'),
    ('财富供应商规模偏小', '财富业务供应商总数较少，嘉信71人为最大非承办供应商，但相比金条的伽玛（279人）、毅航（335人），规模偏小。建议：与金条合并颁奖或在西安顺访海腾（企金已承办方）增加活动丰富度。'),
    ('供应商档期冲突', '4个承办站点需协调供应商老板到场时间，建议提前21天确认档期，预留备选供应商。'),
    ('Q1数据未最终确认', '颁奖名单依赖Q1业绩数据，需确保数据在T-14天前出炉并核对无误。'),
    ('差旅成本控制', '方案B需跑4个城市，方案A仅需3个。如预算紧张，优先考虑方案A。'),
]

for i, (title, desc) in enumerate(risks, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'风险{i}：{title}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC2, 0x50, 0x30)
    p2 = doc.add_paragraph(desc)
    p2.runs[0].font.size = Pt(11)
    doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════
# 8. 待确认事项
# ═══════════════════════════════════════════════════════════

h1 = doc.add_heading('八、待确认事项', level=1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x90)

add_styled_table([
    ['序号', '待确认事项', '负责人', '完成时限'],
    ['1', '中乾供应商全称及洛阳具体站点地址', '服务组', 'T-21天前'],
    ['2', '中弘供应商全称及沈阳具体站点地址', '服务组', 'T-21天前'],
    ['3', '嘉信供应商全称及西安具体站点地址', '服务组', 'T-21天前'],
    ['4', '金条颁奖方案选择（A合并西安 vs B独立亳州）', '服务组负责人', 'T-21天前'],
    ['5', '各承办供应商老板档期确认', '服务组', 'T-14天前'],
    ['6', 'Q1各业务线排名及颁奖名单', '数据组', 'T-14天前'],
    ['7', '差旅预算审批', '行政', 'T-10天前'],
    ['8', '奖杯/证书设计及制作确认', '行政', 'T-10天前'],
], header_color='C26D3A')

doc.add_paragraph('')
doc.add_paragraph('')

# ─── Footer ───
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 方案完 —')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Save
output_path = '/Users/sundanian/Documents/projects/ai-agents/my-agent/workspace/Q1开门红颁奖站点选择方案.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
