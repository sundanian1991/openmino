#!/usr/bin/env python3
"""Generate AI产品定价模式演变商业案例报告.docx"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/Users/sundanian/Documents/projects/ai-agents/my-agent/projects/AI产品定价模式演变-从CodingPlan到TokenPlan_ppt169_20260425/AI产品定价模式演变商业案例报告.docx"
ACCENT = RGBColor(0xB8, 0x54, 0x50)
HEADING = RGBColor(0x33, 0x33, 0x33)
BODY = RGBColor(0x55, 0x55, 0x55)
DIVIDER = RGBColor(0xD0, 0xD0, 0xD0)
BG = RGBColor(0xF9, 0xF9, 0xF9)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Helvetica Neue'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
style.font.size = Pt(10.5)
style.font.color.rgb = BODY
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.5

# Override Title style
title_style = doc.styles['Title']
title_style.font.name = 'Helvetica Neue'
title_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
title_style.font.size = Pt(26)
title_style.font.bold = True
title_style.font.color.rgb = HEADING
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_style.paragraph_format.space_after = Pt(12)

# Override Heading 1
h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Helvetica Neue'
h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
h1_style.font.size = Pt(16)
h1_style.font.bold = True
h1_style.font.color.rgb = HEADING
h1_style.paragraph_format.space_before = Pt(24)
h1_style.paragraph_format.space_after = Pt(12)

# Override Heading 2
h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Helvetica Neue'
h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
h2_style.font.size = Pt(13)
h2_style.font.bold = True
h2_style.font.color.rgb = ACCENT
h2_style.paragraph_format.space_before = Pt(18)
h2_style.paragraph_format.space_after = Pt(9)

# Override Heading 3
h3_style = doc.styles['Heading 3']
h3_style.font.name = 'Helvetica Neue'
h3_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
h3_style.font.size = Pt(11)
h3_style.font.bold = True
h3_style.font.color.rgb = HEADING
h3_style.paragraph_format.space_before = Pt(12)
h3_style.paragraph_format.space_after = Pt(6)

# Set page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_spacer():
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_para(text, bold=False, size=10.5, color=BODY, align=WD_ALIGN_PARAGRAPH.LEFT, indent=0, after=6, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    run = p.add_run(text)
    run.font.name = 'Helvetica Neue'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    return p

def add_mixed_para(parts, align=WD_ALIGN_PARAGRAPH.LEFT, after=6, before=0, indent=0):
    """parts: list of (text, bold, size, color)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    for text, bold, size, color in parts:
        run = p.add_run(text)
        run.font.name = 'Helvetica Neue'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color
    return p

def add_table(headers, rows, col_widths=None):
    """Add a styled table. headers: list of str, rows: list of list of str"""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    if col_widths:
        for idx, w in enumerate(col_widths):
            table.columns[idx].width = Pt(w)

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell.paragraphs[0].runs[0].font.color.rgb = HEADING
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        # Background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F9F9F9')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            cell.paragraphs[0].runs[0].font.color.rgb = BODY
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    add_spacer()
    return table

def add_page_break():
    doc.add_page_break()

def add_numbered_list(items, indent=0):
    """items: list of (text, bold, size, color) tuples for each line"""
    for idx, (text, bold, size, color) in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(indent + 20)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{idx}. ")
        run.font.name = 'Helvetica Neue'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run = p.add_run(text)
        run.font.name = 'Helvetica Neue'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color

# ===== COVER PAGE =====
for _ in range(8):
    add_para("")
add_para("AI产品定价模式演变", bold=True, size=26, color=HEADING, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
add_para("从 Coding Plan 到 Token Plan", size=16, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
add_para("商业案例研究报告", size=12, color=BODY, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
for _ in range(3):
    add_para("")
add_para("日期：2026年4月25日", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
add_para("版本：V1.0", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
add_para("分类：商业分析 / 定价策略", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)

add_page_break()

# ===== EXECUTIVE SUMMARY =====
doc.add_heading("执行摘要", level=1)

add_mixed_para([
    ("核心结论：", True, 10.5, HEADING),
    ("AI大模型定价正从「卖工具订阅」演进为「卖算力资源包」，这个过程精确重演了电信行业从话费套餐到流量套餐的20年历史。", False, 10.5, BODY),
    ("Coding Plan是过渡产品，Token Plan是终态。", True, 10.5, HEADING),
], after=12)

add_para("能从中学到的不是具体的定价数字，而是三件事：", after=6)

add_mixed_para([
    ("计量单位即定价权", True, 10.5, HEADING),
    (" — 谁定义计量单位，谁就掌握定价权和利润空间", False, 10.5, BODY),
], after=3)
add_mixed_para([
    ("补贴-收紧-分层的三段式节奏", True, 10.5, HEADING),
    (" — 所有互联网赛道的定价都走这三步", False, 10.5, BODY),
], after=3)
add_mixed_para([
    ("入口产品到生态产品的升级路径", True, 10.5, HEADING),
    (" — 用高频刚需获客，用多场景提高ARPU", False, 10.5, BODY),
], after=12)

add_mixed_para([
    ("关键数据：", True, 10.5, HEADING),
    ("同样¥40-50/月，不同计量方式下实际可用量差异巨大：", False, 10.5, BODY),
], after=6)

add_table(
    ["平台", "月费", "计量", "实际等效"],
    [
        ["智谱 GLM Lite\n¥49", "¥49/月", "Prompt\n5h窗口", "每天80次提问"],
        ["MiniMax Starter\n¥29", "¥29/月", "Prompt\n5h窗口", "每5h 40次，无月上限"],
        ["阿里百炼 Lite\n¥40", "¥40/月", "API请求", "~20-120次提问/天"],
        ["腾讯云 Token Lite\n¥39", "¥39/月", "Token", "~70轮深度问答"],
    ],
    col_widths=[140, 80, 80, 140]
)

add_mixed_para([
    ("一个开发者8个月Claude Code消耗", False, 10.5, BODY),
    ("100亿Token", True, 10.5, ACCENT),
    ("（API价$15,000+），Max plan只花$800，节省93%。", False, 10.5, BODY),
], after=3)
add_para("Token构成：文件读取35-45%，代码生成仅5-15%。", after=12)

add_page_break()

# ===== PART 1: FACTS =====
doc.add_heading("一、事实层：发生了什么", level=1)

doc.add_heading("1.1 时间线与关键事件", level=2)
add_table(
    ["时间", "事件", "商业信号"],
    [
        ["2024Q4-2025Q2", "DeepSeek带头降价，百万Token从几十元降到几毛钱", "价格战启动，抢用户规模"],
        ["2025Q4", "智谱GLM率先推出Coding Plan（固定月费）", "从按量付费转向订阅制"],
        ["2025H2", "首购优惠取消、低档停售、限时抢购", "补贴收紧，利润压力显现"],
        ["2026-01-28", "Kimi全面从Prompt计量切换为Token计量", "计量单位统一化趋势"],
        ["2026-02", "OpenClaw爆火（GitHub 355K stars）", "Agent自动循环→Token消耗6-10倍增长"],
        ["2026-03", 'MiniMax将"Coding Plan"更名为"Token Plan"', "从编程工具升级为通用算力"],
        ["2026-03-05", "腾讯云最后入局，六大平台全部上线", "市场全面铺开"],
        ["2026-04", "阿里百炼升级企业版Token Plan，¥198/698/1398", "云平台跟进Token Plan趋势"],
        ["2026-04-03", "小米MiMo Token Plan上线¥39-659", "新入局者极低价破局"],
        ['2026-04-23', 'MiMo-V2.5-Pro公测，官方批评"低价是陷阱"', "低价入局后预示收紧"],
        ["2026-04-24", "DeepSeek V4发布，1.6T参数，成本暴降73%", "开源重塑定价天花板"],
    ],
    col_widths=[90, 220, 130]
)

doc.add_heading("1.2 参与者格局", level=2)
add_para("两条路线并行：", after=3)
add_mixed_para([
    ("模型厂商（专卖店）：", True, 10.5, HEADING),
    ("智谱GLM、Kimi、MiniMax、小米MiMo — 只卖自家模型，深度优化", False, 10.5, BODY),
], after=3)
add_mixed_para([
    ("云平台（超市）：", True, 10.5, HEADING),
    ("阿里云百炼、火山方舟、腾讯云 — 聚合多模型，一键切换", False, 10.5, BODY),
], after=12)

doc.add_heading("1.3 五种计量体系并存", level=2)
add_table(
    ["计量方式", "采用者", "换算关系", "透明度"],
    [
        ["API请求次数", "阿里、火山、腾讯", "1提问=5-30次API调用", "低"],
        ["Prompt次数", "智谱、MiniMax早期", "1Prompt≈15-20次模型调用", "中"],
        ["Token计量", "Kimi、MiniMax", "直接按输入输出Token计费", "高"],
        ["Credit计量", "Cursor、MiMo、百炼", "统一Credit池，不同模型不同倍率", "中"],
        ["时间窗口配额", "智谱5h、Claude 5h+7d", "限制节奏而非总量", "中"],
    ],
    col_widths=[100, 100, 160, 60]
)
add_mixed_para([
    ("关键发现：", True, 10.5, ACCENT),
    ("同样¥40-50/月，不同计量方式下实际可用量差异巨大。直接比数字毫无意义，90%开发者踩坑。", False, 10.5, BODY),
], after=12)

doc.add_heading("1.4 OpenClaw：Token消耗的催化剂", level=2)
add_para("OpenClaw作为开源AI Agent框架（GitHub 355K stars），直接催热了Coding Plan需求：", after=3)
add_numbered_list([
    ("MiniMax M2 token用量：2月vs12月增长6倍，Coding Plan消耗增长10倍+", False, 10.5, BODY),
    ("阿里云算力告急，限量供应 → Lite档4-13停售的直接原因", False, 10.5, BODY),
    ("Anthropic 2026-04从标准订阅移除OpenClaw访问 — 成本撑不住了", False, 10.5, BODY),
])

doc.add_heading("1.5 芯片供给线：算力成本的底层变量", level=2)
add_para("英伟达在中国退缩：市占率66%(2024) → 54%(2025)，2025-04全面禁止AI GPU对华出口。", after=6)
add_mixed_para([
    ("华为昇腾崛起：", True, 10.5, HEADING),
], after=3)
add_numbered_list([
    ("2025年累计交付81.2万片，市占率~20%，国产第一", False, 10.5, BODY),
    ("950PR：FP4算力1.56 PFLOPS（H20的2.87倍），定价~7万元（H200的1/3），4月量产", False, 10.5, BODY),
    ("2026目标：市占率50%，出货75-100万片，营收超800亿", False, 10.5, BODY),
])

doc.add_heading("1.6 DeepSeek V4定价冲击", level=2)
add_table(
    ["模型", "总参数", "激活参数", "输入（缓存）", "输入（未缓存）", "输出"],
    [
        ["V4-Pro", "1.6T", "49B", "1元/百万", "12元/百万", "24元/百万"],
        ["V4-Flash", "284B", "13B", "0.2元/百万", "1元/百万", "2元/百万"],
        ["对比GPT-5", "-", "-", "$5/百万", "-", "$15/百万"],
    ],
    col_widths=[80, 60, 60, 80, 80, 80]
)

doc.add_heading("1.7 IPO/资本化进程与定价策略的关联", level=2)
add_mixed_para([
    ("智谱GLM（2513.HK）：", True, 10.5, HEADING),
    ("2026-01-08港交所上市，市值511亿港元。Coding Plan上线两个月付费开发者超15万，ARR快速破亿元。", False, 10.5, BODY),
], after=6)
add_mixed_para([
    ("MiniMax（MMX）：", True, 10.5, HEADING),
    ("成立到IPO仅4年，全球最快AI上市公司。2025营收7904万美元+159%，海外收入占比73%，毛利率转正25.4%。", False, 10.5, BODY),
], after=6)
add_mixed_para([
    ("含义：", True, 10.5, ACCENT),
    ("上市公司必须证明商业可持续性 → Coding Plan/Token Plan是可见的收入增长引擎。定价不能太激进也不能太保守。", False, 10.5, BODY),
], after=12)

add_page_break()

# ===== PART 2: FRAMEWORK =====
doc.add_heading("二、框架层：为什么这么演变", level=1)

doc.add_heading("2.1 类比框架：电信套餐演进史", level=2)
add_table(
    ["电信阶段", "AI对应阶段", "核心逻辑"],
    [
        ["2009话费套餐（充100用不完）", "2024 API按量计费（百万Token几十元）", "供给过剩，用户感知不到成本"],
        ["2013流量套餐（9.9几G）", "2025 Coding Plan（¥29-200/月）", "用量激增，需要包月降低感知"],
        ['2017"不限量"达量降速', "2026分层套餐（极速版/普通版）", "精细化分层，按速度/质量定价"],
        ["2020按GB精打细算", "2026+ Token精细计量", "用户开始算账，厂商利润回归"],
    ],
    col_widths=[170, 160, 150]
)

doc.add_heading("2.2 五力分析", level=2)
add_table(
    ["力量", "分析"],
    [
        ["供应商议价力", "算力（GPU）是刚性成本。NVIDIA占85-90%市场，算力成本传导到Token定价"],
        ["买方议价力", "开发者选择多（7+平台），切换成本低（API Key一换就行），议价力强"],
        ["替代品威胁", "开源模型（DeepSeek、Qwen）+本地部署=潜在替代。压制定价上限"],
        ["新进入者威胁", "低（需要百亿级融资+算力），但小米MiMo等新玩家仍在进入"],
        ["行业内竞争", "极高。7家在一个池子里打价格战，首月优惠从7.9元起"],
    ],
    col_widths=[120, 360]
)

doc.add_heading("2.3 定价模式演进的三层逻辑", level=2)
add_mixed_para([
    ("第一层 成本驱动：", True, 10.5, HEADING),
    ("算力成本 → Token成本 → 套餐定价", False, 10.5, BODY),
], after=6)
add_mixed_para([
    ("第二层 竞争驱动：", True, 10.5, HEADING),
    ("价格战 → 补贴圈用户 → 收紧利润", False, 10.5, BODY),
], after=6)
add_mixed_para([
    ("第三层 价值驱动：", True, 10.5, HEADING),
    ("从「卖Token」到「卖工作流」到「卖生态」", False, 10.5, BODY),
], after=12)

add_page_break()

# ===== PART 3: INSIGHTS =====
doc.add_heading("三、洞察层：能提炼什么", level=1)

doc.add_heading("3.1 计量单位即定价权", level=2)
add_para("谁定义计量单位，谁就掌握定价权。MiniMax从Prompt→Token：从「模糊打包」到「透明计费」，争夺用户信任。云厂商用「API请求次数」：数字看起来大，但实际消耗多，信息不对称有利于厂商。", after=6)
add_mixed_para([
    ("可复用模式：", True, 10.5, ACCENT),
    ("推出新产品时，计量单位的选择决定了用户感知和利润空间。模糊计量有利于早期获客，透明计量有利于长期信任。", False, 10.5, BODY),
], after=12)

doc.add_heading("3.2 补贴-收紧-分层的三段式节奏", level=2)
add_table(
    ["阶段", "目的", "AI案例"],
    [
        ["补贴", "抢用户规模，建立习惯", "2024-2025：百万Token几毛钱，首月7.9元"],
        ["收紧", "测试价格弹性，筛选付费用户", "2025H2：限购、停售低档、限时抢购"],
        ["分层", "利润最大化，匹配不同支付意愿", "2026：MiniMax极速版¥98-899，Kimi Allegro ¥559"],
    ],
    col_widths=[70, 130, 200]
)
add_mixed_para([
    ("可复用模式：", True, 10.5, ACCENT),
    ("进入新市场时，不要一开始就想好定价。先补贴跑量，再收紧测试弹性，最后精细分层。", False, 10.5, BODY),
], after=12)

doc.add_heading("3.3 从「入口」到「生态」的升级路径", level=2)
add_para("Coding是获客入口，Token是生态载体。MiniMax的Token Plan不只卖编程，打包了图片、语音、音乐、视频。从「编程工具」升级为「AI多媒体工作室」。", after=6)
add_mixed_para([
    ("可复用模式：", True, 10.5, ACCENT),
    ("用高频刚需场景（编程）获客，用低频高价值场景（视频/音乐）提高ARPU。", False, 10.5, BODY),
], after=12)

doc.add_heading("3.4 国内外定价逻辑的根本差异", level=2)
add_table(
    ["维度", "国内", "海外"],
    [
        ["竞争环境", "7家在同一池子打价格战", "面对Claude/GPT/Gemini三巨头"],
        ["定价策略", "低价抢量，比谁更便宜", '"证明"我不比Claude差"，比谁更强"'],
        ["用户预期", '"你在国内买到的便宜，是因为有人在替你卷"', "为质量和生态付费"],
    ],
    col_widths=[70, 190, 190]
)
add_spacer()

doc.add_heading("3.5 Agent时代的成本结构剧变", level=2)
add_numbered_list([
    ("70% coding agent token是浪费（DEV Community）", False, 10.5, BODY),
    ("一个开发者8个月消耗100亿Token", False, 10.5, BODY),
    ("Cursor 2025-06转credit-based，重度用户一周超支$350", False, 10.5, BODY),
    ("订阅费是地板不是天花板，实际账单2-5倍", False, 10.5, BODY),
])

doc.add_heading("3.6 跨越鸿沟：AI从实验到生产的信任挑战", level=2)
add_para("Gartner：90%企业推进AI实验，仅41%成功落地生产环境。预计2027年40% Agentic AI项目因无法证明ROI被终止。", after=6)
add_mixed_para([
    ("对定价的含义：", True, 10.5, ACCENT),
    ("企业在实验阶段愿意尝试低价Coding Plan → 进入生产环境后需要SLA、稳定性、合规性 → Token Plan企业版溢价的基础。跨越鸿沟的关键不是模型能力，而是「可信赖度」。", False, 10.5, BODY),
], after=12)

add_page_break()

# ===== PART 4: CHECKLIST =====
doc.add_heading("四、定价策略实战清单", level=1)

doc.add_heading("4.1 定价策略决策清单", level=2)
add_numbered_list([
    ("定义计量单位：用什么计费？（Token/请求/次数/credit）", True, 10, HEADING),
    ("选择获客策略：补贴还是差异化？", True, 10, HEADING),
    ("设计分层结构：几档？价差多少？", True, 10, HEADING),
    ("确定透明度：额度描述要多透明？", True, 10, HEADING),
    ("规划演进路径：从什么到什么？", True, 10, HEADING),
])
add_spacer()

doc.add_heading("4.2 商业模式转型检查点", level=2)
add_numbered_list([
    ("用户成本可预测性：用户能不能算清楚自己要花多少？", True, 10, HEADING),
    ("利润弹性：收紧定价后用户留存如何？", True, 10, HEADING),
    ("生态粘性：从单一场景到多场景的升级是否自然？", True, 10, HEADING),
    ("竞争护城河：计量单位/定价结构是否容易被复制？", True, 10, HEADING),
])
add_spacer()

doc.add_heading("4.3 反模式（踩过的坑）", level=2)
add_numbered_list([
    ("Cursor 2025-06退款风波：突然改变计费模式，用户反弹", True, 10, HEADING),
    ("计量单位混乱：90%用户踩坑，信息不对称最终伤害信任", True, 10, HEADING),
    ("饥饿营销反噬：每天10:30抢购，短期有效，长期伤品牌", True, 10, HEADING),
])

add_page_break()

# ===== PART 5: FUTURE =====
doc.add_heading("五、未来预判", level=1)

doc.add_heading("5.1 Token成本将持续下降", level=2)
add_para("驱动力：国产芯片放量（华为950年内75-100万片）+ DeepSeek V4开源MoE效率革命 + 推理专用芯片竞争。Token Plan价格战继续，利润空间来自「服务溢价」而非「算力差价」。", after=12)

doc.add_heading("5.2 国产芯片替代加速改变供给格局", level=2)
add_para("英伟达市占率持续下降（66%→54%→2026年可能低于50%）。华为昇腾+寒武纪+阿里平头哥形成国产算力三角。算力供给不再被卡脖子 → 国内厂商成本结构和海外产生分化 → 定价策略可以更激进。", after=12)

doc.add_heading("5.3 开源模型重塑定价权", level=2)
add_para('DeepSeek模式：开源+极致低价+不卖订阅。V3.2在OpenRouter月调用5.35万亿token，全球第三。开源模型的免费部署能力持续压制定价上限。厂商定价权来自「服务便利性」而非「模型独占性」。', after=12)

doc.add_heading("5.4 Agent时代的成本指数级增长不可逆", level=2)
add_para("推理需求2026年已达训练的4-5倍，IDC预测2028年推理占比73%。Coding Plan「不限量」模式不可持续。所有厂商最终转向Token精细计量——谁先转谁主动，谁后转谁被动。", after=12)

add_page_break()

# ===== CONCLUSION =====
doc.add_heading("结论：从Coding Plan到Token Plan的本质", level=1)

add_mixed_para([
    ("Coding Plan → Token Plan的转变，本质上是从「卖工具使用权」到「卖算力资源包」的商业模式升级。", False, 10.5, HEADING),
], after=12)

add_para("Coding Plan ≈ 手机合约机（绑定特定工具/场景）", indent=20, after=6)
add_para("Token Plan ≈ 手机流量包（通用算力资源，不限场景）", indent=20, after=12)

add_para('MiniMax改名的逻辑：不只卖编程能力，卖的是全模态AI能力的「算力通票」。这不是一个命名变化，而是一次定位升级。', after=12)

add_spacer()
doc.add_heading("待解决的信息缺口", level=2)
add_numbered_list([
    ("各平台Coding Plan/Token Plan的实际利润率（公开数据缺失）", False, 10.5, BODY),
    ("首月优惠后的续费率（判断补贴效果）", False, 10.5, BODY),
    ("小米MiMo的真实Token消耗数据和成本结构", False, 10.5, BODY),
])

doc.save(OUTPUT)
print(f"Saved to {OUTPUT}")
