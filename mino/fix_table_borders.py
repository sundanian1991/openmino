from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_element = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_element.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_element)

    tcPr.append(tcBorders)

def add_table_borders(doc_path, output_path):
    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell,
                    top={'sz': '4', 'val': 'single', 'color': '000000'},
                    left={'sz': '4', 'val': 'single', 'color': '000000'},
                    bottom={'sz': '4', 'val': 'single', 'color': '000000'},
                    right={'sz': '4', 'val': 'single', 'color': '000000'}
                )

    doc.save(output_path)
    print(f"Updated: {output_path}")

# 处理所有Word文档
import os
base_dir = "/Users/sundanian/Documents/projects/ai-agents/my-agent/mino/企业微信BPO风险管控体系/正式文档"

for file in os.listdir(base_dir):
    if file.endswith('.docx'):
        input_path = os.path.join(base_dir, file)
        output_path = os.path.join(base_dir, f"bordered_{file}")
        try:
            add_table_borders(input_path, output_path)
        except Exception as e:
            print(f"Error processing {file}: {e}")
