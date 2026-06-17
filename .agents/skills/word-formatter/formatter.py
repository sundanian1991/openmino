#!/usr/bin/env python3
"""
Word Formatter - 文档自动排版工具
将杂乱的文档转换为专业排版的Word/PDF文件
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any

# Document processing
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# Charts
import matplotlib.pyplot as plt
import matplotlib
matplotlib.rcParams['font.sans-serif'] = ['PingFang SC', 'Microsoft YaHei', 'SimHei']
matplotlib.rcParams['axes.unicode_minus'] = False

# ==============================================================================
# Content Type Detection
# ==============================================================================

CONTENT_TYPES = {
    "meeting_notes": {
        "keywords": ["会议", "议程", "参会", "决议", "行动项", "纪要", "meeting", "agenda", "minutes"],
        "style": "business_formal",
        "description": "会议纪要"
    },
    "report": {
        "keywords": ["报告", "总结", "分析", "数据", "季度", "年度", "report", "analysis", "summary"],
        "style": "business_formal",
        "description": "工作报告"
    },
    "study_notes": {
        "keywords": ["笔记", "学习", "知识点", "概念", "定义", "notes", "learning", "study"],
        "style": "modern_minimal",
        "description": "学习笔记"
    },
    "proposal": {
        "keywords": ["方案", "计划", "目标", "预算", "规划", "proposal", "plan", "budget"],
        "style": "academic",
        "description": "方案文档"
    }
}


def detect_content_type(text: str) -> Tuple[str, str, str]:
    """
    检测文档类型
    Returns: (type_key, style, description)
    """
    text_lower = text.lower()
    scores = {}

    for type_key, config in CONTENT_TYPES.items():
        score = sum(1 for kw in config["keywords"] if kw in text_lower)
        scores[type_key] = score

    if max(scores.values()) > 0:
        best_type = max(scores, key=scores.get)
        config = CONTENT_TYPES[best_type]
        return best_type, config["style"], config["description"]

    return "report", "business_formal", "通用文档"


# ==============================================================================
# Style Definitions
# ==============================================================================

STYLES = {
    "business_formal": {
        "name": "商务正式",
        "title_font": "Microsoft YaHei",
        "title_size": 22,
        "heading1_size": 16,
        "heading2_size": 14,
        "body_font": "Microsoft YaHei",
        "body_size": 11,
        "line_spacing": 1.5,
        "colors": {
            "primary": "#1a365d",
            "accent": "#2b6cb0",
            "text": "#2d3748",
            "light": "#e2e8f0"
        }
    },
    "modern_minimal": {
        "name": "简约现代",
        "title_font": "PingFang SC",
        "title_size": 24,
        "heading1_size": 18,
        "heading2_size": 14,
        "body_font": "PingFang SC",
        "body_size": 11,
        "line_spacing": 1.8,
        "colors": {
            "primary": "#1a202c",
            "accent": "#4a5568",
            "text": "#2d3748",
            "light": "#f7fafc"
        }
    },
    "academic": {
        "name": "学术风格",
        "title_font": "SimSun",
        "title_size": 18,
        "heading1_size": 15,
        "heading2_size": 13,
        "body_font": "SimSun",
        "body_size": 12,
        "line_spacing": 1.5,
        "colors": {
            "primary": "#1a202c",
            "accent": "#4a5568",
            "text": "#000000",
            "light": "#f0f0f0"
        }
    }
}


def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


# ==============================================================================
# Document Analysis
# ==============================================================================

def analyze_document(file_path: str) -> Dict[str, Any]:
    """
    分析文档内容
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    # Read content based on file type
    if path.suffix.lower() == '.docx':
        content = read_docx(file_path)
    elif path.suffix.lower() in ['.txt', '.md']:
        content = read_text(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {path.suffix}")

    # Detect type and style
    type_key, style_key, description = detect_content_type(content)

    # Count words
    word_count = len(content)

    # Detect data patterns
    has_data = detect_data_patterns(content)

    return {
        "file_path": file_path,
        "content": content,
        "type": type_key,
        "type_description": description,
        "style": style_key,
        "style_name": STYLES[style_key]["name"],
        "word_count": word_count,
        "has_data": has_data
    }


def read_docx(file_path: str) -> str:
    """Read content from Word document"""
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return '\n'.join(paragraphs)


def read_text(file_path: str) -> str:
    """Read content from text file"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def detect_data_patterns(text: str) -> List[Dict]:
    """
    检测文本中的数据模式
    """
    data_patterns = []

    # Percentage patterns: 30%, 增长30%
    percentage_pattern = r'(\d+(?:\.\d+)?)\s*%'
    percentages = re.findall(percentage_pattern, text)
    if percentages:
        data_patterns.append({
            "type": "percentage",
            "values": percentages,
            "suggested_chart": "pie" if len(percentages) >= 3 else "bar"
        })

    # Number comparisons: 从X到Y, X vs Y
    comparison_pattern = r'从\s*(\d+)\s*到\s*(\d+)|(\d+)\s*(?:vs|VS|对比)\s*(\d+)'
    comparisons = re.findall(comparison_pattern, text)
    if comparisons:
        data_patterns.append({
            "type": "comparison",
            "values": comparisons,
            "suggested_chart": "bar"
        })

    # Time series: Q1, Q2, 一月, 二月
    time_pattern = r'(Q[1-4]|[一二三四]季度|[一二三四五六七八九十]+月|\d{4}年)'
    time_refs = re.findall(time_pattern, text)
    if time_refs:
        data_patterns.append({
            "type": "time_series",
            "values": time_refs,
            "suggested_chart": "line"
        })

    return data_patterns


# ==============================================================================
# Document Structuring
# ==============================================================================

def structure_content(content: str, level: str) -> List[Dict]:
    """
    根据处理级别结构化内容
    level: format_only, light_restructure, deep_restructure
    """
    elements = []
    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Detect element type
        if is_title(line):
            elements.append({"type": "title", "text": line})
        elif is_heading1(line):
            elements.append({"type": "heading1", "text": clean_heading(line)})
        elif is_heading2(line):
            elements.append({"type": "heading2", "text": clean_heading(line)})
        elif is_list_item(line):
            elements.append({"type": "list_item", "text": clean_list_item(line)})
        else:
            elements.append({"type": "paragraph", "text": line})

    if level == "light_restructure":
        elements = add_subheadings(elements)
        elements = add_transitions(elements)
    elif level == "deep_restructure":
        elements = add_subheadings(elements)
        elements = add_transitions(elements)
        elements = improve_clarity(elements)

    return elements


def is_title(line: str) -> bool:
    """Check if line is a title"""
    return line.startswith('# ') or (len(line) < 50 and not line.endswith('。'))


def is_heading1(line: str) -> bool:
    """Check if line is a level 1 heading"""
    patterns = [
        r'^##\s+',
        r'^[一二三四五六七八九十]+[、.]\s*',
        r'^\d+[、.]\s+\S',
    ]
    return any(re.match(p, line) for p in patterns)


def is_heading2(line: str) -> bool:
    """Check if line is a level 2 heading"""
    patterns = [
        r'^###\s+',
        r'^[（(][一二三四五六七八九十0-9]+[）)]\s*',
        r'^\d+\.\d+\s+',
    ]
    return any(re.match(p, line) for p in patterns)


def is_list_item(line: str) -> bool:
    """Check if line is a list item"""
    patterns = [
        r'^[-•*]\s+',
        r'^\d+[.)]\s+',
    ]
    return any(re.match(p, line) for p in patterns)


def clean_heading(line: str) -> str:
    """Clean heading markers"""
    line = re.sub(r'^#+\s+', '', line)
    line = re.sub(r'^[一二三四五六七八九十]+[、.]\s*', '', line)
    line = re.sub(r'^\d+[、.]\s+', '', line)
    return line.strip()


def clean_list_item(line: str) -> str:
    """Clean list item markers"""
    line = re.sub(r'^[-•*]\s+', '', line)
    line = re.sub(r'^\d+[.)]\s+', '', line)
    return line.strip()


def add_subheadings(elements: List[Dict]) -> List[Dict]:
    """Add subheadings to long sections"""
    result = []
    paragraph_count = 0

    for elem in elements:
        if elem["type"] in ["title", "heading1", "heading2"]:
            paragraph_count = 0
            result.append(elem)
        elif elem["type"] == "paragraph":
            paragraph_count += 1
            if paragraph_count > 3 and len(elem["text"]) > 100:
                # Try to extract a subheading from the paragraph
                first_sentence = elem["text"].split('。')[0]
                if len(first_sentence) < 30:
                    result.append({"type": "heading2", "text": first_sentence})
                    remaining = '。'.join(elem["text"].split('。')[1:])
                    if remaining:
                        result.append({"type": "paragraph", "text": remaining})
                    paragraph_count = 0
                    continue
            result.append(elem)
        else:
            result.append(elem)

    return result


def add_transitions(elements: List[Dict]) -> List[Dict]:
    """Add transition phrases between sections"""
    transitions = [
        "接下来，",
        "此外，",
        "具体来说，",
        "在此基础上，",
    ]

    result = []
    trans_idx = 0
    prev_was_heading = False

    for elem in elements:
        if elem["type"] in ["heading1", "heading2"]:
            prev_was_heading = True
            result.append(elem)
        elif elem["type"] == "paragraph" and prev_was_heading:
            prev_was_heading = False
            # Don't add transition if paragraph already has one
            if not any(elem["text"].startswith(t) for t in transitions):
                if not elem["text"][0] in "首先第一":
                    elem["text"] = transitions[trans_idx % len(transitions)] + elem["text"]
                    trans_idx += 1
            result.append(elem)
        else:
            prev_was_heading = False
            result.append(elem)

    return result


def improve_clarity(elements: List[Dict]) -> List[Dict]:
    """Improve clarity of paragraphs (placeholder for AI enhancement)"""
    # This would be enhanced by Claude during execution
    return elements


# ==============================================================================
# Chart Generation
# ==============================================================================

def generate_chart(data: Dict, chart_type: str, style: Dict, output_path: str) -> str:
    """
    Generate a chart and save to file
    """
    colors = style["colors"]
    primary = hex_to_rgb(colors["primary"])
    accent = hex_to_rgb(colors["accent"])

    fig, ax = plt.subplots(figsize=(10, 6))

    if chart_type == "bar":
        bars = ax.bar(
            data["labels"],
            data["values"],
            color=[f'#{colors["primary"][1:]}', f'#{colors["accent"][1:]}'] * 5
        )
        ax.set_ylabel(data.get("ylabel", ""))

    elif chart_type == "pie":
        colors_list = [colors["primary"], colors["accent"], colors["light"]]
        ax.pie(
            data["values"],
            labels=data["labels"],
            colors=colors_list[:len(data["values"])],
            autopct='%1.1f%%',
            startangle=90
        )

    elif chart_type == "line":
        ax.plot(
            data["x"],
            data["y"],
            color=colors["primary"],
            linewidth=2,
            marker='o'
        )
        ax.set_xlabel(data.get("xlabel", ""))
        ax.set_ylabel(data.get("ylabel", ""))
        ax.grid(True, alpha=0.3)

    ax.set_title(data.get("title", ""), fontsize=14, fontweight='bold')

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()

    return output_path


# ==============================================================================
# Word Document Generation
# ==============================================================================

def create_formatted_document(
    elements: List[Dict],
    style_key: str,
    images: Optional[List[str]] = None,
    output_path: str = None
) -> str:
    """
    Create a formatted Word document
    """
    style = STYLES[style_key]
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Add content
    for i, elem in enumerate(elements):
        if elem["type"] == "title":
            p = doc.add_heading(elem["text"], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif elem["type"] == "heading1":
            p = doc.add_heading(elem["text"], level=1)

        elif elem["type"] == "heading2":
            p = doc.add_heading(elem["text"], level=2)

        elif elem["type"] == "paragraph":
            p = doc.add_paragraph(elem["text"])
            p.paragraph_format.line_spacing = style["line_spacing"]
            p.paragraph_format.space_after = Pt(12)

        elif elem["type"] == "list_item":
            p = doc.add_paragraph(elem["text"], style='List Bullet')

        elif elem["type"] == "image" and images:
            img_path = elem.get("path")
            if img_path and os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5))
                # Center the image
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add cover image at the beginning if provided
    if images and len(images) > 0 and os.path.exists(images[0]):
        # Insert at beginning
        first_para = doc.paragraphs[0]
        run = first_para.insert_paragraph_before()
        # Add image to the new paragraph
        doc.paragraphs[0].add_run().add_picture(images[0], width=Inches(6))
        doc.add_page_break()

    # Save document
    if output_path is None:
        output_path = "formatted_document.docx"

    doc.save(output_path)
    return output_path


def convert_to_pdf(docx_path: str) -> Optional[str]:
    """
    Convert Word document to PDF
    """
    pdf_path = docx_path.replace('.docx', '.pdf')

    try:
        # Try docx2pdf first (requires MS Word on Mac/Windows)
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return pdf_path
    except ImportError:
        pass
    except Exception:
        pass

    try:
        # Try LibreOffice
        import subprocess
        output_dir = os.path.dirname(docx_path) or '.'
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', output_dir, docx_path
        ], capture_output=True, timeout=60)

        if result.returncode == 0 and os.path.exists(pdf_path):
            return pdf_path
    except FileNotFoundError:
        pass
    except Exception:
        pass

    return None


# ==============================================================================
# Main Entry Point
# ==============================================================================

def format_document(
    input_path: str,
    level: str = "format_only",
    output_dir: Optional[str] = None
) -> Dict[str, str]:
    """
    Main function to format a document

    Args:
        input_path: Path to input document
        level: Processing level (format_only, light_restructure, deep_restructure)
        output_dir: Output directory (default: same as input)

    Returns:
        Dict with paths to output files
    """
    # Analyze document
    analysis = analyze_document(input_path)

    # Structure content
    elements = structure_content(analysis["content"], level)

    # Prepare output paths
    input_file = Path(input_path)
    if output_dir is None:
        output_dir = input_file.parent

    base_name = input_file.stem
    docx_output = Path(output_dir) / f"{base_name}_formatted.docx"

    # Create formatted document
    style_key = analysis["style"]
    docx_path = create_formatted_document(
        elements,
        style_key,
        output_path=str(docx_output)
    )

    # Convert to PDF
    pdf_path = convert_to_pdf(docx_path)

    return {
        "docx": docx_path,
        "pdf": pdf_path,
        "analysis": analysis
    }


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python formatter.py <input_file> [level]")
        print("  level: format_only, light_restructure, deep_restructure")
        sys.exit(1)

    input_file = sys.argv[1]
    level = sys.argv[2] if len(sys.argv) > 2 else "format_only"

    result = format_document(input_file, level)

    print(f"Word文档: {result['docx']}")
    if result['pdf']:
        print(f"PDF文档: {result['pdf']}")
    else:
        print("PDF转换失败，请安装 docx2pdf 或 LibreOffice")
