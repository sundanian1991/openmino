#!/usr/bin/env python3
"""
Word Proofreader - 错别字修正、润色、排版工具
支持 Track Changes (修订模式) 记录每处文字修改
"""

import json
import sys
import os
import copy
import difflib
from datetime import datetime, timezone

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误：请先安装 python-docx")
    print("  pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Revision ID counter (Word requires unique IDs for each tracked change)
# ---------------------------------------------------------------------------
_rev_id = 100


def _next_rev_id():
    global _rev_id
    _rev_id += 1
    return str(_rev_id)


# ---------------------------------------------------------------------------
# Extract: Read document paragraphs as JSON
# ---------------------------------------------------------------------------
def extract_text(docx_path):
    """Extract text from all paragraphs in a Word document."""
    doc = Document(docx_path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if text.strip():
            paragraphs.append({
                "index": i,
                "text": text,
                "style": para.style.name if para.style else "Normal",
            })
    return {"file": os.path.abspath(docx_path), "paragraph_count": len(doc.paragraphs), "paragraphs": paragraphs}


# ---------------------------------------------------------------------------
# Diff helpers
# ---------------------------------------------------------------------------
def _compute_diff_ops(old_text, new_text):
    """Compute character-level diff operations between two strings."""
    matcher = difflib.SequenceMatcher(None, old_text, new_text, autojunk=False)
    ops = []
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == "equal":
            ops.append(("equal", old_text[i1:i2]))
        elif op == "replace":
            ops.append(("delete", old_text[i1:i2]))
            ops.append(("insert", new_text[j1:j2]))
        elif op == "insert":
            ops.append(("insert", new_text[j1:j2]))
        elif op == "delete":
            ops.append(("delete", old_text[i1:i2]))
    return ops


# ---------------------------------------------------------------------------
# XML element builders for Track Changes
# ---------------------------------------------------------------------------
def _make_run(text, rPr=None):
    """Create a normal w:r element."""
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(copy.deepcopy(rPr))
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _make_del_run(text, rPr=None):
    """Create a run with w:delText (for tracked deletions)."""
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(copy.deepcopy(rPr))
    dt = OxmlElement("w:delText")
    dt.text = text
    dt.set(qn("xml:space"), "preserve")
    r.append(dt)
    return r


def _wrap_del(child, author="Claude Proofreader"):
    """Wrap an element in a w:del tracked-change container."""
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    el = OxmlElement("w:del")
    el.set(qn("w:id"), _next_rev_id())
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), now)
    el.append(child)
    return el


def _wrap_ins(child, author="Claude Proofreader"):
    """Wrap an element in a w:ins tracked-change container."""
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    el = OxmlElement("w:ins")
    el.set(qn("w:id"), _next_rev_id())
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), now)
    el.append(child)
    return el


# ---------------------------------------------------------------------------
# Apply tracked changes to a single paragraph
# ---------------------------------------------------------------------------
def _apply_tracked_changes(para_element, new_text, author="Claude Proofreader"):
    """
    Replace paragraph text content with granular tracked changes.
    Uses character-level diff to mark precise deletions and insertions.
    Returns True if any changes were applied.
    """
    # Collect current text from all runs
    old_text = ""
    for r in para_element.findall(qn("w:r")):
        for t in r.findall(qn("w:t")):
            if t.text:
                old_text += t.text

    if old_text == new_text:
        return False

    # Grab formatting from the first run
    rPr = None
    first_run = para_element.find(qn("w:r"))
    if first_run is not None:
        existing = first_run.find(qn("w:rPr"))
        if existing is not None:
            rPr = copy.deepcopy(existing)

    # Remove existing runs (but keep pPr, bookmarks, etc.)
    removable_tags = {"r", "ins", "del", "smartTag", "hyperlink"}
    for child in list(para_element):
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local in removable_tags:
            para_element.remove(child)

    # Compute diff and rebuild
    ops = _compute_diff_ops(old_text, new_text)

    pPr = para_element.find(qn("w:pPr"))
    insert_after = pPr  # insert after pPr, or at beginning if None

    for op_type, text in ops:
        if not text:
            continue

        if op_type == "equal":
            elem = _make_run(text, rPr)
        elif op_type == "delete":
            elem = _wrap_del(_make_del_run(text, rPr), author)
        elif op_type == "insert":
            elem = _wrap_ins(_make_run(text, rPr), author)
        else:
            continue

        if insert_after is not None:
            insert_after.addnext(elem)
        else:
            para_element.insert(0, elem)
        insert_after = elem

    return True


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------
def _apply_formatting(doc, config):
    """Apply document-wide formatting."""
    defaults = {
        "body_font_cn": "SimSun",
        "body_font_en": "Times New Roman",
        "body_size": 12,
        "heading1_size": 22,
        "heading2_size": 16,
        "heading3_size": 14,
        "line_spacing": 1.5,
        "space_after": 6,
    }
    cfg = {**defaults, **(config or {})}

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        # Line / paragraph spacing
        pf = para.paragraph_format
        pf.line_spacing = cfg["line_spacing"]
        pf.space_after = Pt(cfg["space_after"])

        for run in para.runs:
            # Font size by heading level
            if "Heading 1" in style_name:
                run.font.size = Pt(cfg["heading1_size"])
            elif "Heading 2" in style_name:
                run.font.size = Pt(cfg["heading2_size"])
            elif "Heading 3" in style_name:
                run.font.size = Pt(cfg["heading3_size"])
            else:
                run.font.size = Pt(cfg["body_size"])

            # Latin font
            run.font.name = cfg["body_font_en"]

            # East-Asian font
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), cfg["body_font_cn"])

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)


# ---------------------------------------------------------------------------
# Apply: Main correction entry point
# ---------------------------------------------------------------------------
def apply_corrections(docx_path, corrections, output_path, mode="full", format_config=None):
    """
    Apply corrections to a Word document with Track Changes.

    corrections: list of {"index": int, "corrected": str}
    mode: "typo_only" | "typo_polish" | "full"
    """
    doc = Document(docx_path)

    correction_map = {c["index"]: c["corrected"] for c in corrections}

    change_count = 0
    for i, para in enumerate(doc.paragraphs):
        if i in correction_map:
            changed = _apply_tracked_changes(para._element, correction_map[i])
            if changed:
                change_count += 1

    if mode == "full":
        _apply_formatting(doc, format_config)

    doc.save(output_path)

    return {
        "output": os.path.abspath(output_path),
        "paragraphs_changed": change_count,
        "total_paragraphs": len(doc.paragraphs),
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    if len(sys.argv) < 3:
        print("Word Proofreader - 错别字修正、润色、排版工具")
        print()
        print("用法:")
        print("  python proofreader.py extract <docx_path>")
        print("  python proofreader.py apply <docx_path> <corrections.json> [options]")
        print()
        print("选项:")
        print("  --output <path>    输出文件路径 (默认: *_proofread.docx)")
        print("  --mode <mode>      typo_only | typo_polish | full (默认: full)")
        print("  --format <json>    自定义排版配置 JSON 文件")
        sys.exit(1)

    command = sys.argv[1]

    if command == "extract":
        docx_path = sys.argv[2]
        if not os.path.exists(docx_path):
            print(f"错误：文件不存在: {docx_path}", file=sys.stderr)
            sys.exit(1)
        result = extract_text(docx_path)
        print(json.dumps(result, ensure_ascii=False, indent=2))

    elif command == "apply":
        if len(sys.argv) < 4:
            print("错误：缺少 corrections.json 参数", file=sys.stderr)
            sys.exit(1)

        docx_path = sys.argv[2]
        corrections_path = sys.argv[3]

        if not os.path.exists(docx_path):
            print(f"错误：文件不存在: {docx_path}", file=sys.stderr)
            sys.exit(1)
        if not os.path.exists(corrections_path):
            print(f"错误：文件不存在: {corrections_path}", file=sys.stderr)
            sys.exit(1)

        # Parse optional arguments
        output_path = None
        mode = "full"
        format_config = None

        i = 4
        while i < len(sys.argv):
            if sys.argv[i] == "--output" and i + 1 < len(sys.argv):
                output_path = sys.argv[i + 1]
                i += 2
            elif sys.argv[i] == "--mode" and i + 1 < len(sys.argv):
                mode = sys.argv[i + 1]
                i += 2
            elif sys.argv[i] == "--format" and i + 1 < len(sys.argv):
                with open(sys.argv[i + 1], "r", encoding="utf-8") as f:
                    format_config = json.load(f)
                i += 2
            else:
                i += 1

        if output_path is None:
            base, ext = os.path.splitext(docx_path)
            output_path = f"{base}_proofread{ext}"

        # Load corrections
        with open(corrections_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        corrections = data if isinstance(data, list) else data.get("paragraphs", [])

        result = apply_corrections(docx_path, corrections, output_path, mode, format_config)
        print(json.dumps(result, ensure_ascii=False, indent=2))

    else:
        print(f"未知命令: {command}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
